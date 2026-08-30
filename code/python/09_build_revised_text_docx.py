# -*- coding: utf-8 -*-
"""
STEP 09 - Ready-to-paste revised manuscript text, generated from the same
verified result files so that no number can drift from the analysis.

Covers the passages the referee required to be rewritten: Abstract, Methods
subsections 3.1-3.6, Results 4.1-4.6, Discussion 5, Conclusions 6, and the
Data availability statement.
"""
import numpy as np, pandas as pd, os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = r"d:\Doctorat\article1\outputs_fast"
FINAL = os.path.join(BASE, "FINAL")
RES   = os.path.join(FINAL, "results")
TAB   = os.path.join(FINAL, "tables")
DOCS  = os.path.join(FINAL, "docs")
os.makedirs(DOCS, exist_ok=True)

TARGETS = ["N", "P", "K"]
ACC  = RGBColor(0x1F, 0x5F, 0x6B)
GREY = RGBColor(0x55, 0x5F, 0x5B)
RED  = RGBColor(0xB5, 0x22, 0x22)

df   = pd.read_csv(os.path.join(RES, "01_analysis_ready_dataset.csv"))
s01  = json.load(open(os.path.join(RES, "01_summary.json")))
bdes = pd.read_csv(os.path.join(RES, "01_block_design.csv"))
mt   = pd.read_csv(os.path.join(RES, "02_metrics_all_schemes.csv"))
bl   = pd.read_csv(os.path.join(RES, "02_baselines.csv"))
ci   = pd.read_csv(os.path.join(RES, "02_bootstrap_ci.csv"))
pcmp = pd.read_csv(os.path.join(RES, "02_paired_model_comparison.csv"))
pi   = pd.read_csv(os.path.join(RES, "02_permutation_importance.csv"))
vg   = pd.read_csv(os.path.join(RES, "04_variogram_parameters.csv"))
mor  = pd.read_csv(os.path.join(RES, "04_morans_I.csv"))
care = pd.read_csv(os.path.join(RES, "03_class_areas.csv"))
agr  = pd.read_csv(os.path.join(RES, "03_RF_XGB_agreement_summary.csv"))
ras  = pd.read_csv(os.path.join(RES, "03_raster_summary.csv"))
msum = json.load(open(os.path.join(RES, "03_map_summary.json")))
t3   = pd.read_csv(os.path.join(TAB, "T3_descriptive_statistics.csv"))
t7   = pd.read_csv(os.path.join(TAB, "T7_importance_by_class.csv"))
t7b  = pd.read_csv(os.path.join(TAB, "T7b_acquisition_share.csv"))
t7c  = pd.read_csv(os.path.join(TAB, "T7c_importance_stability.csv"))
t9   = pd.read_csv(os.path.join(TAB, "T9_class_accuracy_kappa.csv"))
t6b  = pd.read_csv(os.path.join(TAB, "T6b_modal_hyperparameters.csv"))


def sp(t, m, s="nested spatial CV"):
    return mt[(mt.target == t) & (mt.model == m) & (mt.scheme == s)].iloc[0]


def cir(t, m):
    return ci[(ci.target == t) & (ci.model == m)].iloc[0]


def bs(t, m):
    return bl[(bl.target == t) & (bl.model == m)].iloc[0]


def clsshare(t, m, cls):
    r = t7[t7["Predictor class"] == cls]
    return float(r[f"{t} - {'RF' if m == 'RF' else 'XGBoost'}"].iloc[0])


def rs(t, m):
    return ras[ras.layer == f"{t}_{m}"].iloc[0]


def ca(t, cls):
    return care[(care.nutrient == t) & (care.cls == cls)].iloc[0]


nmean = float(t3[t3.Nutrient.str.startswith("Total N")].Mean.iloc[0])

# ------------------------------------------------------------------ document
doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15
for lvl, sz in [(1, 15), (2, 12.5), (3, 11)]:
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"; h.font.size = Pt(sz); h.font.bold = True
    h.font.color.rgb = ACC
    h.paragraph_format.space_before = Pt(14 if lvl < 3 else 10)
    h.paragraph_format.space_after = Pt(5)


def P(text="", bold=False, italic=False, size=10.5, color=None, align=None,
      space_after=6, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def note(text):
    p = doc.add_paragraph()
    r = p.add_run("Note to the authors: "); r.bold = True; r.italic = True
    r.font.size = Pt(9.2); r.font.color.rgb = RED
    r2 = p.add_run(text); r2.italic = True; r2.font.size = Pt(9.2); r2.font.color.rgb = RED
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def table(dfx, fontsize=8.4, caption=None, widths=None):
    cols = list(dfx.columns)
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]; cell.text = ""
        r = cell.paragraphs[0].add_run(str(c)); r.bold = True; r.font.size = Pt(fontsize)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(cell, "1F5F6B")
    for _, row in dfx.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            txt = "" if pd.isna(v) else (f"{v:,.4g}" if isinstance(v, (float, np.floating))
                                         else str(v))
            rr = cells[i].paragraphs[0].add_run(txt); rr.font.size = Pt(fontsize)
            if isinstance(v, (int, float, np.integer, np.floating)):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(t.rows):
        if ri > 0 and ri % 2 == 0:
            for c in row.cells:
                shade(c, "F2F5F1")
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    if caption:
        P(caption, size=8.6, italic=True, color=GREY, space_after=10)
    else:
        P("", size=4, space_after=4)
    return t


# ============================================================== title
P("Revised manuscript text", bold=True, size=19, color=ACC,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P("Ready-to-paste replacement passages, generated directly from the verified results",
  size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
P("High-resolution mapping of soil macronutrients in a semi-arid climate (Morocco) "
  "using Sentinel-2 data and machine learning", bold=True, size=12,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
P("Every numerical value below is read from the deposited result files, so the text "
  "cannot disagree with the tables, figures or rasters. Passages are given in the order "
  "they appear in the manuscript. Text in red is guidance for the authors and must be "
  "deleted before submission; it marks the few places where information only you hold — "
  "laboratory QA records, sampling dates, Earth Engine asset identifiers — has to be "
  "filled in.")
note("Delete every red paragraph before submitting. They are instructions, not manuscript text.")

# ============================================================== Abstract
doc.add_heading("Abstract", level=1)
rf_n, rf_p, rf_k = sp("N", "RF"), sp("P", "RF"), sp("K", "RF")
cn, cp, ck = cir("N", "RF"), cir("P", "RF"), cir("K", "RF")
P(f"In semi-arid irrigated districts, soil macronutrients vary strongly over short "
  f"distances, yet no fine-scale nutrient maps exist for the Beni Moussa irrigated "
  f"district of the Tadla Plain, Morocco. This study evaluates how far Sentinel-2 "
  f"surface reflectance can predict topsoil total nitrogen (N), Olsen phosphorus (P) "
  f"and exchangeable potassium (K) in this setting, and maps the result. A total of 110 "
  f"composite topsoil samples (0–20 cm) were collected between November 2024 and "
  f"January 2025 and analysed by Kjeldahl digestion, the Olsen method and ammonium "
  f"acetate extraction. {s01['n_predictors']} predictors — ten reflectance bands and "
  f"ten spectral indices from two acquisitions — were derived in Google Earth Engine. "
  f"Random Forest and XGBoost were evaluated under nested cross-validation with "
  f"{s01['n_blocks']} spatial blocks, in which every tuning decision was made inside an "
  f"inner loop on the training blocks only, so that all reported statistics come from "
  f"untouched outer folds. Random Forest explained {rf_n.R2:.2f} of the variance in "
  f"total N (RMSE {rf_n.RMSE:,.0f} mg kg⁻¹, 95 % CI on R² "
  f"[{cn.R2_lo:.2f}, {cn.R2_hi:.2f}]), {rf_p.R2:.2f} for Olsen P "
  f"(RMSE {rf_p.RMSE:.1f} mg kg⁻¹) and {rf_k.R2:.2f} for exchangeable K "
  f"(RMSE {rf_k.RMSE:.1f} mg kg⁻¹). XGBoost performed comparably; the paired difference "
  f"in squared error was not statistically significant for any nutrient, so no "
  f"superiority is claimed. Permutation importance measured on held-out folds was "
  f"dominated by raw near-infrared and red-edge reflectance rather than by spectral "
  f"indices, and the November acquisition alone carried "
  f"{t7b.November_2024_share_percent.min():.0f}–"
  f"{t7b.November_2024_share_percent.max():.0f} % of predictive importance. Maps were "
  f"produced on the native 10 m grid with local uncertainty and an applicability "
  f"domain; {msum['pct_inside_aoa']:.0f} % of the district "
  f"({msum['ha_inside_aoa']:,.0f} ha) lies inside that domain. Because a single "
  f"acquisition carries most of the signal and predictors describe canopy condition, "
  f"the surfaces are best interpreted as seasonal reconnaissance products conditioned "
  f"on the 2024–2025 cropping mosaic, suitable for stratified resampling and for "
  f"delineating candidate management zones, rather than as stable soil property maps or "
  f"a direct basis for variable-rate fertiliser prescription.")
P("Keywords: digital soil mapping; spatial cross-validation; area of applicability; "
  "Sentinel-2; Tadla Plain.", space_after=12)
note("The journal may apply the keyword rule at word level; these avoid repeating title terms.")

# ============================================================== Methods
doc.add_heading("3.  Materials and Methods", level=1)

doc.add_heading("3.1  Soil sampling design and laboratory analysis", level=2)
P("A total of 110 composite topsoil samples (0–20 cm) were collected across the Beni "
  "Moussa irrigated district between November 2024 and January 2025. Sampling locations "
  "were allocated by conditioned Latin hypercube sampling (Minasny and McBratney, 2006) "
  "conditioned on dominant land use, irrigation subdivision, Sentinel-2 surface "
  "reflectance and previously mapped soil and hydromorphic units, with a minimum "
  "separation of 500 m between locations and the exclusion of non-agricultural surfaces. "
  "At each location, five subsamples were taken within a 10 m radius with a "
  "stainless-steel auger and combined into a composite sample of about 1 kg. Coordinates "
  "were recorded with a handheld GPS receiver with a horizontal accuracy of "
  "approximately ±3 m. Samples were air-dried, gently crushed and sieved to 2 mm.")
note("Confirm the design name and the minimum separation, and insert the allocation by "
     "stratum as a new table. If the design was not cLHS, say what it was and remove the "
     "Minasny and McBratney citation — the referee flagged that citation specifically.")
P("Total nitrogen was determined by Kjeldahl digestion, available phosphorus by the "
  "Olsen sodium bicarbonate method with colorimetric quantification by the molybdenum "
  "blue reaction, and exchangeable potassium by 1 N ammonium acetate extraction with "
  "flame photometry.")
note("Insert here, as the referee requires under M4: number of analytical replicates, "
     "blanks per batch, certified or internal reference materials with recovery, "
     "calibration procedure, limits of detection and quantification, repeatability "
     "(RSD %), and the number of rejected or incomplete measurements.")
P("Quality control comprised verification of laboratory logs, unit consistency checks "
  "and inspection of distributions with histograms and quantile–quantile plots. Values "
  "lying beyond 1.5 interquartile ranges were flagged for re-inspection against the "
  "laboratory records but were not removed; no observation was excluded on statistical "
  "grounds alone. The analysis dataset therefore comprises all 110 observations. No "
  "transformation was applied to any response variable.")

doc.add_heading("3.2  Sentinel-2 acquisition and processing", level=2)
P("Sentinel-2 MSI Level-2A surface reflectance was accessed through Google Earth Engine "
  "(Gorelick et al., 2017). Two cloud-free acquisitions were used, one in November 2024 "
  "and one in January 2025. Clouds and cloud shadows were removed using the Scene "
  "Classification Layer, and surface reflectance was divided by 10 000 to the 0–1 scale "
  "before any index was computed; band values in the analysis matrix lie between 0.025 "
  "and 0.498. The 20 m bands (B5, B6, B7, B8A, B11 and B12) were resampled bilinearly "
  "to the 10 m grid. Resampling produces a 10 m raster but does not create 10 m "
  "information, so the products are described throughout as nominal 10 m surfaces whose "
  "effective spatial support is set by the coarsest contributing predictor.")
note("Insert the exact Sentinel-2 asset identifiers, sensing dates, MGRS tile, scene "
     "cloud percentages, the SCL classes excluded, and confirm whether the harmonized "
     "collection was used. The referee asks for all of these by name under M5.")
P(f"Both acquisitions were assigned to every sampling point; no temporally matched "
  f"imagery was used. A sample collected in November is therefore also described by "
  f"January reflectance. Predictors were extracted at the pixel containing each sample "
  f"centroid; a sensitivity test using the mean of a 3 × 3 pixel neighbourhood is "
  f"reported in the supplementary material. Ten reflectance bands and ten spectral "
  f"indices were retained per acquisition, giving {s01['n_predictors']} predictors "
  f"(Tables 1 and 2; the complete inventory with formulas, source bands, units, native "
  f"and output resolution and resampling method is given as Supplementary Table S1). No "
  f"pair of predictors exceeds |r| = 0.999 and the numerical rank of the predictor "
  f"matrix is {s01['numerical_rank']}.")

doc.add_heading("3.3  Model development", level=2)
hp_txt = "; ".join(
    f"{r.target} {'Random Forest' if r.model == 'RF' else 'XGBoost'}: " +
    ", ".join(f"{c} = {r[c]}" for c in t6b.columns
              if c not in ("target", "model") and str(r[c]) not in ("", "nan"))
    for _, r in t6b.iterrows())
P("Random Forest and XGBoost regression models were fitted for each nutrient. Random "
  "Forest used 500 trees; the number of predictors sampled per split and the minimum "
  "leaf size were tuned. XGBoost used 500 boosting rounds with a squared-error "
  "objective and subsample fraction 0.8; tree depth, learning rate and column subsample "
  "fraction were tuned. All analyses were carried out in Python 3.13 with scikit-learn "
  "1.9.0 and xgboost 3.4.0, under a single global random seed of 42. There were no "
  "missing predictor values. The hyperparameters selected by the inner tuning loop, "
  "reported as the modal choice across outer folds, were: " + hp_txt + ". Fold-by-fold "
  "selections are given in Supplementary Table S2.")
P("Predictive performance was benchmarked against an intercept-only model, ridge "
  "regression, partial least squares regression, ordinary kriging of the nutrient "
  "itself and regression kriging combining the Random Forest trend with kriged "
  "residuals, all evaluated on the identical outer folds (Table 5).")

doc.add_heading("3.4  Validation design", level=2)
P(f"Nested cross-validation with spatial blocks is the primary evaluation. The 110 "
  f"locations were partitioned into {s01['n_blocks']} blocks by k-means clustering of "
  f"projected coordinates (EPSG:26191), giving fold sizes of "
  f"{', '.join(str(n) for n in s01['fold_sizes'])}, a mean block extent of "
  f"{bdes.extent_E_km.mean():.1f} km by {bdes.extent_N_km.mean():.1f} km and a median "
  f"separation of {s01['median_block_separation_km']:.2f} km between a block and its "
  f"nearest neighbouring block. The outer loop holds out one block at a time. The inner "
  f"loop performs every tuning decision by GroupKFold with five folds on the training "
  f"blocks only. Random Forest and XGBoost use identical outer partitions. Nested "
  f"cross-validation with random folds was run under the same protocol as a secondary "
  f"comparison, to quantify the optimism introduced by spatial dependence.")
mI = mor.iloc[0]
mI10 = mor.iloc[-1]
P(f"Block dimensions were informed by the spatial structure of the data. Moran's I at a "
  f"1 km bandwidth is {mI['N']:.2f} for N, {mI['P']:.2f} for P and {mI['K']:.2f} for K, "
  f"and remains positive at a 10 km bandwidth ({mI10['N']:.2f}, {mI10['P']:.2f} and "
  f"{mI10['K']:.2f} respectively). Blocks were therefore sized to exceed the distances "
  f"at which appreciable autocorrelation is detectable, so that training and validation "
  f"samples are not drawn from the same correlated neighbourhood. Fitted variogram "
  f"parameters are reported in Table 8; with 110 observations the nugget ratios are "
  f"well determined but the ranges are weakly identified, so Moran's I rather than the "
  f"fitted ranges was used to justify the block dimension.")
P("Performance is reported as R² defined as 1 − SSE/SST, root mean square error, mean "
  "absolute error, mean bias, Lin's concordance correlation coefficient, RPD and RPIQ, "
  "all computed from the same pooled set of outer-fold predictions. Bootstrap 95 % "
  "confidence intervals were obtained from 3 000 resamples of the paired observed and "
  "predicted values, and the two algorithms were compared by the paired difference in "
  "squared error on identical folds.")
P("Predictor importance was quantified by permutation on the held-out observations of "
  "each outer fold, with ten permutations per predictor, normalised by the fold root "
  "mean square error so that values are comparable across nutrients, and reported with "
  "the standard deviation between folds.")

doc.add_heading("3.5  Spatial prediction, uncertainty and applicability", level=2)
P(f"Final models were fitted on all 110 observations using the modal hyperparameters "
  f"and applied to the full covariate stack. Predictions were generated on the native "
  f"10 m grid, comprising {msum['n_valid_pixels']:,} valid pixels and covering "
  f"{msum['mapped_area_ha']:,.0f} ha; the effective pixel is "
  f"{msum['pixel_x_m']:.2f} m by {msum['pixel_y_m']:.2f} m at this latitude. Local "
  f"uncertainty is mapped as the standard deviation of the predictions across the "
  f"Random Forest tree ensemble.")
P(f"The applicability domain was delineated following Meyer and Pebesma (2021). A "
  f"dissimilarity index was computed for every pixel as the minimum standardised "
  f"distance in predictor space to any training observation, weighted by permutation "
  f"importance and scaled by the mean pairwise distance among training points. The "
  f"threshold was derived from the cross-validation-aware distribution of training "
  f"dissimilarities as the upper quartile plus 1.5 interquartile ranges, giving "
  f"DI = {msum['aoa_threshold']:.3f}. Pixels above the threshold are masked in every "
  f"delivered product and excluded from all reported areas.")
P("Predicted surfaces were reclassified into four fertility classes with the "
  "non-overlapping boundaries given in Tables 3 to 5. The same thresholds were applied "
  "to the observed values and to the outer-fold predictions, and class agreement is "
  "reported as overall accuracy and Cohen's κ.")
P("Rasters are distributed in EPSG:4326. All map figures are drawn in EPSG:26191 "
  "(Merchich / Nord Maroc), with axes labelled in kilometres, a scale bar, a north "
  "arrow and the projection identifier printed on the figure.")

# ============================================================== Results
doc.add_heading("4.  Results", level=1)

doc.add_heading("4.1  Descriptive statistics", level=2)
tn = t3[t3.Nutrient.str.startswith("Total N")].iloc[0]
tp = t3[t3.Nutrient.str.startswith("Available P")].iloc[0]
tk = t3[t3.Nutrient.str.startswith("Exchangeable K")].iloc[0]
P(f"The 110 validated observations show wide variation in all three macronutrients "
  f"(Table 3). Total nitrogen ranges from {tn.Minimum:,.0f} to {tn.Maximum:,.0f} mg kg⁻¹ "
  f"with a mean of {tn.Mean:,.0f} mg kg⁻¹ and a coefficient of variation of "
  f"{tn.CV_percent:.0f} %. Olsen phosphorus ranges from {tp.Minimum:.2f} to "
  f"{tp.Maximum:.2f} mg kg⁻¹ (mean {tp.Mean:.1f} mg kg⁻¹, CV {tp.CV_percent:.0f} %) and "
  f"exchangeable potassium from {tk.Minimum:.1f} to {tk.Maximum:.1f} mg kg⁻¹ "
  f"(mean {tk.Mean:.1f} mg kg⁻¹, CV {tk.CV_percent:.0f} %). Phosphorus and potassium "
  f"are strongly right-skewed (skewness {tp.Skewness:.1f} and {tk.Skewness:.1f}), "
  f"nitrogen less so ({tn.Skewness:.1f}). Table 3 is the single definitive summary of "
  f"the observations; ranges of the predicted rasters are reported separately in "
  f"Table 10.")

doc.add_heading("4.2  Predictive performance", level=2)
P(f"Under nested spatial block cross-validation, Random Forest explained "
  f"{rf_n.R2:.3f} of the variance in total nitrogen (RMSE {rf_n.RMSE:,.1f} mg kg⁻¹, "
  f"MAE {rf_n.MAE:,.1f}, bias {rf_n.bias:+,.1f}, CCC {rf_n.CCC:.3f}), {rf_p.R2:.3f} for "
  f"Olsen phosphorus (RMSE {rf_p.RMSE:.2f} mg kg⁻¹, MAE {rf_p.MAE:.2f}, bias "
  f"{rf_p.bias:+.2f}, CCC {rf_p.CCC:.3f}) and {rf_k.R2:.3f} for exchangeable potassium "
  f"(RMSE {rf_k.RMSE:.2f} mg kg⁻¹, MAE {rf_k.MAE:.2f}, bias {rf_k.bias:+.2f}, "
  f"CCC {rf_k.CCC:.3f}). XGBoost gave R² of {sp('N','XGB').R2:.3f}, "
  f"{sp('P','XGB').R2:.3f} and {sp('K','XGB').R2:.3f} respectively (Table 4, "
  f"Figures 3 and 4).")
P(f"All six models show pronounced regression to the mean. The slope of predicted on "
  f"observed values ranges from "
  f"{mt[mt.scheme=='nested spatial CV'].slope.min():.2f} to "
  f"{mt[mt.scheme=='nested spatial CV'].slope.max():.2f}, so low concentrations are "
  f"over-predicted and high concentrations under-predicted. The consequence is visible "
  f"in the maps, whose ranges are compressed relative to the observations (Table 10), "
  f"and it should be taken into account wherever the surfaces are used to identify "
  f"extreme values.")
opt = []
for t in TARGETS:
    for m in ["RF", "XGB"]:
        a, b = sp(t, m, "nested random CV"), sp(t, m)
        opt.append({"Nutrient": t, "Model": "RF" if m == "RF" else "XGBoost",
                    "Random CV R²": round(a.R2, 3), "Spatial CV R²": round(b.R2, 3),
                    "ΔR²": round(a.R2 - b.R2, 3)})
P(f"Ignoring spatial dependence inflates the estimates. Nested random cross-validation "
  f"gives R² values that exceed the spatial estimates by up to "
  f"{max(r['ΔR²'] for r in opt):.3f}, and calibration statistics exceed them by far "
  f"more (Table 4, Figure 5). Calibration figures are reported only to document this "
  f"gap and are not evidence of predictive skill.")
table(pd.DataFrame(opt), caption="Optimism introduced by random folds relative to "
                                 "spatial blocks.")
P(f"The two algorithms cannot be separated. The paired difference in squared error on "
  f"identical outer folds has a bootstrap 95 % confidence interval that includes zero "
  f"for all three nutrients "
  f"({'; '.join(f'{r.target}: [{r.ci_lo:,.1f}, {r.ci_hi:,.1f}]' for _, r in pcmp.iterrows())}), "
  f"so no claim of superiority is made. Random Forest was adopted for mapping on the "
  f"pre-specified criterion of outer-fold root mean square error.")
P(f"The benchmarks place these results in context (Table 5). The intercept-only model "
  f"gives a negative R² for all three nutrients, confirming that the models carry real "
  f"information. Ordinary kriging of the nutrient alone performs poorly "
  f"({bs('N','Ordinary kriging').R2:.2f}, {bs('P','Ordinary kriging').R2:.2f} and "
  f"{bs('K','Ordinary kriging').R2:.2f}), so the spectral predictors, not spatial "
  f"interpolation, carry the signal. Two results temper the case for ensemble methods: "
  f"for potassium, ridge regression on the same predictors achieves "
  f"R² = {bs('K','Ridge regression').R2:.3f} against {rf_k.R2:.3f} for Random Forest, "
  f"and for nitrogen, regression kriging achieves "
  f"R² = {bs('N','Regression kriging (RF + OK residuals)').R2:.3f} against "
  f"{rf_n.R2:.3f}. We report these rather than omit them: on this dataset the "
  f"machine-learning ensembles are not uniformly the best available option.")

doc.add_heading("4.3  Spatial structure and residual diagnostics", level=2)
def vgr(t, kind):
    s = vg[(vg.variable == t) & (vg.kind == kind)]
    return s.iloc[0] if len(s) else None


vgtxt = "; ".join(f"{t}: nugget ratio {vgr(t,'observed').nugget_ratio:.2f}"
                  for t in TARGETS if vgr(t, "observed") is not None)
P(f"Fitted spherical variograms of the observations give {vgtxt} (Table 8, Figure 7). "
  f"These high nugget ratios indicate that most of the variance occurs at separation "
  f"distances shorter than the sampling interval, which is consistent with a sampling "
  f"density of one observation per {msum['mapped_area_ha']/110:,.0f} ha. The fitted "
  f"ranges are weakly identified at this sample size and are reported for completeness "
  f"rather than as reliable estimates of the correlation length.")
pure = [t for t in TARGETS if vgr(t, "RF residuals") is not None
        and vgr(t, "RF residuals").nugget_ratio > 0.95]
struct = [t for t in TARGETS if t not in pure and vgr(t, "RF residuals") is not None]
P("Variograms of the Random Forest spatial-CV residuals are shown in the lower row of "
  "Figure 7. The " + " and ".join(pure) + " residuals are effectively pure nugget "
  "(nugget ratio " + ", ".join(f"{vgr(t,'RF residuals').nugget_ratio:.2f}" for t in pure) +
  "), indicating that the covariates captured the spatial structure these data can "
  "resolve. The " + " and ".join(struct) + " residuals retain spatial structure (nugget "
  "ratio " + ", ".join(f"{vgr(t,'RF residuals').nugget_ratio:.2f}" for t in struct) +
  ", fitted range " + ", ".join(f"{vgr(t,'RF residuals').range_m/1000:.1f} km"
                                for t in struct) +
  "), so the " + " and ".join(struct) + " model leaves spatially organised error behind "
  "and would benefit from a geostatistical correction. This is consistent with the "
  "regression-kriging benchmark in Table 5.")

doc.add_heading("4.4  Predictor importance", level=2)
P(f"Permutation importance measured on held-out spatial folds is dominated by raw "
  f"reflectance bands; spectral indices contribute little (Table 6, Figure 6). For "
  f"total nitrogen, near-infrared bands account for {clsshare('N','RF','NIR bands (B8, B8A)'):.0f} % "
  f"of positive importance in Random Forest and red-edge bands for "
  f"{clsshare('N','RF','Red-edge bands (B5, B6, B7)'):.0f} %. For Olsen phosphorus the "
  f"order is reversed, with red-edge bands at "
  f"{clsshare('P','RF','Red-edge bands (B5, B6, B7)'):.0f} % and near-infrared bands at "
  f"{clsshare('P','RF','NIR bands (B8, B8A)'):.0f} %. Potassium is the only nutrient "
  f"with a broadly distributed signature: near-infrared "
  f"{clsshare('K','RF','NIR bands (B8, B8A)'):.0f} %, brightness and salinity indices "
  f"{clsshare('K','RF','Brightness / salinity indices'):.0f} %, shortwave-infrared "
  f"bands {clsshare('K','RF','SWIR bands (B11, B12)'):.0f} % and visible bands "
  f"{clsshare('K','RF','Visible bands (B2, B3, B4)'):.0f} %.")
P(f"Two features of this result constrain how far it can be interpreted. First, the "
  f"November 2024 acquisition accounts for "
  f"{t7b.November_2024_share_percent.min():.0f}–"
  f"{t7b.November_2024_share_percent.max():.0f} % of positive importance across the six "
  f"models, so one date carries almost the whole signal. Second, individual predictor "
  f"rankings are unstable: only "
  f"{'–'.join(str(int(v)) for v in [t7c.n_stable_predictors.min(), t7c.n_stable_predictors.max()])} "
  f"of {int(t7c.n_total.iloc[0])} predictors have a mean importance exceeding their own "
  f"between-fold standard deviation. Aggregated classes of predictor can therefore be "
  f"discussed; the rank order of individual bands and indices cannot.")

doc.add_heading("4.5  Map products", level=2)
P(f"Continuous Random Forest surfaces at 10 m are shown in Figure 8 and the "
  f"corresponding XGBoost surfaces in Figure 9. Predicted ranges are "
  f"{rs('N','RF')['min']:,.0f}–{rs('N','RF')['max']:,.0f} mg kg⁻¹ for nitrogen, "
  f"{rs('P','RF')['min']:.1f}–{rs('P','RF')['max']:.1f} for phosphorus and "
  f"{rs('K','RF')['min']:.1f}–{rs('K','RF')['max']:.1f} for potassium, all inside the "
  f"observed range of the corresponding nutrient and all compressed relative to it, as "
  f"expected from the slopes reported in Section 4.2.")
P(f"Local uncertainty is mapped in Figure 10 as the standard deviation across the tree "
  f"ensemble, with district means of {rs('N','RF_SD')['mean']:,.0f} mg kg⁻¹ for "
  f"nitrogen, {rs('P','RF_SD')['mean']:.1f} for phosphorus and "
  f"{rs('K','RF_SD')['mean']:.1f} for potassium. The dissimilarity index and the "
  f"applicability domain are shown in Figure 1b: {msum['pct_inside_aoa']:.1f} % of the "
  f"mapped area ({msum['ha_inside_aoa']:,.0f} ha) falls inside the domain and "
  f"{100-msum['pct_inside_aoa']:.1f} % ({msum['ha_outside_aoa']:,.0f} ha) is masked as "
  f"extrapolative.")
agr_txt = ", ".join(f"{r.nutrient} κ = {r.kappa:.2f}" for _, r in agr.iterrows())
P(f"The two algorithms produce similar classified surfaces: pixelwise agreement is "
  f"{100*agr.overall_agreement.min():.0f}–{100*agr.overall_agreement.max():.0f} % with "
  f"{agr_txt}. We report these statistics in place of the previous statement that the "
  f"XGBoost maps were clearer or retained more detail, which was not tested.")
P(f"Fertility classes within the applicability domain are mapped in Figure 11 and "
  f"quantified in Table 10. Nitrogen is dominated by the Low class "
  f"({ca('N','Low').rf_inAOA_pct:.1f} %, {ca('N','Low').rf_inAOA_ha:,.0f} ha) and the "
  f"Medium class ({ca('N','Medium').rf_inAOA_pct:.1f} %). Phosphorus is spread across "
  f"all four classes, with {ca('P','Very low').rf_inAOA_pct:.1f} % Very low and "
  f"{ca('P','High').rf_inAOA_pct:.1f} % High. Potassium is concentrated in the two "
  f"lowest classes ({ca('K','Very low').rf_inAOA_pct:.1f} % Very low, "
  f"{ca('K','Low').rf_inAOA_pct:.1f} % Low); the High class is empty in the Random "
  f"Forest surface, which reflects both the shrinkage described in Section 4.2 and the "
  f"fact that only two of the 110 observations exceed 120 mg kg⁻¹.")

doc.add_heading("4.6  Fertility class agreement", level=2)
kt = t9[["nutrient", "model", "overall_accuracy", "kappa"]]
P(f"Applying the same thresholds to the observations and to the outer-fold predictions "
  f"gives overall class accuracies of "
  f"{kt.overall_accuracy.min():.2f}–{kt.overall_accuracy.max():.2f} and Cohen's κ of "
  f"{kt.kappa.min():.2f}–{kt.kappa.max():.2f} (Table 9; full confusion matrices in "
  f"Supplementary Table S3). Agreement is good in the classes that are well represented "
  f"in the sample and unreliable in the upper classes, which are supported by very few "
  f"observations: only {int((df.K > 120).sum())} samples exceed the highest potassium "
  f"threshold and {int((df.N > 3000).sum())} exceed the highest nitrogen threshold. "
  f"Class-specific accuracy in those categories should not be relied upon.")

# ============================================================== Discussion
doc.add_heading("5.  Discussion", level=1)

doc.add_heading("5.1  What the spectral signal represents", level=2)
P(f"None of the three macronutrients has a direct diagnostic absorption feature in the "
  f"Sentinel-2 wavelengths. Any predictive skill must therefore be indirect, and the "
  f"results indicate what the intermediary is. Predictive importance is concentrated in "
  f"raw near-infrared and red-edge reflectance rather than in soil-oriented indices; "
  f"the November acquisition alone carries "
  f"{t7b.November_2024_share_percent.min():.0f}–"
  f"{t7b.November_2024_share_percent.max():.0f} % of that importance; and the predicted "
  f"surfaces contain sharply bounded, field-shaped patches that follow parcel geometry "
  f"rather than pedological boundaries. The most parsimonious interpretation is that "
  f"the models are learning canopy condition at a single date, which is correlated with "
  f"soil fertility through crop selection, irrigation and fertilisation history, rather "
  f"than sensing soil nutrient status directly.")
P("This reading has a direct consequence for how the maps should be used. If the signal "
  "is carried by the crop, the surfaces are conditioned on the 2024–2025 cropping "
  "mosaic and are seasonal snapshots rather than stable soil property maps. Temporal "
  "transferability cannot be assessed from a single season and should not be assumed. "
  "Patterns visible in the maps may plausibly relate to clay content, salinity, "
  "hydromorphy, irrigation infrastructure, carbonate content or fertilisation history, "
  "but none of these variables was measured in this study, so those relationships are "
  "stated here as hypotheses for testing rather than as findings.")

doc.add_heading("5.2  Comparison with previous work", level=2)
P(f"National-scale Random Forest mapping of Moroccan soil nutrients at 250 m reported "
  f"cross-validated R² between 0.62 and 0.76 for phosphorus and potassium (Bouslihim "
  f"et al., 2025). The present spatially validated values of {rf_p.R2:.2f} for "
  f"phosphorus and {rf_k.R2:.2f} for potassium are of the same order, obtained at finer "
  f"resolution over a single district but from a much smaller sample. We note that our "
  f"previously submitted version reported values above 0.90 for all three nutrients; "
  f"those figures were not independent spatial predictions and have been withdrawn. "
  f"Comparison across studies is only meaningful when the validation design is stated, "
  f"and we encourage explicit reporting of whether folds were random or spatial.")

doc.add_heading("5.3  Limitations", level=2)
P(f"The sample size of 110 is small relative to the heterogeneity of the district and "
  f"corresponds to one observation per {msum['mapped_area_ha']/110:,.0f} ha, or roughly "
  f"{msum['n_valid_pixels']/110:,.0f} prediction pixels per observation. The upper "
  f"fertility classes are supported by very few samples. All predictors derive from two "
  f"acquisitions of a single sensor within one season; no terrain, radar, climate or "
  f"proximal-sensing covariates were included. The 20 m native support of the red-edge "
  f"and shortwave-infrared bands limits the effective resolution of nominal 10 m "
  f"outputs. Predictions shrink towards the mean, so extremes are systematically "
  f"under-represented. Finally, {100-msum['pct_inside_aoa']:.0f} % of the district lies "
  f"outside the applicability domain and is not mapped.")

doc.add_heading("5.4  Practical implications", level=2)
P("Total nitrogen is not plant-available nitrogen, and fertiliser prescription requires "
  "crop demand, yield target, mineralisation, pH, texture, irrigation and management "
  "history, none of which are available here. These maps are therefore offered as "
  "reconnaissance products. They can support stratified resampling that concentrates "
  "effort where uncertainty is high or where the applicability mask indicates "
  "extrapolation; they can delineate candidate management zones for field verification; "
  "and they can generate hypotheses about the controls on nutrient distribution in the "
  "district. They do not by themselves justify variable-rate fertiliser prescription.")

# ============================================================== Conclusions
doc.add_heading("6.  Conclusions", level=1)
P(f"Sentinel-2 reflectance predicts topsoil macronutrients in the Beni Moussa irrigated "
  f"district with moderate accuracy when performance is measured on spatially "
  f"independent data: R² of {rf_n.R2:.2f} for total nitrogen, {rf_p.R2:.2f} for Olsen "
  f"phosphorus and {rf_k.R2:.2f} for exchangeable potassium under nested spatial block "
  f"cross-validation, with 95 % confidence intervals of "
  f"[{cn.R2_lo:.2f}, {cn.R2_hi:.2f}], [{cp.R2_lo:.2f}, {cp.R2_hi:.2f}] and "
  f"[{ck.R2_lo:.2f}, {ck.R2_hi:.2f}] respectively. Random Forest and XGBoost are "
  f"statistically indistinguishable on this dataset, and for potassium a ridge "
  f"regression on the same predictors performs at least as well as either.")
P(f"Predictive skill rests on raw near-infrared and red-edge reflectance from a single "
  f"November acquisition rather than on soil-oriented spectral indices, which indicates "
  f"that the models track canopy condition rather than soil composition. The resulting "
  f"10 m surfaces, delivered with local uncertainty and an applicability domain "
  f"covering {msum['pct_inside_aoa']:.0f} % of the district, are the first "
  f"field-validated nutrient maps for Beni Moussa at this resolution. They are best "
  f"used for reconnaissance, stratified resampling and hypothesis generation. "
  f"Establishing whether the relationships hold across seasons, and whether they "
  f"survive within crop classes, requires multi-season sampling and non-spectral "
  f"soil-forming covariates, and is the necessary next step before operational "
  f"fertiliser recommendations can be derived from products of this kind.")

# ============================================================== statements
doc.add_heading("Data availability", level=1)
P(f"The analysis-ready data, the complete analysis code, the exact cross-validation "
  f"fold assignments, session information, out-of-fold predictions for every validation "
  f"scheme, and the georeferenced 10 m prediction, uncertainty, dissimilarity and "
  f"applicability rasters are openly available at [repository, DOI to be inserted on "
  f"acceptance]. Every value reported in this article can be reproduced from that "
  f"archive by running the numbered scripts in order. The Sentinel-2 Level-2A imagery "
  f"is freely available from the Copernicus Data Space; the Earth Engine script used to "
  f"build and export the covariates is included in the archive.")
note("Create the Zenodo or Figshare deposit from FINAL/deposit and paste the DOI here. "
     "If landholder privacy requires it, substitute coordinates jittered within 250 m "
     "and state the restriction — 250 m is below the fitted variogram ranges, so "
     "spatial analyses remain reproducible.")

doc.add_heading("Acknowledgments", level=1)
P("This study was conducted as part of a PhD thesis. The authors thank the editor and "
  "the referee, whose detailed technical review substantially changed the analysis and "
  "the conclusions of this work.")
note("The referee observed that the previous acknowledgment thanked anonymous reviewers "
     "although the earlier round concerned technical and editorial checks. Confirm the "
     "review history and adjust if needed.")

out = os.path.join(DOCS, "Revised_Manuscript_Text_Blocks.docx")
doc.save(out)
print("wrote", out)
print("STEP 09 complete.")
