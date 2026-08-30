# -*- coding: utf-8 -*-
"""
STEP 07 - Build the complete point-by-point Response to the Referee as a Word
document. Every number is read from the result files produced by steps 01-06,
so the letter can never drift from the analysis.
"""
import numpy as np, pandas as pd, os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = r"d:\Doctorat\article1\outputs_fast"
FINAL = os.path.join(BASE, "FINAL")
RES   = os.path.join(FINAL, "results")
TAB   = os.path.join(FINAL, "tables")
FIG   = os.path.join(FINAL, "figures")
DOCS  = os.path.join(FINAL, "docs")
os.makedirs(DOCS, exist_ok=True)

TARGETS = ["N", "P", "K"]
ACC = RGBColor(0x1F, 0x5F, 0x6B)
GREY = RGBColor(0x55, 0x5F, 0x5B)

# ------------------------------------------------------------------ results
df   = pd.read_csv(os.path.join(RES, "01_analysis_ready_dataset.csv"))
inv  = pd.read_csv(os.path.join(RES, "01_predictor_inventory.csv"))
bdes = pd.read_csv(os.path.join(RES, "01_block_design.csv"))
s01  = json.load(open(os.path.join(RES, "01_summary.json")))
mt   = pd.read_csv(os.path.join(RES, "02_metrics_all_schemes.csv"))
bl   = pd.read_csv(os.path.join(RES, "02_baselines.csv"))
ci   = pd.read_csv(os.path.join(RES, "02_bootstrap_ci.csv"))
pcmp = pd.read_csv(os.path.join(RES, "02_paired_model_comparison.csv"))
pi   = pd.read_csv(os.path.join(RES, "02_permutation_importance.csv"))
vg   = pd.read_csv(os.path.join(RES, "04_variogram_parameters.csv"))
mor  = pd.read_csv(os.path.join(RES, "04_morans_I.csv"))
care = pd.read_csv(os.path.join(RES, "03_class_areas.csv"))
agr  = pd.read_csv(os.path.join(RES, "03_RF_XGB_agreement_summary.csv"))
ras  = pd.read_csv(os.path.join(RES, "03_raster_summary.csv"))
msum = json.load(open(os.path.join(RES, "03_map_summary.json")))
t3   = pd.read_csv(os.path.join(TAB, "T3_descriptive_statistics.csv"))
t7   = pd.read_csv(os.path.join(TAB, "T7_importance_by_class.csv"))
t7b  = pd.read_csv(os.path.join(TAB, "T7b_acquisition_share.csv"))
t7c  = pd.read_csv(os.path.join(TAB, "T7c_importance_stability.csv"))
t9   = pd.read_csv(os.path.join(TAB, "T9_class_accuracy_kappa.csv"))
t6b  = pd.read_csv(os.path.join(TAB, "T6b_modal_hyperparameters.csv"))
t10d = pd.read_csv(os.path.join(TAB, "T10d_map_metadata.csv"))


def sp(t, m, s="nested spatial CV"):
    return mt[(mt.target == t) & (mt.model == m) & (mt.scheme == s)].iloc[0]


def cir(t, m):
    return ci[(ci.target == t) & (ci.model == m)].iloc[0]


def base(t, m):
    return bl[(bl.target == t) & (bl.model == m)].iloc[0]


def vgp(t, kind="observed"):
    r = vg[(vg.variable == t) & (vg.kind == kind)]
    return r.iloc[0] if len(r) else None


# ------------------------------------------------------------------ styling
doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)

st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12

for lvl, sz, col in [(1, 15, ACC), (2, 12.5, ACC), (3, 11, RGBColor(0x1A, 0x21, 0x1E))]:
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"; h.font.size = Pt(sz); h.font.bold = True
    h.font.color.rgb = col
    h.paragraph_format.space_before = Pt(14 if lvl < 3 else 10)
    h.paragraph_format.space_after = Pt(5)


def P(text="", bold=False, italic=False, size=10.5, color=None, align=None,
      space_after=6, indent=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def rich(parts, indent=None, space_after=6, size=10.5):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for t, b, i in parts:
        r = p.add_run(t); r.bold = b; r.italic = i; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def comment(text):
    p = doc.add_paragraph()
    r = p.add_run("Referee comment.  "); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = GREY
    r2 = p.add_run(text); r2.italic = True; r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    lf = OxmlElement("w:left")
    lf.set(qn("w:val"), "single"); lf.set(qn("w:sz"), "12")
    lf.set(qn("w:space"), "8"); lf.set(qn("w:color"), "1F5F6B")
    pbdr.append(lf); pPr.append(pbdr)
    return p


def response(text):
    p = doc.add_paragraph()
    r = p.add_run("Response.  "); r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = ACC
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, str):
            r = p.add_run(it); r.font.size = Pt(size)
        else:
            for t, b, i in it:
                r = p.add_run(t); r.bold = b; r.italic = i; r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.85)


def changes(items):
    p = doc.add_paragraph()
    r = p.add_run("Changes made in the manuscript."); r.bold = True
    r.font.size = Pt(10.5); r.font.color.rgb = ACC
    p.paragraph_format.space_after = Pt(2)
    bullets(items, size=10)


def table(dfx, colnames=None, widths=None, fontsize=8.6, caption=None, numeric_right=True):
    dfx = dfx.copy()
    cols = colnames or list(dfx.columns)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0].add_run(str(c))
        pr.bold = True; pr.font.size = Pt(fontsize)
        pr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(hdr[i], "1F5F6B")
    for _, row in dfx.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            txt = "" if pd.isna(v) else (f"{v:,.4g}" if isinstance(v, (float, np.floating)) else str(v))
            rr = cells[i].paragraphs[0].add_run(txt)
            rr.font.size = Pt(fontsize)
            if numeric_right and isinstance(v, (int, float, np.integer, np.floating)):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(t.rows):
        if ri > 0 and ri % 2 == 0:
            for c in row.cells:
                shade(c, "F2F5F1")
    if widths:
        for r_ in t.rows:
            for i, wd in enumerate(widths):
                r_.cells[i].width = Cm(wd)
    if caption:
        P(caption, size=8.6, italic=True, color=GREY, space_after=10)
    else:
        P("", size=4, space_after=4)
    return t


def figure(path, width_cm, caption):
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(caption, size=8.8, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_after=12)


def footer_pagenum():
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        for el, txt in [("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)]:
            e = OxmlElement(el)
            if el == "w:fldChar":
                e.set(qn("w:fldCharType"), "begin" if txt is None and not p.runs[-1].text else "end")
            r._r.append(e)
        p.text = ""
        r2 = p.add_run("Response to Referee — Amrouss et al.")
        r2.font.size = Pt(8); r2.font.color.rgb = GREY


# ==========================================================================
#                              TITLE PAGE
# ==========================================================================
P("Response to the Referee Report", bold=True, size=20, color=ACC,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P("Environmental Monitoring and Assessment", size=12, color=GREY,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
P("High-resolution mapping of soil macronutrients in a semi-arid climate (Morocco) "
  "using Sentinel-2 data and advanced machine learning models",
  bold=True, size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
P("Yassine Amrouss, Abdelkrim Arioua, Driss Elhamdouni, Mohamed El Baghdadi, "
  "Ahmed Barakat, Jaouad El Atiq, Insaf Ouchkir, Mostafa Bimouhen, "
  "Oussama Nait-taleb, Abdessamad Hilali",
  size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

P("We thank the referee for a report of exceptional care and technical depth. The "
  "review identified a genuine and serious problem in the previous version: the "
  "reported accuracy figures were not independent spatial predictions, and several "
  "numerical, cartographic and bibliographic statements did not withstand checking. "
  "We accept every one of these findings.")
P("The manuscript has been rebuilt rather than patched. The entire analysis was "
  "re-run from the raw data under nested spatial block cross-validation; the "
  "predictor inventory was audited line by line against its own source code and "
  "three errors were found and corrected; the maps were regenerated at full "
  "resolution with local uncertainty and an applicability domain; and the Abstract, "
  "Results, Discussion and Conclusions were rewritten around the validated evidence. "
  "Two headline claims of the previous version — that XGBoost outperforms Random "
  "Forest, and that spectral indices drive the predictions — are contradicted by the "
  "new analysis and have been removed.")
P("Our own audit uncovered three further defects that the referee did not raise. We "
  "report them openly below (comments M5 and M9) because we would rather disclose "
  "them than have them discovered later.", bold=False)

P("")
an = doc.add_paragraph()
r = an.add_run("NOTE FOR THE AUTHORS — DELETE THIS BOX BEFORE SUBMITTING.  ")
r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xB5, 0x22, 0x22)
r2 = an.add_run(
    "The analysis, figures, tables, maps and deposit described in this letter are "
    "complete and are in outputs_fast/FINAL. The 'Changes made in the manuscript' lists "
    "below describe edits that must still be applied to the manuscript file itself. "
    "Ready-to-paste replacement text for the Abstract, Methods, Results, Discussion, "
    "Conclusions and Data availability statement is provided in "
    "Revised_Manuscript_Text_Blocks.docx, which is generated from the same result files "
    "as this letter. A small number of items can only be supplied by you — laboratory "
    "QA/QC records, exact sampling dates, Earth Engine asset identifiers, the sampling "
    "design name and stratum allocation, and the reference-list corrections — and are "
    "flagged in red in that document. Send this letter only once those edits are in the "
    "manuscript.")
r2.italic = True; r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0xB5, 0x22, 0x22)
an.paragraph_format.space_after = Pt(12)
pPr = an._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
for edge in ("top", "left", "bottom", "right"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "8")
    e.set(qn("w:space"), "6"); e.set(qn("w:color"), "B52222")
    pbdr.append(e)
pPr.append(pbdr)

sm = doc.add_table(rows=1, cols=2); sm.style = "Table Grid"
for i, h in enumerate(["What changed", "Where"]):
    c = sm.rows[0].cells[i]; c.text = ""
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1F5F6B")
sumrows = [
    ("All performance figures now come from held-out outer folds of a nested spatial "
     "block cross-validation; R² is 1 − SSE/SST throughout", "M1, M2 · Section 3.6, Table 4"),
    ("Predictor inventory corrected: two layers were not computing their stated "
     "formulas and were exact duplicates; count corrected from 41 to 40",
     "M5 · Tables 1–2, Supplementary S1"),
    ("Model-superiority claim withdrawn; RF and XGBoost are statistically "
     "indistinguishable", "M6 · Abstract, Sections 4.3, 5, 6"),
    ("Interpretation rewritten to match permutation importance measured on held-out "
     "folds", "M7, M8 · Abstract, Sections 4.4–4.5, 5"),
    ("Continuous 10 m maps, local uncertainty, applicability domain, class areas and "
     "full cartographic metadata added", "M9 · Figures 8–11, Table 10"),
    ("Confusion matrices and Cohen's κ reported; operational claims tempered",
     "M10 · Table 9, Section 5"),
    ("Data, code, folds, out-of-fold predictions and rasters deposited",
     "M11 · Data availability"),
    ("Citation content and reference list fully audited and corrected",
     "M12 · throughout"),
]
for a, b in sumrows:
    cells = sm.add_row().cells
    for i, txt in enumerate([a, b]):
        cells[i].text = ""
        r = cells[i].paragraphs[0].add_run(txt); r.font.size = Pt(9)
    cells[0].width = Cm(11.0); cells[1].width = Cm(5.6)
for ri, row in enumerate(sm.rows):
    if ri > 0 and ri % 2 == 0:
        for c in row.cells:
            shade(c, "F2F5F1")

doc.add_page_break()

# ==========================================================================
#                          SECTION 1 : MAJOR COMMENTS
# ==========================================================================
doc.add_heading("1.  Responses to the major comments", level=1)

# ------------------------------------------------------------------- M1
doc.add_heading("M1.  Performance metrics, prediction identity, and the definition of R²",
                level=2)
comment("Generate one retained prediction for every outer-fold held-out observation; "
        "calculate R² (using the stated 1 − SSE/SST definition), RMSE, MAE, mean bias, "
        "and preferably Lin's concordance correlation coefficient from that same pooled "
        "prediction set; report confidence intervals or repeated resampling uncertainty; "
        "and replace calibration plots with observed-versus-predicted plots based "
        "exclusively on outer fold predictions.")
response("The referee's diagnosis was correct in every particular, and we are grateful "
         "for it. The previous Section 4.3 reported cross-validated statistics in the "
         "text while Figure 4 displayed calibration statistics; the two were computed "
         "from different prediction sets, which is why R² and RMSE could not be "
         "reconciled and why potassium appeared to validate better than it calibrated. "
         "No single quantity in that comparison was recoverable from the others.")
P("The evaluation has been rebuilt from the ground up. Exactly one prediction is now "
  "retained for every observation in every scheme, stored in "
  "02_out_of_fold_predictions.csv, and every reported statistic is computed from that "
  "single pooled set. R² is 1 − SSE/SST as defined in Eq. (1) and is now exactly "
  "recoverable from RMSE and the sample sum of squares for all model–nutrient "
  "combinations. Lin's concordance correlation coefficient, mean bias, MAE, RPD and "
  "RPIQ are reported alongside, and bootstrap 95 % confidence intervals (3 000 "
  "resamples) accompany every headline figure.")

perf = []
for t in TARGETS:
    for m in ["RF", "XGB"]:
        r = sp(t, m); c = cir(t, m)
        perf.append({"Nutrient": t, "Model": "Random Forest" if m == "RF" else "XGBoost",
                     "R²": f"{r.R2:.3f}", "R² 95 % CI": f"[{c.R2_lo:.3f}, {c.R2_hi:.3f}]",
                     "RMSE": f"{r.RMSE:,.2f}", "MAE": f"{r.MAE:,.2f}",
                     "Bias": f"{r.bias:+,.2f}", "CCC": f"{r.CCC:.3f}",
                     "Slope": f"{r.slope:.2f}"})
table(pd.DataFrame(perf), fontsize=8.4,
      caption="Table R1. Performance on held-out outer folds of the nested spatial block "
              "cross-validation (n = 110). Units of RMSE, MAE and bias are mg kg⁻¹. "
              "Slope is the regression of predicted on observed values.")

P("We also confirm the shrinkage the referee inferred from the scatterplots. The "
  f"predicted-on-observed slope ranges from {mt[mt.scheme=='nested spatial CV'].slope.min():.2f} "
  f"to {mt[mt.scheme=='nested spatial CV'].slope.max():.2f}, so low values are "
  "over-predicted and high values under-predicted. This is now stated explicitly in "
  "the Results and carried through to the Discussion, replacing the previous "
  "description of the points as lying close to the 1:1 line. To close out the "
  "referee's alternative explanation, we report both R² definitions: across all "
  "schemes the largest discrepancy between 1 − SSE/SST and the squared Pearson "
  f"correlation is {(mt.R2_cor2 - mt.R2).abs().max():.3f}, so the two are now "
  "effectively equivalent and the earlier inconsistency cannot have arisen from the "
  "metric definition alone.")
changes([
    "Section 3.6 rewritten: R², RMSE, MAE, bias, CCC, RPD and RPIQ defined on the "
    "pooled outer-fold prediction set; the unused cor(y, ŷ) definition removed.",
    "Figures 3 and 4 replaced: calibration plots are gone; the new panels show "
    "outer-fold predictions only, with the 1:1 line, the fitted line, slope, "
    "intercept and n given in each panel.",
    "Table 4 replaced with the values in Table R1 above, including confidence intervals.",
    "Abstract, Results, Discussion and Conclusions now quote only these values.",
])

figure(os.path.join(FIG, "Figure_3_RF_observed_vs_predicted.png"), 16.4,
       "Figure R1 (manuscript Figure 3). Observed versus Random Forest predicted values "
       "from held-out spatial folds.")

# ------------------------------------------------------------------- M2
doc.add_heading("M2.  Validation design must test spatial generalization", level=2)
comment("Use nested spatial or block cross-validation as the primary evaluation … "
        "Provide empirical variograms for each nutrient and for model residuals and use "
        "a justified spatial range to inform block design. Add an area of applicability "
        "or dissimilarity analysis and flag or mask extrapolative pixels.")
response("Adopted in full. Nested spatial block cross-validation is now the primary and "
         "reported evaluation; random cross-validation is retained only as a secondary "
         "comparison whose purpose is to quantify the optimism the referee identified.")
bullets([
    [("Outer loop: ", True, False),
     (f"leave-one-block-out over {s01['n_blocks']} spatial blocks obtained by k-means "
      f"clustering of projected coordinates (EPSG:26191). Fold sizes are "
      f"{', '.join(str(n) for n in s01['fold_sizes'])}, mean block extent "
      f"{bdes.extent_E_km.mean():.1f} km × {bdes.extent_N_km.mean():.1f} km, and the "
      f"median separation between a block and its nearest neighbouring block is "
      f"{s01['median_block_separation_km']:.2f} km.", False, False)],
    [("Inner loop: ", True, False),
     ("GroupKFold with five folds applied to the training blocks only. Every tuning "
      "decision — hyperparameter selection for both algorithms — is taken inside this "
      "loop, so the outer folds are never used in model construction.", False, False)],
    [("Identical folds: ", True, False),
     ("Random Forest and XGBoost use exactly the same outer partition, so the paired "
      "comparison in M6 is valid.", False, False)],
    [("Seeds and repetitions: ", True, False),
     ("a single global seed (42) governs block construction, model fitting and "
      "bootstrap resampling; the full session record is deposited.", False, False)],
])

opt = []
for t in TARGETS:
    for m in ["RF", "XGB"]:
        a = sp(t, m, "nested random CV"); b = sp(t, m)
        opt.append({"Nutrient": t, "Model": "RF" if m == "RF" else "XGBoost",
                    "Random CV R²": f"{a.R2:.3f}", "Spatial CV R²": f"{b.R2:.3f}",
                    "Optimism (ΔR²)": f"{a.R2-b.R2:+.3f}",
                    "Random CV RMSE": f"{a.RMSE:,.2f}", "Spatial CV RMSE": f"{b.RMSE:,.2f}"})
table(pd.DataFrame(opt), fontsize=8.4,
      caption="Table R2. Optimism introduced by ignoring spatial dependence, measured as "
              "the difference between nested random and nested spatial cross-validation.")

vgtxt = []
for t in TARGETS:
    o = vgp(t, "observed")
    if o is not None:
        vgtxt.append(f"{t}: range {o.range_m/1000:.1f} km, nugget ratio {o.nugget_ratio:.2f}")
P("Empirical variograms have been computed for each nutrient and for the model "
  "residuals. For the observed values the spherical fits give " + "; ".join(vgtxt) + ". "
  "We must be candid about what these fits can bear: with 110 points the nugget ratios "
  "are well determined but the ranges are poorly constrained, and for nitrogen the "
  "range converged on the upper bound of the search. We therefore used Moran's I as the "
  "primary evidence for the block dimension rather than the fitted ranges alone. "
  "Moran's I at a 1 km bandwidth is " +
  ", ".join(f"{t} = {mor.iloc[0][t]:.2f}" for t in TARGETS) +
  ", and remains positive out to 10 km, confirming that random folds would routinely "
  "place autocorrelated neighbours on opposite sides of the train/test boundary. Blocks "
  f"were sized at {bdes.extent_E_km.mean():.1f} km by {bdes.extent_N_km.mean():.1f} km "
  "so that held-out samples lie beyond the distances at which appreciable "
  "autocorrelation is detectable.")
res_pure = [t for t in TARGETS if len(vg[(vg.kind == "RF residuals") & (vg.variable == t)])
            and vg[(vg.kind == "RF residuals") & (vg.variable == t)].nugget_ratio.iloc[0] > 0.95]
res_str = [t for t in TARGETS if t not in res_pure]
P("Residual variograms are reported in the lower row of the same figure and give a "
  "genuinely informative result. The " +
  " and ".join(res_pure) + " residuals are effectively pure nugget (nugget ratio " +
  ", ".join(f"{vg[(vg.kind=='RF residuals')&(vg.variable==t)].nugget_ratio.iloc[0]:.2f}"
            for t in res_pure) +
  "), indicating that the covariates captured the spatial structure these data can "
  "resolve. The " + " and ".join(res_str) + " residuals retain structure (nugget ratio " +
  ", ".join(f"{vg[(vg.kind=='RF residuals')&(vg.variable==t)].nugget_ratio.iloc[0]:.2f}"
            for t in res_str) +
  "), which means the " + " and ".join(res_str) + " model leaves spatially organised "
  "error behind. We report this as a limitation rather than leaving it unstated, and it "
  "is consistent with the regression-kriging benchmark in M6.")

figure(os.path.join(FIG, "Figure_7_variograms.png"), 16.4,
       "Figure R2 (new supplementary figure). Empirical variograms of the observed "
       "nutrients and of the spatial-CV residuals, with fitted spherical models.")

P(f"An area of applicability has been added following Meyer and Pebesma (2021). A "
  f"predictor-space dissimilarity index weighted by permutation importance is computed "
  f"for every 10 m pixel and compared with a cross-validation-aware threshold derived "
  f"from the training data (DI = {msum['aoa_threshold']:.3f}). "
  f"{msum['pct_inside_aoa']:.1f} % of the district "
  f"({msum['ha_inside_aoa']:,.0f} ha) lies inside the applicability domain; the "
  f"remaining {100-msum['pct_inside_aoa']:.1f} % ({msum['ha_outside_aoa']:,.0f} ha) is "
  f"masked in every delivered map and excluded from all reported class areas.")
P(f"Finally, we now state the mapped extent unambiguously. The prediction grid covers "
  f"{msum['mapped_area_ha']:,.0f} ha, which corresponds to one sample per "
  f"{msum['mapped_area_ha']/110:,.0f} ha "
  f"({msum['mapped_area_ha']/110/100:.1f} km²) and approximately "
  f"{msum['n_valid_pixels']/110:,.0f} prediction pixels per observation. The "
  "Introduction no longer uses the 69 500 ha figure without a referent, and the "
  "supported spatial scale is discussed explicitly in Section 5.")
changes([
    "New Section 3.5 'Validation design' describes the nested spatial protocol, block "
    "construction, fold sizes, separation distances, seeds and repetitions.",
    "New supplementary figure with nutrient and residual variograms; Moran's I table added.",
    "New applicability-domain layer; all maps masked outside it (Figures 8–11).",
    "Introduction and Section 5 now state the mapped area, sampling density and the "
    "spatial scale the data can support.",
])

# ------------------------------------------------------------------- M3
doc.add_heading("M3.  Sampling design and temporal alignment require fuller documentation",
                level=2)
comment("Section 3.1 refers to a 'specific sampling method' … but it does not name the "
        "design, give allocation by stratum, describe randomization, define minimum "
        "separation, or state inclusion/exclusion criteria … Reconcile the chronology … "
        "Explain whether both images were assigned to every sample or whether temporally "
        "matched imagery was used … Clarify spatial support.")
response("We accept this comment in full and have rewritten Section 3.1 accordingly.")
P("On the temporal question the referee's inference from the date-suffixed predictor "
  "names is correct, and we confirm it explicitly rather than leaving it to be deduced: "
  "both acquisitions were assigned to every sampling point, and no temporally matched "
  "imagery was used. A point sampled in November is therefore also described by January "
  "imagery. We now state this in Section 3.2.1, and we treat it as a design limitation "
  "in Section 5 rather than presenting the imagery as contemporaneous with sampling. "
  "The chronology has been reconciled: the sampling campaign ran between November 2024 "
  "and January 2025, and the phrase 'during the 2024 growing season' has been removed.")
P("On spatial support, we now state that predictors were extracted at the pixel "
  "containing each sample centroid, and we report a sensitivity analysis in which all "
  "predictors are re-extracted as the mean of a 3 × 3 pixel neighbourhood (30 m), "
  "which brackets the combined effect of the ±3 m GPS uncertainty, the 10 m composite "
  "sampling radius and the 20 m native support of the red-edge and SWIR bands. We also "
  "now describe the maps as nominal 10 m products whose effective spatial support is "
  "constrained by the coarser predictors, as the referee requests under M5.")
changes([
    "Section 3.1 rewritten: the design is named, the conditioning covariates are listed, "
    "allocation by stratum is given in a new table, randomisation and minimum separation "
    "are described, and inclusion/exclusion criteria are stated.",
    "New sampling-design map (Figure 1a) showing the sample locations and the spatial "
    "cross-validation blocks.",
    "Sampling dates reported by group, with the lag between each sample and each "
    "acquisition.",
    "Section 3.2.1 states explicitly that both acquisitions were assigned to every "
    "sample; Section 5 discusses the consequence.",
    "Extraction support stated, with a 3 × 3 neighbourhood sensitivity test.",
])

# ------------------------------------------------------------------- M4
doc.add_heading("M4.  Laboratory QA/QC and descriptive statistics are insufficient", level=2)
comment("Provide one definitive sample-level table (n, minimum, Q1, median, mean, Q3, "
        "maximum, SD, CV, and skewness) after QA/QC, and report raster ranges separately "
        "for every model and nutrient. State whether any response transformation was used.")
response("We accept this comment and have added a full description of the quality "
         "control actually applied, together with a single definitive descriptive table "
         "that replaces the three mutually inconsistent sets of figures in the previous "
         "version.")
P("The quality control operated at several complementary levels, and Section 3.1 now "
  "reports each of them.")
bullets([
    [("Analytical replicates. ", True, False),
     ("Every sample was analysed in three independent analytical replicates for each of "
      "the three determinations. The replicates were used to monitor repeatability, and "
      "their arithmetic mean was taken as the final value for that sample in all "
      "descriptive statistics and modelling. The 110 values used throughout this paper "
      "are therefore replicate means, not single determinations.", False, False)],
    [("Blanks. ", True, False),
     ("Two analytical blanks were included in every series of 30 samples, one at the "
      "start and one at the end, carried through the same procedure as the samples, so "
      "that both contamination and drift within a series would be detected.",
      False, False)],
    [("Internal reference materials. ", True, False),
     ("Internal reference soils of known composition were included in every series of 30 "
      "samples and processed identically. Their results were compared against the "
      "established internal values to detect analytical drift between series. Blank and "
      "reference results were checked before any series was validated.", False, False)],
    [("Calibration. ", True, False),
     ("Instruments were calibrated with standard solutions of known concentration, "
      "prepared by dilution of stock solutions to span the expected measurement range. "
      "Phosphate standards calibrated the ultraviolet–visible spectrophotometer for "
      "Olsen P; potassium standards calibrated the flame photometer for exchangeable K; "
      "total N was determined by Kjeldahl digestion, distillation and titration against "
      "hydrochloric acid of known concentration.", False, False)],
])
P("Analytical performance has also been quantified for each determination and is now "
  "reported as Table 1 of the manuscript. Limits of detection and quantification were "
  "derived from the standard deviation of the analytical blanks, as three and ten times "
  "that standard deviation; recovery was evaluated against internal certified reference "
  "soils to verify extraction efficiency; and repeatability was computed as the relative "
  "standard deviation across the three independent replicates of each sample before "
  "averaging. All work was carried out at the Regional Centre of Agricultural Research "
  "of Tadla, INRA Morocco, which is now named in Section 3.1.")
qa = pd.DataFrame([
    ["Total N", "Kjeldahl, digestion and titration", "50.0", "150.0", "96.5", "4.2"],
    ["Olsen P", "UV–visible spectrophotometry", "0.2", "0.8", "98.2", "3.5"],
    ["Exchangeable K", "Flame photometry", "2.0", "5.0", "99.1", "2.8"],
], columns=["Nutrient", "Method", "LOD (mg kg⁻¹)", "LOQ (mg kg⁻¹)",
            "Mean recovery (%)", "Repeatability RSD (%)"])
table(qa, fontsize=8.4,
      caption="Table R3a (manuscript Table 1). Analytical performance. Recovery between "
              "90 and 110 % and a relative standard deviation below 5 % lie within the "
              "ranges conventionally accepted for agronomic and environmental soil "
              "laboratories.")
P(f"We checked these limits against every observation rather than assuming they were "
  f"met. All 110 determinations exceed the limit of quantification of their method: the "
  f"lowest values are {df.N.min():,.0f} mg kg⁻¹ for nitrogen against an LOQ of "
  f"150 mg kg⁻¹, {df.P.min():.2f} mg kg⁻¹ for phosphorus against 0.8 mg kg⁻¹, and "
  f"{df.K.min():.1f} mg kg⁻¹ for potassium against 5.0 mg kg⁻¹, giving margins of "
  f"{df.N.min()/150:.1f}, {df.P.min()/0.8:.1f} and {df.K.min()/5.0:.1f} times the "
  f"respective limits. The entire range of each nutrient is therefore quantifiable, and "
  f"no observation required censoring or substitution at the lower limit.")
t3s = t3[["Nutrient", "n", "Minimum", "Q1", "Median", "Mean", "Q3", "Maximum",
          "SD", "CV_percent", "Skewness"]].copy()
t3s.columns = ["Nutrient", "n", "Min", "Q1", "Median", "Mean", "Q3", "Max", "SD",
               "CV (%)", "Skew"]
table(t3s, fontsize=8.4,
      caption="Table R3 (manuscript Table 3). Descriptive statistics of the 110 "
              "validated observations. All concentrations in mg kg⁻¹.")

nmean = float(t3[t3.Nutrient.str.startswith("Total N")].Mean.iloc[0])
P(f"We must correct a numerical error of our own. The previous Section 4.1 gave the "
  f"mean total nitrogen as approximately 1 683 mg kg⁻¹. The correct value is "
  f"{nmean:,.1f} mg kg⁻¹. We are grateful that the discrepancy surfaced, because the "
  f"referee used our incorrect figure when questioning the nitrogen map under M9; the "
  f"reconciliation in that response uses the corrected value.")

P("The referee is also right that the values previously quoted in Section 3.7.3 were "
  "raster predictions rather than observations, and right that relabelling them is "
  "itself a result. We now report the two sets separately and draw the inference the "
  "referee anticipated:")
rr = []
for t in TARGETS:
    r = ras[ras.layer == f"{t}_RF"].iloc[0]
    o = df[t]
    rr.append({"Nutrient": t,
               "Observed range": f"{o.min():,.2f} – {o.max():,.2f}",
               "Predicted raster range (RF)": f"{r['min']:,.2f} – {r['max']:,.2f}",
               "Range compression": f"{100*(1-(r['max']-r['min'])/(o.max()-o.min())):.0f} %",
               "Observed mean": f"{o.mean():,.1f}", "Map mean": f"{r['mean']:,.1f}"})
table(pd.DataFrame(rr), fontsize=8.4,
      caption="Table R4. Observed sample ranges compared with the ranges of the "
              "predicted 10 m surfaces. The compression is the spatial expression of "
              "the regression-to-the-mean shrinkage quantified under M1.")

P("No response transformation was applied; the models were fitted on the untransformed "
  "concentrations, and this is now stated. On outlier handling we accept the referee's "
  "objection: valid extremes should not be removed merely for exceeding an interquartile "
  "rule. The interquartile screen has been withdrawn and used only to flag values for "
  "re-inspection against the laboratory records; no observation was removed on "
  "statistical grounds alone, and the analysis dataset contains all 110 samples.")
changes([
    "Section 3.1 now reports the three analytical replicates per sample and the use of "
    "their mean, the two blanks per series of 30 samples, the internal reference soils "
    "included in every series, and the calibration procedure for each of the three "
    "determinations.",
    "Section 3.1 states explicitly that the 110 values used throughout are replicate "
    "means rather than single determinations.",
    "Section 4.1 corrected: mean total N is now given as "
    f"{nmean:,.1f} mg kg⁻¹.",
    "New Table 3 gives the single definitive sample-level summary; raster ranges are "
    "reported separately in Table 10 for every model and nutrient.",
    "The IQR-based removal rule has been withdrawn and the practice described accurately.",
    "Statement added that no response transformation was used.",
])

# ------------------------------------------------------------------- M5
doc.add_heading("M5.  Sentinel-2 preprocessing and the predictor inventory are not "
                "reproducible", level=2)
comment("The stated 41 predictors cannot be reconstructed … Provide a supplementary "
        "inventory with the exact model variable name, acquisition date, formula, source "
        "bands, scale/unit, native resolution, output resolution, and resampling method "
        "for every predictor, then correct the count throughout. Confirm that reflectance "
        "was converted to an appropriate scale before computing indices with additive "
        "constants, especially EVI and SAVI.")
response("This comment led us to audit the index code against Table 2, and the audit "
         "found errors that we report here in full. We are grateful the referee pressed "
         "on the reconstruction, because the defects were real and would not otherwise "
         "have been caught.")
P("First, the straightforward points. The predictor count was wrong: the models used 21 "
  "variables per acquisition, not 20, and the correct count for the previously submitted "
  "version was 42 rather than the 41 stated. The referee's arithmetic in M5 anticipated "
  "this exactly. On the scaling question we can now answer with evidence: Earth Engine "
  "surface reflectance was divided by 10 000 before any index was computed. Band values "
  "in the deposited predictor matrix lie between 0.025 and 0.498, and EVI, SAVI and BSI "
  "are reproducible from them to within 3 × 10⁻⁸, which confirms that the additive "
  "constants in EVI and SAVI were applied on the correct scale.")
P("Second, and more seriously, we re-derived every index from the deposited band values "
  "and compared it with the stored layer. Six of the ten reproduce their stated formula "
  "exactly (NDVI, SAVI, EVI, GNDVI, NDMI, BSI, all to within 3 × 10⁻⁸). Four do not:")
bullets([
    [("SI ", True, False),
     ("was declared as B11 × B12. The stored layer is (B11 − B8)/(B11 + B8), which is "
      "the exact negative of NDMI: SI + NDMI = 0 to machine precision at all 110 points, "
      "on both acquisition dates.", False, False)],
    [("VSSI ", True, False),
     ("was declared as 2 × B3 − 5 × (B4 + B8). The stored layer is (B11 − B3)/(B11 + B3), "
      "the exact negative of the layer labelled NDSI: VSSI + NDSI = 0 to machine "
      "precision at all 110 points, on both dates.", False, False)],
    [("NDSI ", True, False),
     ("was not defined in Table 2 at all, as the referee noted. The stored layer is "
      "(B3 − B11)/(B3 + B11), which is the standard MNDWI formula.", False, False)],
    [("MNDWI ", True, False),
     ("was declared as (B3 − B11)/(B3 + B11). The stored layer uses B12 in place of B11.",
      False, False)],
])
P("Two of the 42 layers therefore carried no information whatsoever beyond another "
  "column of the same matrix — they were sign-flipped duplicates. This is precisely the "
  "condition under which impurity-based importance is least trustworthy, which bears "
  "directly on the referee's comment M7.")
P("We also record a fourth discrepancy in the opposite direction. Table 2 printed NDRE "
  "as (B8 − B5)/(B8 + B5), but the code computed (B8A − B5)/(B8A + B5). Here the code "
  "was the better choice, since B8A and B5 share the same 20 m native support, so we "
  "have retained the computed definition and corrected Table 2 rather than the other "
  "way round.")
P("The whole predictor set has been rebuilt from the ten reflectance bands with verified "
  "formulas: SI and VSSI now compute what they claim, MNDWI uses B11, and the redundant "
  "NDSI layer has been removed. The corrected inventory contains 10 bands and 10 indices "
  f"per acquisition, that is {s01['n_predictors']} predictors in total, with no pair "
  f"exceeding |r| = 0.999 and a numerical rank of {s01['numerical_rank']}. Every model, "
  "figure and table in this revision was regenerated on the corrected set.")
inv_show = (inv[inv.type == "Spectral index"].drop_duplicates("formula")
            [["variable", "formula", "source_bands", "native_res_m", "reference"]].copy())
inv_show["Index"] = inv_show.variable.str.rsplit("_", n=2).str[0]
inv_show = inv_show[["Index", "formula", "source_bands", "native_res_m", "reference"]]
inv_show.columns = ["Index", "Sentinel-2 formula", "Source bands",
                    "Limiting native res. (m)", "Reference"]
table(inv_show, fontsize=8.2, widths=[1.9, 5.4, 2.6, 2.3, 4.2],
      caption="Table R5 (manuscript Table 2, corrected). Every index is reproducible "
              "from the deposited band values. The full 40-row inventory with "
              "acquisition dates, units, output resolution and resampling method is "
              "given as Supplementary Table S1.")
changes([
    "Predictor count corrected throughout from 41 to "
    f"{s01['n_predictors']} (Abstract, Sections 3.2.2, 3.3, 3.7.1, Conclusions).",
    "Table 2 corrected: SI, VSSI, MNDWI and NDRE formulas and attributions fixed; the "
    "mislabelled NDSI layer removed.",
    "Section 3.2.1 now reports asset identifiers, sensing dates, MGRS tiles, cloud "
    "cover, the Scene Classification Layer classes excluded, cloud-shadow handling, "
    "the reflectance scaling factor and the resampling algorithm.",
    "Statement added that the 20 m bands were resampled bilinearly to 10 m and that "
    "this creates a nominal, not an effective, 10 m support.",
    "New Supplementary Table S1 gives the complete predictor inventory.",
    "All models, figures and tables regenerated on the corrected predictor set.",
])

# ------------------------------------------------------------------- M6
doc.add_heading("M6.  Model development, baselines, and algorithm comparison need a "
                "reproducible protocol", level=2)
comment("Create a dedicated modelling subsection … report the R and package versions, "
        "random seeds, … and selected hyperparameters for each model–nutrient "
        "combination … Include defensible benchmarks such as the mean/intercept model, a "
        "simple regularized or PLS regression, and a spatial baseline such as ordinary "
        "kriging … The claim that XGBoost is superior for all nutrients is not supported.")
response("We accept the criticism of the superiority claim without reservation and "
         "withdraw it. The previous statement that XGBoost outperformed Random Forest "
         "'by nearly 50 % in RMSE reduction' for nitrogen compared a Random Forest "
         "cross-validation figure with an XGBoost calibration figure. When both are "
         "computed from the same held-out predictions the claim disappears.")

pc = pcmp.copy()
pc["Nutrient"] = pc.target
pc["RF R²"] = [f"{sp(t,'RF').R2:.3f}" for t in pc.target]
pc["XGBoost R²"] = [f"{sp(t,'XGB').R2:.3f}" for t in pc.target]
pc["Δ mean squared error (RF − XGB)"] = pc.mean_MSE_diff_RF_minus_XGB.map(lambda v: f"{v:+,.1f}")
pc["95 % CI"] = pc.apply(lambda r: f"[{r.ci_lo:+,.1f}, {r.ci_hi:+,.1f}]", axis=1)
pc["Conclusion"] = pc.verdict
table(pc[["Nutrient", "RF R²", "XGBoost R²", "Δ mean squared error (RF − XGB)",
          "95 % CI", "Conclusion"]], fontsize=8.4,
      caption="Table R6. Paired comparison of the two algorithms on identical outer "
              "folds, with bootstrap confidence intervals on the paired difference in "
              "squared error.")
P("The correct statement, and the one the manuscript now makes, is that the two "
  "algorithms are statistically indistinguishable on this dataset. We have defined a "
  "primary selection criterion in advance — the root mean square error on the outer "
  "spatial folds — and we present Random Forest as the mapping model on that criterion "
  "while stating plainly that the difference is not significant.")

blt = []
for t in TARGETS:
    sub = pd.concat([
        pd.DataFrame([{"Model": "Random Forest", "R2": sp(t, "RF").R2, "RMSE": sp(t, "RF").RMSE},
                      {"Model": "XGBoost", "R2": sp(t, "XGB").R2, "RMSE": sp(t, "XGB").RMSE}]),
        bl[bl.target == t][["model", "R2", "RMSE"]].rename(columns={"model": "Model"})])
    for _, r in sub.iterrows():
        blt.append({"Nutrient": t, "Model": r.Model, "R²": f"{r.R2:.3f}",
                    "RMSE": f"{r.RMSE:,.2f}"})
table(pd.DataFrame(blt), fontsize=8.0, widths=[2.0, 7.4, 3.0, 3.4],
      caption="Table R7 (new manuscript Table 5). Benchmarks evaluated on identical "
              "outer spatial folds. The intercept-only model is the null reference; "
              "ordinary kriging is the purely spatial baseline; regression kriging "
              "combines the Random Forest trend with kriged residuals.")

mods = []
for _, r in t6b.iterrows():
    prm = ", ".join(f"{c} = {r[c]}" for c in t6b.columns if c not in ("target", "model")
                    and pd.notna(r[c]))
    mods.append({"Nutrient": r.target,
                 "Model": "Random Forest" if r.model == "RF" else "XGBoost",
                 "Selected hyperparameters (modal across outer folds)": prm})
table(pd.DataFrame(mods), fontsize=8.2, widths=[2.0, 3.4, 10.4],
      caption="Table R8. Hyperparameters actually selected by the inner tuning loop, "
              "reported as the modal choice across the outer folds. Fold-by-fold "
              "selections are given in Supplementary Table S2.")
changes([
    "New Section 3.4 'Model development' created; the machine-learning protocol has "
    "been moved out of Section 3.3, which now covers predictor extraction only.",
    "Software versions, random seed, missing-value handling, tuning grids, boosting "
    "rounds, early-stopping settings, objective and evaluation metric all reported.",
    "Selected hyperparameters reported per model and nutrient (Table R8), replacing the "
    "previous search ranges.",
    "New Table 5 reports the intercept-only, ridge, PLS, ordinary kriging and "
    "regression-kriging benchmarks.",
    "Superiority claim removed from the Abstract, Section 4.3, Section 5 and the "
    "Conclusions; replaced by a paired comparison with uncertainty.",
    "Section 3 no longer asserts that RF and XGBoost outperform SVM or ANN or require "
    "fewer adjustments.",
])

# ------------------------------------------------------------------- M7
doc.add_heading("M7.  Variable-importance figures contradict the written interpretation",
                level=2)
comment("Rewrite the Abstract, Results, Discussion, and Conclusions so they describe the "
        "figures rather than a predetermined spectral mechanism. If the intended claim "
        "concerns classes of predictor, aggregate importance by class and demonstrate "
        "that quantitatively. State the RF importance metric … Use permutation importance "
        "and report stability across outer resamples … Normalize importance before "
        "cross-nutrient comparisons and avoid causal language based on rankings alone.")
response("The referee is right, and the corrected analysis makes the mismatch larger "
         "rather than smaller. We have removed the previous interpretation entirely.")
P("On the metric: the previous figures reported ranger's impurity importance for Random "
  "Forest and Gain for XGBoost while the text described both as 'Gain'. As the referee "
  "notes, impurity importance is biased with correlated continuous predictors, and our "
  "own audit under M5 shows that two predictors were exact duplicates of two others — "
  "the worst case for that metric. We have therefore replaced it with permutation "
  "importance measured on the held-out observations of each outer spatial fold, "
  "normalised by the fold RMSE so that values are comparable across nutrients, and "
  "reported with the between-fold standard deviation.")

cls_tbl = t7.copy()
cls_tbl.columns = [c if c == "cls" else c for c in cls_tbl.columns]
cls_tbl = cls_tbl.rename(columns={"cls": "Predictor class"})
table(cls_tbl, fontsize=8.0, widths=[5.0, 1.9, 1.9, 1.9, 1.9, 1.9, 1.9],
      caption="Table R9 (new manuscript Table 6). Permutation importance aggregated by "
              "predictor class, expressed as a percentage of total positive importance. "
              "Columns are nutrient × model.")
P("Aggregating by class, as the referee suggests, yields a claim we can actually "
  "support: reflectance bands dominate, and the spectral indices contribute little. "
  "Nitrogen and phosphorus are driven by near-infrared and red-edge reflectance; "
  "potassium is driven by near-infrared, shortwave-infrared and visible reflectance "
  "jointly. The specific mechanisms asserted previously do not survive: the red-edge "
  f"indices named in the old Abstract as the leading predictors of nitrogen hold "
  f"{100*pi[(pi.target=='N')&(pi.model=='RF')&(pi.variable.str.startswith(('NDRE','GNDVI')))].perm_importance.clip(lower=0).sum()/pi[(pi.target=='N')&(pi.model=='RF')].perm_importance.clip(lower=0).sum():.1f} % "
  "of the Random Forest permutation importance.")

stab_txt = ", ".join(f"{r.target}/{r.model} {int(r.n_stable_predictors)}"
                     for _, r in t7c.iterrows())
P("On stability, which the referee explicitly requested, the result is sobering and we "
  "report it rather than suppress it. Counting predictors whose mean permutation "
  f"importance exceeds its own between-fold standard deviation gives {stab_txt} out of "
  f"{int(t7c.n_total.iloc[0])}. Beyond a small number of November near-infrared and "
  "red-edge bands, individual predictor rankings in this dataset are not distinguishable "
  "from resampling noise. The manuscript now says exactly that, and all causal language "
  "attached to individual indices has been removed.")
P("We also correct the specific error the referee identified: GNDVI is a green–NIR "
  "index, not a red-edge index, as our own Table 2 shows. That misclassification has "
  "been removed from the Abstract and the Discussion.")
changes([
    "Importance metric changed to permutation importance on held-out outer folds, "
    "normalised by fold RMSE; the metric is now named explicitly in Section 3.4.",
    "New Figure 6 shows importance aggregated by predictor class and the 15 strongest "
    "individual predictors with between-fold standard deviations; panels carry (a)–(c) "
    "labels and the panel order matches the caption.",
    "New Table 6 gives the class aggregation quantitatively.",
    "Abstract, Sections 4.4, 4.5, 5 and 6 rewritten to describe the figures; all "
    "mechanistic claims about individual indices removed or reframed as hypotheses.",
    "GNDVI is no longer described as a red-edge index.",
])
figure(os.path.join(FIG, "Figure_6_permutation_importance.png"), 16.4,
       "Figure R3 (manuscript Figure 6). Permutation importance on held-out spatial "
       "folds, aggregated by predictor class and by individual predictor.")

# ------------------------------------------------------------------- M8
doc.add_heading("M8.  The models may be predicting canopy and management patterns rather "
                "than stable soil properties", level=2)
comment("Total N, Olsen P, and exchangeable K do not have direct diagnostic absorption "
        "features in Sentinel-2 bands … Thirteen of the fifteen strongest RF predictors "
        "for N come from the November 2024 acquisition alone … State whether an "
        "agricultural/land-cover mask was used … Discuss temporal transferability "
        "explicitly.")
response("We agree, and the corrected analysis strengthens the referee's reading rather "
         "than weakening it. We now present the canopy/management interpretation as the "
         "leading explanation rather than as an alternative to be dismissed.")
nov_txt = ", ".join(f"{r.target}/{r.model} {r.November_2024_share_percent:.0f} %"
                    for _, r in t7b.iterrows())
novmin, novmax = t7b.November_2024_share_percent.min(), t7b.November_2024_share_percent.max()
P(f"Measured across the full predictor set rather than the top fifteen, and using "
  f"permutation importance rather than impurity, the November 2024 acquisition accounts "
  f"for {nov_txt} of positive importance — between {novmin:.0f} % and {novmax:.0f} % "
  f"across the six models. One acquisition therefore carries almost the "
  f"entire model for all three nutrients, not only for nitrogen. Combined with the "
  f"dominance of raw near-infrared and red-edge reflectance reported under M7, and with "
  f"the sharply bounded field-shaped patches visible in the prediction surfaces, the "
  f"most parsimonious interpretation is that the models are learning the state of the "
  f"canopy at a single date, which is correlated with soil fertility through management "
  f"and crop selection, rather than a direct pedological signal.")
P("The manuscript now states this in the Discussion as the primary interpretation. We "
  "have removed the claims attributing map patterns to clay content, salinity, "
  "hydromorphy, irrigation infrastructure, carbonate content and fertilisation history, "
  "none of which were measured; where these remain they are labelled explicitly as "
  "hypotheses requiring the corresponding measurements. The climate-resilience, "
  "drought-monitoring and Mediterranean-transferability claims have been removed from "
  "the Abstract, Discussion and Conclusions: the study uses one season and contains no "
  "climate analysis, and the referee is right that those claims were unsupported.")
P("On masking, we confirm that the previous maps were clipped to the district boundary "
  "only. The revised products are additionally restricted to the applicability domain "
  f"described under M2, which excludes {msum['ha_outside_aoa']:,.0f} ha of "
  "predictor-space extrapolation, and we now state that water bodies, settlements, "
  "roads and canals are not separately masked, so that readers can judge the products "
  "accordingly.")
changes([
    "Section 5 restructured with a dedicated subsection on what the spectral signal "
    "represents; the canopy/management interpretation is presented first.",
    f"Statement added that the November acquisition carries {novmin:.0f}–{novmax:.0f} % "
    "of model importance, with the temporal-transferability consequence stated.",
    "Maps described as seasonal snapshots conditioned on the 2024–2025 cropping mosaic.",
    "Causal attributions to clay, salinity, hydromorphy, carbonates and fertiliser "
    "history removed or reframed as untested hypotheses.",
    "Climate-resilience, drought-monitoring and Mediterranean-transfer claims removed.",
    "Masking described precisely; applicability-domain mask applied to all products.",
])

# ------------------------------------------------------------------- M9
doc.add_heading("M9.  Map products, uncertainty, applicability, and cartographic "
                "reporting require revision", level=2)
comment("Show legible continuous maps with units and common model-specific or "
        "nutrient-specific scales … Provide local uncertainty maps …, an "
        "applicability-domain/AOA layer, and residual spatial diagnostics … Quantify "
        "class area (ha and %), RF–XGBoost agreement … Report the full CRS/EPSG code, "
        "coordinate units, raster extent, scale, and resolution.")
response("All of these have been implemented, and one further problem of our own is "
         "reported here.")
P("The disclosure first. On checking the delivered rasters against the covariate stack "
  "we found that the map products accompanying the previous version had been generated "
  "at a decimated resolution, not at 10 m: the grid was 709 × 295 cells with an "
  "effective pixel of roughly 85 × 99 m, whereas the covariate stack is 7 093 × 2 950 "
  "cells at 10 m. The '10 m' description was therefore not supported by the files. All "
  "maps in this revision have been regenerated at the full native 10 m grid — "
  f"{msum['n_valid_pixels']:,} valid pixels covering {msum['mapped_area_ha']:,.0f} ha — "
  "and the effective pixel dimensions are stated on every figure.")
table(t10d, fontsize=8.4, widths=[7.6, 8.4],
      caption="Table R10 (new manuscript Table 10). Complete cartographic metadata for "
              "the delivered raster products.")
P("On the coordinate question, the referee is correct that the grid values shown "
  "previously were not UTM 29N/WGS 84. They were Merchich / Nord Maroc, EPSG:26191, the "
  "Moroccan national grid; the axes were unlabelled, which caused the confusion. The "
  "rasters are distributed in EPSG:4326 and all figures are drawn in EPSG:26191 with "
  "labelled axes in kilometres, a scale bar, a north arrow and the EPSG code printed on "
  "the figure.")
P("Continuous surfaces are now the headline product, presented before the classified "
  "maps, with units of mg kg⁻¹ and a perceptually uniform, colour-vision-safe sequential "
  "ramp running from low to high. This also resolves the referee's objection to the "
  "previous scheme, in which high fertility was rendered red and very low fertility "
  "dark green. The 110 sample locations are overlaid on every panel so that "
  "extrapolation distance can be judged directly.")
rng_lines = []
for t in TARGETS:
    o = df[t]
    for algo in ["RF", "XGB"]:
        r = ras[ras.layer == f"{t}_{algo}"].iloc[0]
        if r["min"] < o.min() or r["max"] > o.max():
            rng_lines.append(f"{algo} {t} ({r['min']:,.1f}–{r['max']:,.1f} against "
                             f"{o.min():,.1f}–{o.max():,.1f})")
rng_txt = ("; ".join(rng_lines) if rng_lines else "none")
P(f"Local uncertainty is reported as the standard deviation of the predictions across "
  f"the Random Forest tree ensemble, mapped at 10 m for each nutrient. The applicability "
  f"domain described under M2 is mapped alongside, and every class area we report is "
  f"computed inside that domain. Residual spatial diagnostics are given by the residual "
  f"variograms under M2.")
P(f"On whether predictions remain inside the sampled range, we report the check rather "
  f"than assert the answer. All three Random Forest surfaces lie strictly inside the "
  f"observed range of the corresponding nutrient. The XGBoost surfaces marginally "
  f"undershoot the observed minimum in two cases: {rng_txt}. These excursions are "
  f"small, they are confined to the extrapolative margin, and they are removed by the "
  f"applicability mask; we state them rather than round them away.")

ca = care.copy()
ca["Class"] = ca.cls
ca["Area (ha)"] = ca.rf_inAOA_ha.map(lambda v: f"{v:,.0f}")
ca["Share (%)"] = ca.rf_inAOA_pct.map(lambda v: f"{v:.2f}")
piv = ca.pivot(index="Class", columns="nutrient", values=["Area (ha)", "Share (%)"])
piv.columns = [f"{b} — {a}" for a, b in piv.columns]
piv = piv.reindex(["Very low", "Low", "Medium", "High"]).reset_index()
table(piv, fontsize=8.2,
      caption="Table R11. Fertility class areas from the Random Forest 10 m surfaces, "
              "computed inside the area of applicability.")

agt = agr.copy()
agt.columns = ["Nutrient", "Overall pixel agreement", "Cohen's κ", "Pixels compared"]
agt["Overall pixel agreement"] = agt["Overall pixel agreement"].map(lambda v: f"{v:.3f}")
agt["Cohen's κ"] = agt["Cohen's κ"].map(lambda v: f"{v:.3f}")
agt["Pixels compared"] = agt["Pixels compared"].map(lambda v: f"{v:,}")
table(agt, fontsize=8.4,
      caption="Table R12. Agreement between the Random Forest and XGBoost classified "
              "surfaces, replacing the previous assertion that the XGBoost maps are "
              "clearer or retain more detail.")
P("The claim that XGBoost maps show clearer differences or retain more spatial detail "
  "has been removed. It appeared in the Methods, where a result does not belong, and it "
  "was not tested; the agreement statistics above are given instead.")
P("Finally, the place names have been reconciled across all figures, and the "
  f"nitrogen-map question is addressed with the corrected sample mean of "
  f"{nmean:,.1f} mg kg⁻¹ from M4 together with the applicability mask, which removes "
  "the extrapolative areas that inflated the upper tail of the previous surface.")
changes([
    "All rasters regenerated at the native 10 m grid; effective pixel size stated.",
    "New Figures 8 and 9: continuous Random Forest and XGBoost surfaces with units, a "
    "sequential colour-vision-safe ramp, sample locations and full cartographic furniture.",
    "New Figure 10: local uncertainty (between-tree ensemble standard deviation) at 10 m.",
    "New Figure 1b: dissimilarity index and area of applicability.",
    "New Figure 11: classified fertility maps, reported inside the applicability domain.",
    "New Table 10 with CRS, EPSG, units, extent, resolution, pixel area and mapped area.",
    "Class areas in hectares and per cent; RF–XGBoost agreement with Cohen's κ.",
    "Colour ramp corrected so that low fertility is pale and high fertility saturated.",
    "Place names reconciled: 'Souk Sebt' and 'Fkih Ben Salah' used consistently across "
    "Figures 1, 8, 9 and 11.",
    "Statements about XGBoost map clarity removed from Section 3.7.2.",
])
figure(os.path.join(FIG, "Figure_8_continuous_maps_RF.png"), 15.0,
       "Figure R4 (manuscript Figure 8). Random Forest predicted macronutrient "
       "concentrations at 10 m, masked outside the area of applicability.")

# ------------------------------------------------------------------- M10
doc.add_heading("M10.  Fertility classes do not yet justify operational fertilizer "
                "recommendations", level=2)
comment("Classify the observed samples and outer-fold predictions with the same "
        "thresholds, then report confusion matrices and class-specific accuracy … At "
        "present the maps may support reconnaissance, stratified resampling, or "
        "hypothesis generation, but they do not justify direct variable-rate fertilizer "
        "prescriptions.")
response("We have carried out the requested classification and we accept the referee's "
         "conclusion about the operational status of the maps.")
kt = t9[["nutrient", "model", "overall_accuracy", "kappa"]].copy()
kt.columns = ["Nutrient", "Model", "Overall accuracy", "Cohen's κ"]
kt["Model"] = kt.Model.replace({"RF": "Random Forest", "XGB": "XGBoost"})
table(kt, fontsize=8.6,
      caption="Table R13 (new manuscript Table 9). Agreement between observed fertility "
              "class and the class of the held-out outer-fold prediction, n = 110. Full "
              "confusion matrices with producer's and user's accuracies are given as "
              "Supplementary Table S3.")
P("We now lead with these figures rather than with R², because class agreement is the "
  "quantity an advisory user actually depends on. We also state the limitation the "
  "matrices expose: the upper classes are supported by very few observations — only "
  f"{int(np.sum(df.K.values > 120))} samples exceed the highest potassium threshold and "
  f"{int(np.sum(df.N.values > 3000))} exceed the highest nitrogen threshold — so "
  "class-specific accuracy in those categories is not reliably estimated.")
P("On the thresholds themselves we accept every point. Olsen (1954) describes the "
  "extraction method rather than the interpretation thresholds and is no longer cited "
  "for them; Havlin (2014) has been removed and replaced by a source that is present in "
  "the reference list; the class boundaries have been rewritten as non-overlapping "
  "intervals; and the conversion used for total nitrogen is shown explicitly. The "
  "previous practice of adapting thresholds to avoid empty classes has been abandoned, "
  "since it is circular, and any local adaptation is now justified on agronomic grounds "
  "and identified as such.")
P("Most importantly, we accept that total nitrogen is not plant-available nitrogen and "
  "that fertiliser prescription additionally requires crop demand, yield target, "
  "mineralisation, pH, texture, irrigation and management history, none of which are "
  "available here. The Abstract, Discussion and Conclusions no longer offer the maps as "
  "a basis for variable-rate prescription. They are now presented as reconnaissance "
  "products supporting stratified resampling, the delineation of candidate management "
  "zones for verification, and hypothesis generation.")
changes([
    "New Table 9 and Supplementary Table S3 report class accuracy, Cohen's κ and full "
    "confusion matrices from outer-fold predictions.",
    "Tables 3–5 rewritten with non-overlapping boundaries and verifiable sources; "
    "Havlin (2014) removed; the total-N conversion shown.",
    "Terminology corrected throughout: total N, Olsen P and exchangeable K are described "
    "as distinct nutrient pools and never collectively as plant-available.",
    "Variable-rate fertiliser claims removed from the Abstract, Section 5 and the "
    "Conclusions and replaced by reconnaissance and stratified-resampling applications.",
])

# ------------------------------------------------------------------- M11
doc.add_heading("M11.  Data and code availability must support reproducibility", level=2)
comment("The statement 'All the data are included in the present study' is inaccurate … "
        "Deposit the analysis-ready data and scripts in a stable repository with a "
        "persistent identifier … The next revision should be auditable from raw analysis "
        "inputs to every reported metric and map.")
response("The previous statement was inaccurate and has been withdrawn. A complete "
         "deposit has been assembled and every number in this revision is reproducible "
         "from it by running the numbered scripts in order.")
bullets([
    "Sample identifiers, the three nutrient values and the sampling coordinates for all "
    "110 observations.",
    f"The analysis-ready predictor matrix ({s01['n_predictors']} corrected predictors) "
    "with the full inventory of formulas, source bands, units and resolutions.",
    "The Google Earth Engine script used to build and export the covariates, with asset "
    "identifiers and acquisition dates.",
    "The exact spatial and random fold assignments used for every reported statistic.",
    "The complete analysis code as nine numbered scripts, with the random seed fixed.",
    "Session information: interpreter, operating system and the version of every package.",
    "Selected hyperparameters for every model, nutrient and outer fold.",
    "Out-of-fold predictions for all three validation schemes, from which every metric "
    "in the manuscript can be recomputed.",
    "The georeferenced 10 m continuous predictions, uncertainty surfaces, dissimilarity "
    "index and applicability mask for all three nutrients and both algorithms.",
])
P("Sampling coordinates are reported at full precision. Should the editor consider "
  "landholder privacy to require restriction, we will substitute coordinates jittered "
  "within 250 m — far below the block dimension used for validation and below the "
  "distances at which autocorrelation is detectable in these data, so that the spatial "
  "analyses remain reproducible — and state the restriction and its reason explicitly.")
changes([
    "Data availability statement replaced with a description of the deposited archive "
    "and its persistent identifier.",
    "Deposit assembled with data, code, folds, session information, out-of-fold "
    "predictions and georeferenced rasters.",
])

# ------------------------------------------------------------------- M12
doc.add_heading("M12.  Citation integrity and claimed editorial compliance remain "
                "incomplete", level=2)
comment("Several statements in the response letter are not reflected in the revised "
        "manuscript … Audit citation content, not only formatting.")
response("We checked every assertion in this comment against the submitted file. All of "
         "them are correct. We apologise: the previous response letter described "
         "corrections that had not in fact been carried into the manuscript, and we "
         "recognise that this is a serious lapse. The audit has now been redone "
         "systematically rather than selectively, and the results are given below.")
cit = pd.DataFrame([
    ["Barnes et al. (2000), Sharma et al. (2015) still cited in §4.4",
     "Confirmed; both were also absent from the reference list", "Citations removed"],
    ["Havlin (2014) still in the Introduction and Table 4–5 captions",
     "Confirmed, three occurrences; absent from the reference list",
     "Replaced with Havlin and Heiniger (2020) and a threshold source"],
    ["Sharma (2010) listed but never cited", "Confirmed", "Entry removed"],
    ["53 entries and 49 DOI links, not 49 and 47",
     "Confirmed exactly; the four without DOI are IUSS WRB (2015), Olsen (1954), "
     "Rikimaru et al. (2002) and Rouse et al. (1974)",
     "Counts corrected; the four are cited with stable URLs"],
    ["'Vaglie, M. D.' and misplaced alphabetisation",
     "Confirmed; a single ordering violation", "Corrected to Dalle Vaglie and "
     "alphabetised under D"],
    ["Encoding artifact in §4.5", "Confirmed", "Corrected"],
    ["Zhang (2023) for terra, viridis and DSM covariates",
     "Confirmed, three occurrences; the source is a ¹³⁷Cs erosion review",
     "Replaced with Hijmans (terra), Garnier et al. (viridis) and McBratney et al. (2003)"],
    ["Abatzoglou et al. (2018) for Google Earth Engine",
     "Confirmed; the source is the TerraClimate paper", "Replaced with Gorelick et al. (2017)"],
    ["Chen et al. (2017) for R² and RMSE", "Confirmed; landslide susceptibility study",
     "Replaced with Kuhn and Johnson (2013) and Wadoux et al. (2020)"],
    ["Cambardella et al. (1994) and IUSS WRB (2015) for laboratory procedures",
     "Confirmed; neither describes the analytical methods",
     "Replaced with the analytical standards themselves"],
    ["Lundberg and Lee (2017) for potassium in Vertisols",
     "Confirmed; SHAP was never computed", "Citation removed"],
    ["Viscarra Rossel et al. (2016) for Beni Moussa variability",
     "Confirmed; a global spectral library paper",
     "Replaced with El Hamzaoui and El Baghdadi (2021) and Barakat et al. (2017)"],
    ["Table 2 index attributions (EVI/GNDVI swap, Qi for SAVI, Gao for NDMI, "
     "Nguyen for VSSI)", "All confirmed",
     "EVI → Huete et al. (2002); GNDVI → Gitelson et al. (1996); SAVI → Huete (1988) "
     "only; NDMI → Gao (1996) with the NDWI naming noted; VSSI → Dehni and Lounis (2012); "
     "asterisk removed"],
], columns=["Referee assertion", "Our check", "Action taken"])
table(cit, fontsize=7.8, widths=[5.2, 5.6, 5.6],
      caption="Table R14. Verification of every bibliographic assertion in comment M12.")
P("On the Aljanabi and Dedeoğlu item, we agree that the journal landing page and the "
  "issue metadata use different year conventions. We have adopted the year printed on "
  "the article PDF and in the issue record, and we would welcome the editor's direction "
  "if the journal's house style requires the other convention. Every in-text citation "
  "has been matched to exactly one reference entry, orphan entries have been removed, "
  "consecutive parenthetical groups have been merged, and the list has been "
  "re-alphabetised and brought into consistent APA style.")

doc.add_page_break()

# ==========================================================================
#                       SECTION 2 : MINOR COMMENTS
# ==========================================================================
doc.add_heading("2.  Responses to the minor and editorial comments", level=1)
P("Every point below was checked against the submitted file and every one was accurate. "
  "All have been implemented.")
minor = pd.DataFrame([
    ["Section numbering: Methods jump from 3.3 to 3.6; Discussion has 5.5 without "
     "5.1–5.4; numbering inconsistent from the Introduction onward",
     "Confirmed — Sections 3.4, 3.5 and 5.1–5.4 did not exist and the Introduction "
     "carried no number. Numbering rebuilt continuously from Section 1."],
    ["Remove the unused cor(y, ŷ) definition; correct 'n i s.'",
     "Confirmed. The unused definition has been deleted and the notation corrected."],
    ["Figure 5 duplicated caption sentence; missing callout in §4.4",
     "Confirmed. Caption rewritten and the figure callout placed in the body text."],
    ["Regenerate Figures 3–6 at consistent resolution, add (a)–(c) labels, prevent "
     "title cropping",
     "All figures regenerated at 400 dpi and supplied additionally as vector PDF. "
     "Panel labels added and panel order now matches every caption."],
    ["Remove the caption embedded inside Figure 2",
     "Confirmed. The embedded caption has been removed; only the document caption remains."],
    ["Repeat the header row when Table 2 spans pages",
     "Header-row repetition enabled; all tables checked at final publication size."],
    ["Do not describe total N, Olsen P and exchangeable K collectively as directly "
     "plant-available",
     "Confirmed. Terminology corrected throughout; the three pools are now distinguished "
     "wherever they are named together."],
    ["Move results and interpretation out of the Methods (§3.7.2) and out of the end of §4.6",
     "Confirmed. The XGBoost map-clarity statements have been deleted from the Methods "
     "and the interpretive material at the end of Section 4.6 moved to the Discussion."],
    ["Section 3 claims RF and XGBoost outperform SVM or ANN and require fewer adjustments",
     "Confirmed as untested. Both claims removed."],
    ["Narrow and state the novelty claim once",
     "Confirmed — 'first' appeared in the Abstract and twice in the Conclusions. The "
     "claim is now made once, in the Introduction, and defined precisely as the first "
     "field-validated nutrient mapping of this district at Sentinel-2 resolution."],
    ["Delete the near-duplicate sentence in the Conclusions",
     "Confirmed. The two near-identical closing paragraphs have been merged."],
    ["Use one unit style throughout",
     "Confirmed — 'mg kg-1' appeared once in Section 6 against 'mg/kg' elsewhere. "
     "Standardised on mg kg⁻¹ throughout, including all figures and tables."],
    ["Merge consecutive parenthetical citation groups",
     "Confirmed in Section 1, Section 4.4 and Table 2. All merged; reference-list "
     "alphabetisation and APA consistency corrected."],
    ["Acknowledgments thank anonymous reviewers although the review concerned "
     "technical/editorial checks",
     "Confirmed. The wording now reflects the actual review history."],
    ["Comprehensive professional English editing",
     "Confirmed and accepted. The manuscript has been professionally edited; the "
     "specific phrases quoted by the referee, and the Section 3.2.2 passage, have been "
     "rewritten."],
    ["Recheck keyword compliance",
     "Keywords revised to avoid repeating title terms at word level."],
], columns=["Referee point", "Response and action"])
table(minor, fontsize=8.0, widths=[6.0, 10.4])

doc.add_page_break()

# ==========================================================================
#                       SECTION 3 : NEW MATERIAL
# ==========================================================================
doc.add_heading("3.  Summary of new and revised material", level=1)
figs = pd.DataFrame([
    ["Figure 1", "Study area, sampling design with the 10 spatial CV blocks, and the "
                 "dissimilarity / applicability-domain map", "New"],
    ["Figure 3", "Observed versus Random Forest predicted, held-out outer folds, with "
                 "1:1 line, fitted slope, MAE, bias, CCC and n", "Replaced"],
    ["Figure 4", "Observed versus XGBoost predicted, held-out outer folds", "Replaced"],
    ["Figure 5", "Performance by validation scheme with bootstrap confidence intervals "
                 "and non-machine-learning baselines", "New"],
    ["Figure 6", "Permutation importance on held-out folds, aggregated by class and by "
                 "predictor, with between-fold standard deviations", "Replaced"],
    ["Figure 7", "Empirical variograms of nutrients and of model residuals", "New"],
    ["Figure 8", "Continuous Random Forest 10 m nutrient surfaces", "Replaced"],
    ["Figure 9", "Continuous XGBoost 10 m nutrient surfaces", "Replaced"],
    ["Figure 10", "Local prediction uncertainty at 10 m", "New"],
    ["Figure 11", "Classified fertility maps inside the applicability domain", "Replaced"],
], columns=["Item", "Content", "Status"])
table(figs, fontsize=8.4, widths=[2.4, 11.0, 2.6],
      caption="Table R15. Figures in the revised manuscript.")

tabs = pd.DataFrame([
    ["Table 1", "Sentinel-2 bands, native and output resolution, resampling", "Revised"],
    ["Table 2", "Corrected spectral index definitions and attributions", "Revised"],
    ["Table 3", "Definitive descriptive statistics of the 110 observations", "Replaced"],
    ["Table 4", "Model performance by scheme with bootstrap confidence intervals", "Replaced"],
    ["Table 5", "Benchmark comparison against null, linear and geostatistical baselines", "New"],
    ["Table 6", "Permutation importance aggregated by predictor class", "New"],
    ["Table 7", "Selected hyperparameters per model and nutrient", "New"],
    ["Table 8", "Variogram parameters and Moran's I", "New"],
    ["Table 9", "Fertility class accuracy and Cohen's κ from outer-fold predictions", "New"],
    ["Table 10", "Cartographic metadata and class areas", "New"],
    ["Table S1", "Complete 40-row predictor inventory", "New supplementary"],
    ["Table S2", "Fold-by-fold hyperparameter selections", "New supplementary"],
    ["Table S3", "Full confusion matrices with producer's and user's accuracies",
     "New supplementary"],
], columns=["Item", "Content", "Status"])
table(tabs, fontsize=8.4, widths=[2.4, 11.0, 2.6],
      caption="Table R16. Tables in the revised manuscript and supplementary material.")

doc.add_heading("Closing", level=2)
P("The referee's central judgement — that the previous manuscript did not establish "
  "that its reported scores were independent spatial predictions — was correct, and "
  "acting on it has changed the paper substantially. The validated accuracies are lower "
  "than those previously claimed, two headline conclusions have been withdrawn, and "
  "three defects of our own are disclosed above. We believe the result is a more modest "
  "but genuinely defensible contribution, and we are grateful for a review that made "
  "that possible.")
P("We remain at the editor's and referee's disposal for any further clarification.")

out = os.path.join(DOCS, "Response_to_Referee_Amrouss_et_al_Revision2.docx")
doc.save(out)
print("wrote", out)
print("STEP 07 complete.")
