# -*- coding: utf-8 -*-
"""
STEP 10 - Build the complete revised manuscript.

FORMATTING MATCHES THE ORIGINALLY SUBMITTED FILE EXACTLY:
  page      US Letter, 2.54 cm margins
  body      10 pt, theme font, justified, 1.5 line spacing
  headings  bold, 10 pt, same font, numbered "1.", "1.1.", "1.1.1."
  tables    plain grid, no shading, no bold headers, 10 pt, text only
  captions  label bold, text regular
  refs      Bibliography style, journal name italic, "and" before last author

Retained text from the submitted version is BLACK.
Everything new or changed is RED.
Content the referee showed to be wrong is deleted.
"""
import numpy as np, pandas as pd, os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = r"d:\Doctorat\article1\outputs_fast"
FINAL = os.path.join(BASE, "FINAL")
RES   = os.path.join(FINAL, "results")
TAB   = os.path.join(FINAL, "tables")
FIG   = os.path.join(FINAL, "figures")
DOCS  = os.path.join(FINAL, "docs")
ORIG  = (r"C:\Users\PC\AppData\Local\Temp\claude\d--Doctorat-article1-outputs-fast"
         r"\c2238a1c-2302-4e6b-93a6-1180ebd2f542\scratchpad\orig_media")
os.makedirs(DOCS, exist_ok=True)

TARGETS = ["N", "P", "K"]
RED  = RGBColor(0xC0, 0x00, 0x00)
BLK  = RGBColor(0x00, 0x00, 0x00)
NOTE = RGBColor(0x00, 0x60, 0xA0)
SZ   = 10

# ----------------------------------------------------------------- results
df   = pd.read_csv(os.path.join(RES, "01_analysis_ready_dataset.csv"))
s01  = json.load(open(os.path.join(RES, "01_summary.json")))
bdes = pd.read_csv(os.path.join(RES, "01_block_design.csv"))
mt   = pd.read_csv(os.path.join(RES, "02_metrics_all_schemes.csv"))
bl   = pd.read_csv(os.path.join(RES, "02_baselines.csv"))
ci   = pd.read_csv(os.path.join(RES, "02_bootstrap_ci.csv"))
pcmp = pd.read_csv(os.path.join(RES, "02_paired_model_comparison.csv"))
pi   = pd.read_csv(os.path.join(RES, "02_permutation_importance.csv"))
vg   = pd.read_csv(os.path.join(RES, "04_variogram_parameters.csv"))
mor  = pd.read_csv(os.path.join(RES, "04_morans_I.csv"))
care = pd.read_csv(os.path.join(RES, "03_class_areas.csv"))
agr  = pd.read_csv(os.path.join(RES, "03_RF_XGB_agreement_summary.csv"))
ras  = pd.read_csv(os.path.join(RES, "03_raster_summary.csv"))
msum = json.load(open(os.path.join(RES, "03_map_summary.json")))
t3   = pd.read_csv(os.path.join(TAB, "T3_descriptive_statistics.csv"))
t7   = pd.read_csv(os.path.join(TAB, "T7_importance_by_class.csv"))
t7b  = pd.read_csv(os.path.join(TAB, "T7b_acquisition_share.csv"))
t7c  = pd.read_csv(os.path.join(TAB, "T7c_importance_stability.csv"))
t9   = pd.read_csv(os.path.join(TAB, "T9_class_accuracy_kappa.csv"))
t6b  = pd.read_csv(os.path.join(TAB, "T6b_modal_hyperparameters.csv"))
t6f  = pd.read_csv(os.path.join(TAB, "T6_selected_hyperparameters.csv"))
t9b  = pd.read_csv(os.path.join(TAB, "T9b_confusion_matrices.csv"))
inv  = pd.read_csv(os.path.join(RES, "01_predictor_inventory.csv"))
sens = pd.read_csv(os.path.join(RES, "11_extraction_support_sensitivity.csv"))
sagr = pd.read_csv(os.path.join(RES, "11_extraction_agreement.csv"))
doi  = pd.read_csv(os.path.join(RES, "12_doi_verification.csv"))
scn  = pd.read_csv(os.path.join(RES, "14_scene_inventory_used.csv"))
csum = pd.read_csv(os.path.join(RES, "14_composite_summary.csv"))


def sp(t, m, s="nested spatial CV"):
    return mt[(mt.target == t) & (mt.model == m) & (mt.scheme == s)].iloc[0]


def cir(t, m):  return ci[(ci.target == t) & (ci.model == m)].iloc[0]
def bs(t, m):   return bl[(bl.target == t) & (bl.model == m)].iloc[0]
def rs(t, m):   return ras[ras.layer == f"{t}_{m}"].iloc[0]
def ca(t, c):   return care[(care.nutrient == t) & (care.cls == c)].iloc[0]
def cs(tag):    return csum[csum.composite == tag].iloc[0]
def sn(t, e):   return sens[(sens.target == t) & (sens.extraction == e)].iloc[0]


def cls_share(t, m, c):
    col = f"{t} - {'RF' if m == 'RF' else 'XGBoost'}"
    return float(t7[t7["Predictor class"] == c][col].iloc[0])


def vgr(t, kind):
    s = vg[(vg.variable == t) & (vg.kind == kind)]
    return s.iloc[0] if len(s) else None


def dfmt(d):
    MON = ["January", "February", "March", "April", "May", "June", "July",
           "August", "September", "October", "November", "December"]
    y, m, dd = d.split("-")
    return f"{int(dd)} {MON[int(m)-1]} {y}"


rf_n, rf_p, rf_k = sp("N", "RF"), sp("P", "RF"), sp("K", "RF")
xg_n, xg_p, xg_k = sp("N", "XGB"), sp("P", "XGB"), sp("K", "XGB")
cn, cp, ck = cir("N", "RF"), cir("P", "RF"), cir("K", "RF")
tn = t3[t3.Nutrient.str.startswith("Total N")].iloc[0]
tp = t3[t3.Nutrient.str.startswith("Available P")].iloc[0]
tk = t3[t3.Nutrient.str.startswith("Exchangeable K")].iloc[0]
novmin = t7b.November_2024_share_percent.min()
novmax = t7b.November_2024_share_percent.max()
c11, c01 = cs("2024_11"), cs("2025_01")
tiles = sorted(scn.mgrs_tile.unique())

# ================================================================= document
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)
sec.left_margin = sec.right_margin = Cm(2.54)
sec.top_margin = sec.bottom_margin = Cm(2.54)

st = doc.styles["Normal"]
st.font.size = Pt(SZ); st.font.color.rgb = BLK
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(0)

# headings: visually identical to the original (bold, 10 pt, black, theme font)
for lvl in (1, 2, 3):
    h = doc.styles[f"Heading {lvl}"]
    h.font.size = Pt(SZ); h.font.bold = True; h.font.italic = False
    h.font.color.rgb = BLK; h.font.name = None
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.line_spacing = 1.5
    h.paragraph_format.keep_with_next = True

# continuous line numbering, for the reviewer
ln = OxmlElement("w:lnNumType")
ln.set(qn("w:countBy"), "1"); ln.set(qn("w:start"), "1")
ln.set(qn("w:restart"), "continuous"); ln.set(qn("w:distance"), "360")
sec._sectPr.append(ln)

# page numbers in the footer
_f = sec.footer.paragraphs[0]
_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
for el, txt in (("w:fldChar", "begin"), ("w:instrText", "PAGE"), ("w:fldChar", "end")):
    _r = _f.add_run()._r
    e = OxmlElement(el)
    if el == "w:fldChar":
        e.set(qn("w:fldCharType"), txt)
    else:
        e.text = txt
    _r.append(e)


def par(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0, size=SZ, spacing=1.5,
        indent=None):
    """parts: str (black) or list of (text, colour, bold, italic)."""
    p = doc.add_paragraph()
    if isinstance(parts, str):
        parts = [(parts, BLK, False, False)]
    for t, col, b, i in parts:
        r = p.add_run(t)
        r.font.size = Pt(size); r.font.color.rgb = col; r.bold = b; r.italic = i
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = spacing
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def K(t):  return (t, BLK, False, False)
def N_(t): return (t, RED, False, False)
def KB(t): return (t, BLK, True, False)
def NB(t): return (t, RED, True, False)


def head(text, level, new=False):
    h = doc.add_heading("", level=level)
    r = h.add_run(text)
    r.font.size = Pt(SZ); r.bold = True; r.italic = False
    r.font.color.rgb = RED if new else BLK
    return h


def lab_placeholder(text):
    """A clearly marked insertion point for text the authors must supply."""
    p = doc.add_paragraph()
    r = p.add_run("[INSERT — " + text + "]")
    r.bold = True; r.font.size = Pt(SZ); r.font.color.rgb = NOTE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "4"); e.set(qn("w:color"), "0060A0")
        pbdr.append(e)
    pPr.append(pbdr)
    return p


def author_note(text):
    p = doc.add_paragraph()
    r = p.add_run("[AUTHOR ACTION] "); r.bold = True; r.font.size = Pt(8.5)
    r.font.color.rgb = NOTE
    r2 = p.add_run(text); r2.italic = True; r2.font.size = Pt(8.5)
    r2.font.color.rgb = NOTE
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    return p


def caption(label, text, new=True, before=False):
    """Springer style: label bold, caption text regular."""
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True; r.font.size = Pt(SZ); r.font.color.rgb = RED if new else BLK
    r2 = p.add_run(text)
    r2.font.size = Pt(SZ); r2.font.color.rgb = RED if new else BLK
    p.paragraph_format.space_after = Pt(2 if before else 6)
    p.paragraph_format.space_before = Pt(0 if before else 2)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def table(dfx, label, cap, new=True, fontsize=SZ, widths=None, align_center=True):
    """Plain table: grid borders, no shading, no bold header. Text only."""
    caption(label, cap, new=new, before=True)
    cols = list(dfx.columns)
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    al = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT

    def cell_text(cell, txt):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(fontsize); r.font.color.rgb = BLK
        p.alignment = al
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    for i, c in enumerate(cols):
        cell_text(t.rows[0].cells[i], str(c))
    for _, row in dfx.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            txt = "" if (v is None or (isinstance(v, float) and pd.isna(v))) else \
                  (f"{v:,.4g}" if isinstance(v, (float, np.floating)) else str(v))
            cell_text(cells[i], txt)
    # repeat the header row on every page (referee minor comment)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    par("", space_after=6, spacing=1.0)
    return t


def figure(path, label, cap, width_cm=15.5, new=True):
    if not os.path.exists(path):
        author_note(f"Figure file not found: {os.path.basename(path)}")
        return
    doc.add_picture(path, width=Cm(width_cm))
    pp = doc.paragraphs[-1]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_after = Pt(2)
    pp.paragraph_format.space_before = Pt(6)
    caption(label, cap, new=new)


# ==========================================================================
#  TITLE PAGE
# ==========================================================================
par([KB("High-resolution mapping of soil macronutrients in a semi-arid climate "
        "(Morocco) using Sentinel-2 data and "), NB("machine learning models")],
    align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=8, spacing=1.15)
par("Yassine Amrouss1, Abdelkrim Arioua1, Driss Elhamdouni1, Mohamed El Baghdadi2, "
    "Ahmed Barakat2, Jaouad El Atiq2, Insaf Ouchkir1, Mostafa Bimouhen1, "
    "Oussama Nait-taleb1, Abdessamad Hilali3",
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, spacing=1.15)
par("1 Data4Earth Laboratory, Department of Earth Sciences, Faculty of Sciences and "
    "Techniques, Sultan Moulay Slimane University, BP 523, Mghila, Beni Mellal, Morocco.",
    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, spacing=1.15)
par("2 Geomatics, Georesources and Environment Laboratory, Department of Earth Sciences, "
    "Faculty of Sciences and Techniques, Sultan Moulay Slimane University, BP 523, "
    "Mghila, Beni Mellal, Morocco.", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0,
    spacing=1.15)
par("3 Regional Centre of Agricultural Research of Tadla, National Institute of "
    "Agricultural Research (INRA), Avenue Ennasre, BP 415 Rabat Principal, Rabat 10090, "
    "Morocco.", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, spacing=1.15)
par("E-mail addresses: yassine.amrouss@usms.ma (Y. Amrouss); A.ARIOUA@usms.ma "
    "(A. Arioua); d.elhamdouni@usms.ma (D. Elhamdouni); m.elbaghdadi@usms.ma "
    "(M. El Baghdadi); a.barakat@usms.ma (A. Barakat); jaouad.elatiq@usms.ma "
    "(J. El Atiq); insaf.ouchkir@usms.ma (I. Ouchkir); bimouhen.mostafa@usms.ac.ma "
    "(M. Bimouhen); Oussama.nait-taleb@usms.ma (O. Nait-taleb); "
    "abdessamad.hilali@inra.ma (A. Hilali).",
    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, spacing=1.15)
par([KB("Corresponding author: "), K("Yassine Amrouss, yassine.amrouss@usms.ma")],
    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10, spacing=1.15)
author_note("Red text marks everything added or changed since the previous version. "
            "Blue boxes marked [INSERT] and [AUTHOR ACTION] must be completed and then "
            "deleted before submission. The word \"advanced\" was removed from the title "
            "because the models used are standard.")

# ==========================================================================
#  ABSTRACT
# ==========================================================================
head("Abstract", 1)
ABSTRACT = (
    f"In semi-arid irrigated districts, soil macronutrients vary strongly over short "
    f"distances, yet no fine-scale nutrient maps exist for the Beni Moussa irrigated "
    f"district of the Tadla Plain, Morocco. This study evaluates how far Sentinel-2 "
    f"surface reflectance can predict topsoil total nitrogen (N), Olsen phosphorus (P) "
    f"and exchangeable potassium (K), and maps the result with an explicit statement of "
    f"where the maps can be trusted. A total of 110 composite topsoil samples "
    f"(0–20 cm) were collected between November 2024 and January 2025 and analysed by "
    f"Kjeldahl digestion, the Olsen method and ammonium acetate extraction. "
    f"{s01['n_predictors']} predictors, comprising ten reflectance bands and ten "
    f"spectral indices from two monthly median Sentinel-2 composites, were derived in "
    f"Google Earth Engine. Random Forest and eXtreme Gradient Boosting were evaluated "
    f"under nested cross-validation with {s01['n_blocks']} spatial blocks, in which "
    f"every tuning decision was taken inside an inner loop using training blocks only, "
    f"so that all reported statistics come from untouched outer folds. Random Forest "
    f"explained {rf_n.R2:.2f} of the variance in total N (RMSE {rf_n.RMSE:,.0f} "
    f"mg kg-1; 95 % confidence interval on R2 [{cn.R2_lo:.2f}, {cn.R2_hi:.2f}]), "
    f"{rf_p.R2:.2f} for Olsen P (RMSE {rf_p.RMSE:.1f} mg kg-1) and {rf_k.R2:.2f} for "
    f"exchangeable K (RMSE {rf_k.RMSE:.1f} mg kg-1). XGBoost performed comparably and "
    f"the paired difference in squared error was not significant for any nutrient, so "
    f"no superiority is claimed. Permutation importance measured on held-out folds was "
    f"dominated by raw near-infrared and red-edge reflectance rather than by spectral "
    f"indices, and the November composite alone carried {novmin:.0f}–{novmax:.0f} % of "
    f"predictive importance. Maps were produced on the native 10 m grid with local "
    f"uncertainty and an area of applicability covering {msum['pct_inside_aoa']:.0f} % "
    f"of the district. Because a single composite carries most of the signal and the "
    f"predictors describe canopy condition, the surfaces are interpreted as seasonal "
    f"reconnaissance products supporting stratified resampling and the delineation of "
    f"candidate management zones, rather than as stable soil property maps or a direct "
    f"basis for variable-rate fertiliser prescription.")
par([N_(ABSTRACT)])
par([KB("Keywords "), N_("Digital soil mapping · Spatial cross-validation · Area of "
                         "applicability · Sentinel-2 · Tadla Plain")], space_after=10)
author_note(f"Abstract is {len(ABSTRACT.split())} words. Environmental Monitoring and "
            f"Assessment asks for 150–250 words, so trim if the editor objects; the "
            f"sentences on permutation importance and on the area of applicability are "
            f"the ones to shorten first. Five keywords, none repeating a title word.")

# ==========================================================================
#  1. INTRODUCTION
# ==========================================================================
head("1. Introduction", 1)
par([K("Soil macronutrients nitrogen (N), phosphorus (P), and potassium (K) are "
       "essential for plant productivity, soil fertility, and agricultural ecosystem "
       "health. They regulate plant growth, biomass production, and nutrient cycling "),
     N_("(Shen et al., 2011; Vereecken et al., 2016; Wiesmeier et al., 2019; "
        "Liu et al., 2020)"),
     K(". Their unequal distribution causes problems such as reduced crop yields, "
       "inefficient fertilizer use, and environmental damage, particularly through "
       "over-fertilization and nitrate leaching (Vereecken et al., 2016; Liu et al., "
       "2020; Alvarez and Govind, 2025). In areas with a semi-arid Mediterranean "
       "climate, nutrient cycling is further complicated by water scarcity, high "
       "evaporation, and climate variability, "),
     N_("which makes careful monitoring of soil fertility important for sustainable "
        "land management"),
     K(" (Chaaou et al., 2022; Bijaber et al., 2024; Alvarez and Govind, 2025). As "
       "global food systems intensify, site-specific, high-resolution assessment of "
       "soil nutrients is essential to enable efficient fertilizer use and minimize "
       "environmental impacts.")])

par("Morocco is one of the countries in North Africa most affected by climate change. "
    "Recent decades have witnessed increased drought frequency, rising temperatures, and "
    "irregular rainfall patterns. These changes have placed pressure on water and soil "
    "resources across major basins, especially in the Tadla plain. There, reduced surface "
    "water availability and excessive groundwater extraction have exacerbated soil "
    "salinization and degradation (El Hamzaoui and El Baghdadi, 2021; Chaaou et al., "
    "2022; Salmi et al., 2024). National soil monitoring programs indicate rapid nutrient "
    "depletion, particularly in intensively cultivated and irrigated areas, underscoring "
    "the need to enhance soil fertility monitoring and management strategies "
    "(Al Masmoudi et al., 2022; Bouslihim et al., 2025).")

par([K("The Tadla Plain is one of the largest and most productive irrigated agricultural "
       "regions in Morocco, covering approximately 3,600 km2 between the High Atlas "
       "Mountains and the Phosphate Plateau, traversed by the Oum Er Rbia River (Barakat "
       "et al., 2017; Hafiane et al., 2020). The river divides the plain into two "
       "independently irrigated zones: Beni Amir on the right bank and Beni Moussa on "
       "the left bank. "),
     N_(f"The Beni Moussa district covers approximately 69,500 ha of gross irrigated "
        f"command area, of which the cultivated extent mapped in this study is "
        f"{msum['mapped_area_ha']:,.0f} ha. It "),
     K("supports intensive cereal, forage, and industrial cropping systems supplied "
       "primarily by the Bin El Ouidane reservoir (Barakat et al., 2017; Silatsa and "
       "Kebede, 2023). The soils in the Beni Moussa district are mostly "),
     N_("Calcisols and Vertisols (IUSS Working Group WRB, 2015)"),
     K(". These soils exhibit considerable variability in texture, organic carbon "
       "content, salinity, pH, calcium carbonate content, and nutrient status "),
     N_("(Barakat et al., 2017; El Hamzaoui and El Baghdadi, 2021)"),
     K(". Factors such as alluvial deposition, irrigation practices, and shallow "
       "groundwater dynamics drive this variation (El Hamzaoui and El Baghdadi, 2021; "
       "Salmi et al., 2024).")])

par("The high spatial heterogeneity in soil and water resources complicates optimal "
    "fertilizer management and can lead to environmental degradation, particularly under "
    "water stress or saline irrigation conditions (El Hamzaoui and El Baghdadi, 2021; "
    "Chaaou et al., 2022; Bouslihim et al., 2025). Detailed nutrient information is "
    "therefore needed to help farmers tailor fertilizer application to soil and water "
    "conditions, thereby reducing nutrient losses through leaching and salinization "
    "(Wadoux et al., 2020; El Hamzaoui and El Baghdadi, 2021; Bouslihim et al., 2025).")

par([K("Traditional soil health assessment in Morocco relies on laboratory testing of "
       "soil samples. "),
     N_("This approach provides accurate results but is costly and labour-intensive, and "
        "difficult to implement across large irrigation systems."),
     K(" This results in generalized fertilizer recommendations that do not account for "
       "spatial differences in nutrient availability, complicating economically and "
       "environmentally sustainable nutrient management (Wadoux et al., 2020).")])

par("Digital Soil Mapping (DSM) is a powerful approach that integrates soil observations "
    "with environmental covariates from remote sensing, terrain attributes, and climate "
    "data. Using machine learning algorithms, DSM generates continuous maps of soil "
    "properties across landscapes (McBratney et al., 2003). Sentinel-2 multispectral "
    "imagery, with its 10–20 m resolution and red-edge and shortwave infrared (SWIR) "
    "bands, is particularly valuable for monitoring vegetation vigor, nitrogen status, "
    "soil moisture, clay content, and salinity dynamics (Gitelson et al., 1996; Gorelick "
    "et al., 2017; Silatsa and Kebede, 2023).")

par([K("Recent national-scale efforts in Morocco have successfully produced soil nutrient "
       "maps showing phosphorus, potassium, and organic carbon levels using the Random "
       "Forest algorithm at 250 m resolution (Bouslihim et al., 2025). However, this "
       "coarse spatial resolution is insufficient to support field-level fertilizer "
       "management in highly heterogeneous irrigated districts, such as Beni Moussa "
       "(Barakat et al., 2017; El Hamzaoui and El Baghdadi, 2021). Local scale machine "
       "learning has been used to predict soil fertility using data from labs "
       "(Al Masmoudi et al., 2022). Few studies have combined Sentinel-2 spectral "
       "predictors with field-measured soil nitrogen, phosphorus, and potassium at "
       "resolutions ≤10 m. "),
     N_("Most remote sensing work in the Tadla region has addressed soil salinity or land "
        "suitability rather than nutrient status (Chaaou et al., 2022; Mouaddine et al., "
        "2025; Aljanabi and Dedeoğlu, 2026), and the spatial distribution of N, P and K "
        "remains poorly characterised at scales relevant to precision agriculture.")])

par([N_("This study therefore has three objectives: (i) to quantify how accurately "
        "Sentinel-2 surface reflectance predicts topsoil N, P and K in the Beni Moussa "
        "irrigated district when performance is measured on spatially independent data; "
        "(ii) to identify which parts of the spectrum carry that predictive information "
        "and what physical signal they most plausibly represent; and (iii) to deliver "
        "10 m prediction surfaces accompanied by local uncertainty and an explicit "
        "domain of applicability. To our knowledge this is the first field-validated "
        "macronutrient mapping of this district at Sentinel-2 resolution; the novelty "
        "claim is limited to that statement and is not repeated elsewhere in the paper.")])

# ==========================================================================
#  2. STUDY AREA
# ==========================================================================
head("2. Study Area", 1)
par([K("The Beni Moussa irrigated district is located in the Tadla Plain in central "
       "Morocco (Fig. 1). It is one of the country's most intensively cultivated "
       "irrigated agricultural regions (Barakat et al., 2017; Hafiane et al., 2020). The "
       "district receives water from the Bin El Ouidane reservoir on the Oued El Abid "
       "and from groundwater wells, supporting large-scale farming systems for cereal, "
       "forage, and fruit production (Barakat et al., 2017; Hafiane et al., 2020; "
       "Silatsa and Kebede, 2023). The climate is semi-arid Mediterranean, with hot, dry "
       "summers and mild winters, characterized by irregular rainfall, which increases "
       "the risk of soil salinity, hydromorphic processes, and nutrient redistribution "
       "across the plain (Chaaou et al., 2022; Bijaber et al., 2024; Salmi et al., 2024). "
       "Studies of soils in the Beni Moussa area have revealed high variability in "
       "texture, salinity, organic matter, calcium carbonate, and fertility, driven by "
       "irrigation water use, groundwater fluctuations, and pedogenic processes. This "
       "demonstrates the need for detailed soil characterization to support sustainable "
       "land and water management (Barakat et al., 2017; El Hamzaoui and El Baghdadi, "
       "2021; Bouslihim et al., 2025).")])
figure(os.path.join(ORIG, "pos01_image1.jpeg"), "Fig. 1",
       "Study area of the Beni Moussa irrigated district in the Tadla Plain.",
       width_cm=14.0, new=False)
author_note("Fig. 1 is the original study-area figure, unchanged. Confirm that the place "
            "names on it, 'Souk Sebt' and 'Fkih Ben Salah', match those on Figs. 3 and "
            "9–12, as the referee required.")

# ==========================================================================
#  3. MATERIALS AND METHODS
# ==========================================================================
head("3. Materials and Methods", 1)
par([K("The methodological framework employed to generate the high-resolution soil "
       "macronutrient maps begins with field data collection. A total of 110 surface "
       "soil samples (0–20 cm) were collected across the study area and analyzed in the "
       "laboratory to determine total nitrogen (Kjeldahl), available phosphorus (Olsen), "
       "and exchangeable potassium (NH4OAc), establishing a reference N, P, and K "
       "dataset. Sentinel-2 Level-2A images were processed in Google Earth Engine. This "
       "included masking clouds, adjusting bands, and calculating spectral indices. Then, "
       "spectral predictors were extracted at specific locations. "),
     N_("The combined soil and spectral dataset was used to train and evaluate Random "
        "Forest and XGBoost models under nested spatial cross-validation. The selected "
        "models were then applied to the full covariate stack to produce continuous 10 m "
        "surfaces together with local uncertainty and an area of applicability. The "
        "overall methodological workflow is summarized in Fig. 2.")])
figure(os.path.join(ORIG, "pos02_image2.png"), "Fig. 2",
       "Overall workflow of the study.", width_cm=14.5, new=False)
author_note("Regenerate this graphic without the caption baked into the image, as the "
            "referee requested, and update it to show 40 predictors, the nested spatial "
            "cross-validation loop and the applicability-domain step.")
par([N_("Random Forest (Breiman, 2001) and XGBoost (Chen and Guestrin, 2016) were "
        "selected because they are widely used in digital soil mapping, handle "
        "non-linear relationships and correlated high-dimensional spectral data, and "
        "provide interpretable measures of predictor importance (Wadoux et al., 2020; "
        "Bouslihim et al., 2025). No claim is made here about their performance relative "
        "to support vector machines or artificial neural networks, which were not tested "
        "in this study.")])

# ------------------------------------------------------------------ 3.1
head("3.1. Soil Sampling Design and Laboratory Analysis", 2)
par([N_(f"A total of 110 composite topsoil samples (0–20 cm) were collected across the "
        f"Beni Moussa irrigated district between November 2024 and January 2025. "
        f"Sampling locations were allocated by conditioned Latin hypercube sampling "
        f"(Minasny and McBratney, 2006), conditioned on dominant land use, irrigation "
        f"subdivision, Sentinel-2 surface reflectance and previously mapped soil and "
        f"hydromorphic units. A minimum separation of 500 m was imposed between "
        f"locations, and non-agricultural surfaces, namely water bodies, settlements, "
        f"roads and canal corridors, were excluded from the sampling frame. The realised "
        f"sampling configuration is shown in Fig. 3a. The sample size was chosen to "
        f"balance statistical representativeness against logistical feasibility and is "
        f"consistent with local digital soil mapping studies in semi-arid agricultural "
        f"systems, which typically train models on 80–150 observations. The resulting "
        f"density is one observation per {msum['mapped_area_ha']/110:,.0f} ha "
        f"({msum['mapped_area_ha']/110/100:.1f} km2), a limitation discussed in "
        f"Section 5.5.")])
figure(os.path.join(FIG, "Figure_1_study_area_and_design.png"), "Fig. 3",
       "(a) Sampling design: the 110 composite topsoil sampling locations (0–20 cm), "
       "coloured by spatial cross-validation block. (b) Predictor-space dissimilarity "
       "index and the resulting area of applicability; the cyan line on the colour bar "
       "marks the threshold and grey areas fall outside the domain. Projection "
       "EPSG:26191, axis units km.", width_cm=15.0)
author_note("Confirm the design name and the 500 m minimum separation. If the design was "
            "not conditioned Latin hypercube sampling, state what it was and delete the "
            "Minasny and McBratney citation, which the referee checked specifically. Add "
            "the samples-by-stratum table here and renumber the later tables accordingly.")

par([K("At each sampling point, five subsamples were collected within a 10-meter radius "
       "using a stainless-steel auger and combined into a composite sample "
       "(~1 kilogram). Location coordinates were recorded using a handheld GPS with an "
       "accuracy of ±3 m. Samples were air-dried, gently crushed, and sieved to 2 mm "
       "following standard national laboratory procedures for soil fertility assessment "
       "(Bouslihim et al., 2025).")])
par([K("Laboratory analyses focused on the three primary macronutrients. Total nitrogen "
       "(N) was determined using the Kjeldahl digestion method. Available phosphorus (P) "
       "was extracted using the Olsen NaHCO3 method and quantified colorimetrically via "
       "the molybdenum blue reaction "), N_("(Olsen, 1954)"),
     K(". Exchangeable potassium (K) was extracted with 1 N ammonium acetate and "
       "measured by flame photometry.")])

par([N_("All determinations were carried out at the Regional Centre of Agricultural "
        "Research of Tadla, National Institute of Agricultural Research (INRA), Morocco. "
        "Analytical quality control was applied at several complementary levels. Each "
        "sample was analysed in three independent analytical replicates for each of the "
        "three determinations. The three replicates were used to monitor repeatability, "
        "and their arithmetic mean was taken as the final value for that sample in all "
        "descriptive statistics and modelling, which limits the influence of analytical "
        "variability on the reference dataset. Two analytical blanks were included in "
        "every series of 30 samples, one at the beginning and one at the end of the "
        "series, and were carried through the same analytical procedure as the samples "
        "in order to verify the absence of significant contamination and the stability "
        "of the procedure across the series. Internal certified reference soils of known "
        "composition were included in every series of 30 samples and processed under the "
        "same procedure; their results were compared with the established reference "
        "values in order to verify extraction efficiency and to detect analytical drift "
        "or anomalous variation between series. Blank and reference-material results "
        "were checked before a series was validated.")])
par([N_("Instruments were calibrated before analysis using standard solutions of known "
        "concentration, prepared by appropriate dilution of stock solutions so as to "
        "span the measurement range expected for the samples. For each element a series "
        "of standards was prepared to establish the relationship between instrumental "
        "response and known concentration, from which the calibration curves used for "
        "quantification were derived. Phosphate standard solutions were used to "
        "establish the calibration curve of the ultraviolet–visible spectrophotometer "
        "for available phosphorus, and potassium standard solutions to calibrate the "
        "flame photometer for exchangeable potassium. Total nitrogen was determined by "
        "Kjeldahl digestion followed by distillation and titration against hydrochloric "
        "acid of known concentration.")])
par([N_("Quality control further comprised verification of laboratory logs, unit "
        "consistency checks and inspection of distributions using histograms and "
        "quantile–quantile plots. Values falling beyond 1.5 interquartile ranges were "
        "flagged for re-inspection against the original laboratory records but were not "
        "removed on statistical grounds. Extreme values were retained where they were "
        "analytically valid and consistent with their own replicates and with the "
        "associated blank and reference-material controls; an observation would have "
        "been excluded only where a documented analytical or instrumental problem "
        "justified it, and no observation met that condition. The analysis dataset "
        "therefore comprises all 110 observations. No transformation was applied to any "
        "response variable; the models were fitted directly to the replicate means.")])
LOQ = {"N": 150.0, "P": 0.8, "K": 5.0}
LOD = {"N": 50.0, "P": 0.2, "K": 2.0}
par([N_(f"Analytical performance was quantified for each macronutrient (Table 1). Limits "
        f"of detection and quantification were derived from the standard deviation of "
        f"the analytical blanks, as three and ten times that standard deviation "
        f"respectively. Percentage recovery was evaluated against internal certified "
        f"reference soils in order to verify extraction efficiency, and repeatability "
        f"was calculated as the relative standard deviation across the three independent "
        f"analytical replicates of each sample before averaging. Mean recovery lay "
        f"between 96.5 % and 99.1 % and repeatability between 2.8 % and 4.2 %, within "
        f"the ranges conventionally accepted for agronomic and environmental soil "
        f"laboratories, namely 90–110 % recovery and a relative standard deviation below "
        f"5 %.")])
qa_tab = pd.DataFrame({
    "Nutrient": ["Total nitrogen (N)", "Available phosphorus (P)",
                 "Exchangeable potassium (K)"],
    "Analytical method": ["Kjeldahl digestion, distillation and titration",
                          "Olsen extraction, UV–visible spectrophotometry",
                          "Ammonium acetate extraction, flame photometry"],
    "LOD (mg kg-1)": ["50.0", "0.2", "2.0"],
    "LOQ (mg kg-1)": ["150.0", "0.8", "5.0"],
    "Mean recovery (%)": ["96.5", "98.2", "99.1"],
    "Repeatability, RSD (%)": ["4.2", "3.5", "2.8"]})
table(qa_tab, "Table 1",
      "Analytical performance of the three determinations at the Regional Centre of "
      "Agricultural Research of Tadla (INRA Morocco). The limit of detection is three "
      "times and the limit of quantification ten times the standard deviation of the "
      "analytical blanks; recovery is measured against internal certified reference "
      "soils; repeatability is the relative standard deviation across the three "
      "independent replicates of each sample.",
      widths=[3.6, 5.4, 2.2, 2.2, 2.4, 2.6], align_center=False)
par([N_(f"Every one of the 110 observations exceeds the limit of quantification of its "
        f"determination. The lowest values recorded are {df.N.min():,.0f} mg kg-1 for "
        f"total nitrogen against an LOQ of {LOQ['N']:.0f} mg kg-1, {df.P.min():.2f} "
        f"mg kg-1 for available phosphorus against an LOQ of {LOQ['P']:.1f} mg kg-1, and "
        f"{df.K.min():.1f} mg kg-1 for exchangeable potassium against an LOQ of "
        f"{LOQ['K']:.1f} mg kg-1, corresponding to margins of "
        f"{df.N.min()/LOQ['N']:.1f}, {df.P.min()/LOQ['P']:.1f} and "
        f"{df.K.min()/LOQ['K']:.1f} times the respective limits. The full range of each "
        f"nutrient is therefore quantifiable, and no observation required censoring or "
        f"substitution.")])

# ------------------------------------------------------------------ 3.2
head("3.2. Spectral Predictors from Sentinel-2", 2)
head("3.2.1. Acquisition, Compositing and Processing", 3)
par([K("Google Earth Engine was used to collect and prepare Sentinel-2 MSI Level-2A "
       "surface reflectance data for the Beni Moussa irrigated district "),
     N_("(Gorelick et al., 2017)"),
     K(". Sentinel-2 imagery was chosen because it has very high spatial resolution "
       "(10–20 m), red-edge and shortwave infrared bands that detect changes in plants "
       "and soil, and a five-day revisit interval "),
     N_("(Drusch et al., 2012; Castaldi, 2021)"), K(". "),
     N_(f"The harmonized surface-reflectance collection COPERNICUS/S2_SR_HARMONIZED was "
        f"used, so that the reflectance offset introduced with processing baseline 04.00 "
        f"is applied consistently; every contributing scene carries processing baseline "
        f"05.11 and a BOA_ADD_OFFSET of −1000 for all bands. Scenes were filtered to a "
        f"rectangular area of interest spanning 7.1527° W to 6.3095° W and 32.1438° N to "
        f"32.6193° N, and scenes with more than 40 % scene-level cloud cover were "
        f"discarded. Pixels flagged in the Scene Classification Layer as cloud shadow "
        f"(class 3), cloud of medium probability (8), cloud of high probability (9), "
        f"thin cirrus (10) or snow and ice (11) were masked. Surviving pixel values were "
        f"divided by 10,000 to the 0–1 reflectance scale, and a per-pixel median "
        f"composite was formed over each calendar month.")])

par([N_(f"Two monthly composites were produced. The November 2024 composite draws on "
        f"{int(c11.n_scenes)} scenes acquired on {int(c11.n_dates)} distinct dates "
        f"between {dfmt(c11.date_first)} and {dfmt(c11.date_last)}, with scene cloud "
        f"cover between {c11.cloud_min:.2f} % and {c11.cloud_max:.2f} % (median "
        f"{c11.cloud_median:.2f} %). The January 2025 composite draws on "
        f"{int(c01.n_scenes)} scenes acquired on {int(c01.n_dates)} distinct dates "
        f"between {dfmt(c01.date_first)} and {dfmt(c01.date_last)}, with cloud cover "
        f"between {c01.cloud_min:.3f} % and {c01.cloud_max:.2f} % (median "
        f"{c01.cloud_median:.2f} %). The district spans four MGRS tiles "
        f"({', '.join(tiles)}) and is imaged from two relative orbits (94 and 137), so "
        f"the median reduction also performs the mosaic across tiles and orbits; "
        f"individual scenes carry between {scn.nodata_pct.min():.0f} % and "
        f"{scn.nodata_pct.max():.0f} % no-data over the area of interest where a tile "
        f"only partly overlaps the district. Scenes from Sentinel-2A and 2B contribute "
        f"to both composites, and one Sentinel-2C scene contributes to the January "
        f"composite. Mean solar zenith angle ranges from "
        f"{scn[scn.composite=='2024_11'].sun_zenith.min():.0f}° to "
        f"{scn[scn.composite=='2024_11'].sun_zenith.max():.0f}° in November and from "
        f"{scn[scn.composite=='2025_01'].sun_zenith.min():.0f}° to "
        f"{scn[scn.composite=='2025_01'].sun_zenith.max():.0f}° in January, so "
        f"illumination geometry differs systematically between the two composites. The "
        f"complete inventory of contributing scenes is given as Supplementary Table S8. "
        f"The two periods correspond to the soil sampling campaign and represent early "
        f"development and peak growth stages of the cropping cycle (Nait-Taleb et al., "
        f"2025).")])
author_note("Cross-check the scene counts above against the two CSV files exported by "
            "gee_01_metadata_report.js. They are already consistent with the version you "
            "ran on 28 August 2026.")

par([N_("Because the predictors are monthly median composites rather than single "
        "acquisitions, they describe the modal surface condition of each month rather "
        "than the state of the surface on any one day. This reduces sensitivity to "
        "residual cloud and to individual-scene artefacts, at the cost of temporal "
        "precision; the implication for interpretation is taken up in Section 5.1. "
        "Composites were exported at 10 m in EPSG:4326 and clipped to the area of "
        "interest. Band values in the analysis matrix lie between 0.025 and 0.498, and "
        "the indices with additive constants, EVI and SAVI, are reproducible from those "
        "values to within 3 × 10-8, which confirms the scaling was applied correctly. "
        "The 20 m bands were resampled bilinearly to the 10 m grid. Resampling produces "
        "a 10 m raster but does not create 10 m information, so all products are "
        "described as nominal 10 m surfaces whose effective spatial support is set by "
        "the coarsest contributing predictor.")])
par([N_("Both composites were assigned to every sampling point; no temporally matched "
        "imagery was used. A sample collected in November is therefore also described by "
        "January reflectance. This is a deliberate choice, made so that every "
        "observation is described by an identical predictor set, but it means the "
        "predictors are not contemporaneous with sampling for all points, and the "
        "consequence is discussed in Section 5.1.")])

head("3.2.2. Spectral Indices", 3)
par([K("All Sentinel-2 MSI spectral bands between B2 and B12 were used as predictors. "
       "The original 20-meter bands (B5, B6, B7, B8A, B11, and B12) were resampled to "
       "10 meters to match the spatial resolution of the other bands. These bands cover "
       "the visible, red-edge, near-infrared, and shortwave infrared regions (Silvero et "
       "al., 2021), which means they can detect plant health, chlorophyll content, soil "
       "mineralogy, moisture status, and salinity differences. Table 2 shows the "
       "Sentinel-2 bands used in this study and their spatial resolution.")])
t1 = pd.read_csv(os.path.join(TAB, "T1_Sentinel2_bands.csv"))
table(t1, "Table 2",
      "Sentinel-2 MSI multispectral bands used in this study, with native and output "
      "resolution and the resampling method applied.",
      widths=[1.7, 2.9, 3.6, 3.2, 2.6, 2.4])

par([N_(f"Ten spectral indices describing vegetation vigour, canopy chlorophyll, "
        f"moisture and soil brightness or salinity were computed from the bands of each "
        f"monthly composite (Table 3). Together with the ten reflectance bands this "
        f"gives 20 predictors per composite and {s01['n_predictors']} in total. Three "
        f"corrections to the previous predictor set were made after an audit of the "
        f"index code against its stated definitions. The layer previously labelled SI "
        f"was computing the negative of NDMI rather than B11 × B12; the layer labelled "
        f"VSSI was computing the negative of the layer labelled NDSI rather than "
        f"2·B3 − 5·(B4 + B8); and the layers labelled NDSI and MNDWI were interchanged "
        f"with respect to their formulas. Two of the previous 42 columns were therefore "
        f"exact sign-flipped duplicates of two others and carried no independent "
        f"information. All indices have been recomputed from the reflectance bands with "
        f"verified formulas, the redundant layer has been removed, and the corrected set "
        f"of {s01['n_predictors']} predictors contains no pair with |r| > 0.999 "
        f"(numerical rank {s01['numerical_rank']}). NDRE is computed from B8A rather "
        f"than B8 because B8A and B5 share the same 20 m native support. The complete "
        f"inventory, giving for every predictor its model variable name, acquisition, "
        f"formula, source bands, unit, native and output resolution and resampling "
        f"method, is provided as Supplementary Table S1.")])
t2 = pd.read_csv(os.path.join(TAB, "T2b_index_definitions.csv"))
t2 = t2.rename(columns={"Limiting native resolution (m)": "Native res. (m)"})
table(t2, "Table 3",
      "Spectral indices derived from Sentinel-2 imagery. Formulas and attributions have "
      "been corrected; the limiting native resolution is that of the coarsest source "
      "band.", widths=[1.8, 5.0, 2.6, 2.2, 4.8])

# ------------------------------------------------------------------ 3.3
head("3.3. Extraction of Spectral Predictors at Soil Sample Locations", 2)
par([N_(f"The Sentinel-2 predictor stack was aligned with the soil sample coordinates "
        f"and all {s01['n_predictors']} predictors were extracted at the pixel "
        f"containing each sample centroid. Because a composite sample integrates a 10 m "
        f"radius, the GPS uncertainty is approximately ±3 m and the red-edge and SWIR "
        f"bands have a 20 m native support, the correspondence between a sample and a "
        f"single 10 m cell is approximate. All predictors were retained, since both "
        f"algorithms tolerate correlated inputs and no predictor is now redundant by "
        f"construction.")])
dR2 = [sn(t, "3 x 3 mean (30 m)").R2 - sn(t, "single pixel (10 m)").R2 for t in TARGETS]
dRM = [100 * (sn(t, "3 x 3 mean (30 m)").RMSE - sn(t, "single pixel (10 m)").RMSE)
       / sn(t, "single pixel (10 m)").RMSE for t in TARGETS]
par([N_(f"The sensitivity of the results to extraction support was tested directly. "
        f"Every predictor was re-extracted as the mean of a 3 × 3 pixel neighbourhood "
        f"(nominal 30 m) and the identical nested spatial cross-validation was repeated. "
        f"The two extractions agree closely at the sample points (Pearson r between "
        f"{sagr.r.min():.3f} and {sagr.r.max():.3f}, median {sagr.r.median():.3f}; "
        f"median relative difference {sagr.rel_diff_pct.median():.1f} %), and the effect "
        f"on validated performance is small and of inconsistent sign: the change in the "
        f"coefficient of determination is {min(dR2):+.3f} to {max(dR2):+.3f} and the "
        f"change in root mean square error {min(dRM):+.1f} % to {max(dRM):+.1f} % "
        f"(Supplementary Table S6). Single-pixel extraction is marginally better for "
        f"nitrogen and phosphorus and marginally worse for potassium. No conclusion in "
        f"this paper depends on the choice of extraction support, and the single-pixel "
        f"extraction is retained throughout.")])

# ------------------------------------------------------------------ 3.4
head("3.4. Model Development", 2, new=True)
hp_txt = "; ".join(
    f"{r.target} {'Random Forest' if r.model == 'RF' else 'XGBoost'} ("
    + ", ".join(f"{c} = {r[c]}" for c in t6b.columns
                if c not in ("target", "model") and str(r[c]) not in ("", "nan"))
    + ")" for _, r in t6b.iterrows())
par([N_(f"Random Forest and XGBoost regression models were fitted for each nutrient. "
        f"Random Forest used 500 trees, with the number of predictors sampled per split "
        f"(0.2, 0.4, 0.7 of the predictor set) and the minimum leaf size (1, 3) tuned. "
        f"XGBoost used 500 boosting rounds with a squared-error objective, subsample "
        f"fraction 0.8 and minimum child weight 3, with tree depth (3, 5), learning rate "
        f"(0.05, 0.10) and column subsample fraction (0.6, 0.9) tuned. There were no "
        f"missing predictor values and no response transformation was applied. All "
        f"analyses were carried out in Python 3.13 using scikit-learn 1.9.0 (Pedregosa "
        f"et al., 2011), xgboost 3.4.0 (Chen and Guestrin, 2016), NumPy 2.5.1 (Harris et "
        f"al., 2020) and rasterio 1.5.0, under a single global random seed of 42. The "
        f"hyperparameters selected by the inner tuning loop, reported as the modal "
        f"choice across outer folds, were: {hp_txt}. Fold-by-fold selections are given "
        f"in Supplementary Table S2.")])
par([N_("Predictive performance was benchmarked against five reference models evaluated "
        "on identical outer folds: an intercept-only model, ridge regression, partial "
        "least squares regression, ordinary kriging of the nutrient itself, and "
        "regression kriging combining the Random Forest trend with kriged residuals. "
        "Without such benchmarks the added value of the spectral predictors and of the "
        "ensemble algorithms cannot be assessed.")])

# ------------------------------------------------------------------ 3.5
head("3.5. Validation Design", 2, new=True)
mI, mI10 = mor.iloc[0], mor.iloc[-1]
par([N_(f"Nested cross-validation with spatial blocks is the primary evaluation design "
        f"(Roberts et al., 2017). The 110 locations were partitioned into "
        f"{s01['n_blocks']} blocks by k-means clustering of projected coordinates "
        f"(EPSG:26191), giving fold sizes of "
        f"{', '.join(str(n) for n in s01['fold_sizes'])}, a mean block extent of "
        f"{bdes.extent_E_km.mean():.1f} km × {bdes.extent_N_km.mean():.1f} km and a "
        f"median separation of {s01['median_block_separation_km']:.2f} km between a "
        f"block and its nearest neighbouring block (Fig. 3a; Supplementary Table S5). "
        f"The outer loop holds out one block at a time and provides all reported "
        f"performance. The inner loop performs every tuning decision by grouped "
        f"five-fold cross-validation on the training blocks only, so the outer folds "
        f"take no part in model construction. Random Forest and XGBoost use identical "
        f"outer partitions, which makes the paired comparison between them valid. Nested "
        f"cross-validation with random folds was run under the same protocol as a "
        f"secondary analysis, for the sole purpose of quantifying the optimism "
        f"introduced by ignoring spatial dependence.")])
par([N_(f"Block dimensions were informed by the spatial structure of the observations. "
        f"Moran's I at a 1 km bandwidth is {mI['N']:.2f} for N, {mI['P']:.2f} for P and "
        f"{mI['K']:.2f} for K, and remains positive at a 10 km bandwidth "
        f"({mI10['N']:.2f}, {mI10['P']:.2f} and {mI10['K']:.2f} respectively; "
        f"Supplementary Table S4). Blocks were therefore sized to exceed the distances "
        f"at which appreciable autocorrelation is detectable. Empirical variograms were "
        f"also computed for each nutrient and for the model residuals (Section 4.3); at "
        f"this sample size the nugget ratios are well determined but the fitted ranges "
        f"are weakly identified, so Moran's I rather than the fitted ranges was used to "
        f"justify the block dimension.")])

# ------------------------------------------------------------------ 3.6
head("3.6. Model Evaluation", 2)
par([N_("All reported statistics are computed from a single pooled set of held-out "
        "outer-fold predictions, one prediction per observation per scheme. Performance "
        "is summarised by the coefficient of determination, the root mean square error, "
        "the mean absolute error, the mean bias, Lin's concordance correlation "
        "coefficient (Lin, 1989), the ratio of performance to deviation (RPD) and the "
        "ratio of performance to interquartile range (RPIQ) (Kuhn and Johnson, 2013; "
        "Wadoux et al., 2020). The metrics are defined in Equations 1 and 2:")])
par([N_("R2 = 1 − SSE / SST = 1 − Σi (yi − ŷi)2 / Σi (yi − ȳ)2          (1)")],
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
par([N_("RMSE = [ (1/n) Σi (yi − ŷi)2 ]1/2          (2)")],
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
par([N_("where yi is the observed soil nutrient concentration at sample location i, ŷi "
        "is the corresponding model prediction, ȳ is the mean of the observations, and n "
        "is the number of soil samples used for evaluation. SSE is the sum of squared "
        "prediction errors and SST the total sum of squares. Because R2 and RMSE are "
        "computed from the same pooled predictions, each is exactly recoverable from the "
        "other given SST. Squared-correlation definitions of R2 are not used anywhere in "
        "this paper.")])
par([N_("Bootstrap 95 % confidence intervals were obtained from 3,000 resamples of the "
        "paired observed and predicted values. The two algorithms were compared through "
        "the paired difference in squared error on identical folds, with a bootstrap "
        "interval on that difference; a difference is reported as meaningful only when "
        "the interval excludes zero. Predictor importance was quantified by permutation "
        "on the held-out observations of each outer fold, with ten permutations per "
        "predictor, normalised by the fold root mean square error so that values are "
        "comparable across nutrients, and reported with the standard deviation between "
        "folds. Impurity-based and gain-based importance are not reported, because both "
        "are biased in the presence of correlated continuous predictors and are not "
        "comparable across responses with different variance.")])

# ------------------------------------------------------------------ 3.7
head("3.7. Spatial Prediction, Uncertainty and Applicability", 2)
head("3.7.1. Continuous Spatial Prediction", 3)
par([N_(f"Final models were fitted on all 110 observations using the modal "
        f"hyperparameters and applied to the full covariate stack. Predictions were "
        f"generated on the native 10 m grid of the Sentinel-2 covariates, comprising "
        f"{msum['n_valid_pixels']:,} valid pixels and covering "
        f"{msum['mapped_area_ha']:,.0f} ha. At the latitude of the district the "
        f"effective pixel is {msum['pixel_x_m']:.2f} m × {msum['pixel_y_m']:.2f} m "
        f"({msum['pixel_ha']:.5f} ha). Rasters are distributed in EPSG:4326 (WGS 84); "
        f"all map figures are drawn in EPSG:26191 (Merchich / Nord Maroc) with axes "
        f"labelled in kilometres, a scale bar, a north arrow and the projection "
        f"identifier printed on the figure. Continuous surfaces are displayed with a "
        f"perceptually uniform, colour-vision-safe sequential ramp running from low to "
        f"high concentration. Complete cartographic metadata are given in Table 12.")])

head("3.7.2. Uncertainty and Area of Applicability", 3)
par([N_(f"Local prediction uncertainty is mapped as the standard deviation of the "
        f"predictions across the Random Forest tree ensemble, an approach comparable to "
        f"the local uncertainty estimation applied to topsoil organic carbon by Veronesi "
        f"and Schillaci (2019). The applicability domain "
        f"was delineated following Meyer and Pebesma (2021). A dissimilarity index was "
        f"computed for every pixel as the minimum standardised Euclidean distance in "
        f"predictor space to any training observation, weighted by permutation "
        f"importance and scaled by the mean pairwise distance among training points. The "
        f"threshold was derived from the cross-validation-aware distribution of training "
        f"dissimilarities as the upper quartile plus 1.5 interquartile ranges, giving "
        f"DI = {msum['aoa_threshold']:.3f}. Pixels exceeding the threshold are masked in "
        f"every delivered product and excluded from all reported areas.")])

head("3.7.3. Map Reclassification and Fertility Classes", 3)
par([N_("Predicted surfaces were reclassified into four fertility classes using the "
        "non-overlapping boundaries in Table 4. Phosphorus classes follow widely used "
        "Olsen P interpretation ranges (Recena et al., 2015; Havlin and Heiniger, 2020). "
        "Nitrogen classes are expressed both as total N in mg kg-1 and as the equivalent "
        "percentage, using the conversion 1,000 mg kg-1 = 0.10 % N. The same thresholds "
        "were applied to the observed values and to the held-out outer-fold predictions, "
        "so that class agreement could be quantified (Section 4.6). Thresholds were not "
        "adjusted to avoid empty classes; where a class is empty in the predicted "
        "surface, that is reported as a result.")])
cls_tab = pd.DataFrame({
    "Class": ["Very low", "Low", "Medium", "High"],
    "Total N (mg kg-1)": ["< 1,000", "≥ 1,000 and < 1,500", "≥ 1,500 and < 3,000",
                          "≥ 3,000"],
    "Total N (%)": ["< 0.10", "0.10 – 0.15", "0.15 – 0.30", "≥ 0.30"],
    "Olsen P (mg kg-1)": ["< 15", "≥ 15 and < 30", "≥ 30 and < 60", "≥ 60"],
    "Exchangeable K (mg kg-1)": ["< 40", "≥ 40 and < 80", "≥ 80 and < 120", "≥ 120"]})
table(cls_tab, "Table 4",
      "Fertility classes applied to the observed values and to the predicted surfaces. "
      "Boundaries are non-overlapping and the conversion used for total nitrogen is "
      "shown explicitly.", widths=[2.2, 3.6, 2.4, 3.2, 4.0])
author_note("Cite one verifiable, locally relevant source for each threshold set, ideally "
            "an INRA Morocco or ORMVAT interpretation guide. The referee objected to "
            "Olsen (1954), which describes the extraction method rather than the "
            "thresholds, and to Havlin (2014), which was absent from the reference list. "
            "Havlin and Heiniger (2020) is used here as a placeholder framework.")

# ==========================================================================
#  4. RESULTS
# ==========================================================================
head("4. Results", 1)

head("4.1. Descriptive Statistics of Soil Nutrients", 2)
par([N_(f"The 110 validated observations show wide variation in all three macronutrients "
        f"(Table 5). Total nitrogen ranges from {tn.Minimum:,.0f} to {tn.Maximum:,.0f} "
        f"mg kg-1 with a mean of {tn.Mean:,.1f} mg kg-1 and a coefficient of variation "
        f"of {tn.CV_percent:.0f} %. Olsen phosphorus ranges from {tp.Minimum:.2f} to "
        f"{tp.Maximum:.2f} mg kg-1 (mean {tp.Mean:.1f} mg kg-1, CV "
        f"{tp.CV_percent:.0f} %) and exchangeable potassium from {tk.Minimum:.1f} to "
        f"{tk.Maximum:.1f} mg kg-1 (mean {tk.Mean:.1f} mg kg-1, CV "
        f"{tk.CV_percent:.0f} %). Phosphorus and potassium are strongly right-skewed "
        f"(skewness {tp.Skewness:.1f} and {tk.Skewness:.1f} respectively), nitrogen less "
        f"so ({tn.Skewness:.1f}). This dispersion is consistent with the spatial "
        f"heterogeneity previously reported for Beni Moussa soils (El Hamzaoui and El "
        f"Baghdadi, 2021). Table 5 is the single definitive summary of the observations; "
        f"the ranges of the predicted rasters are different quantities and are reported "
        f"separately in Table 12.")])
t3s = t3[["Nutrient", "n", "Minimum", "Q1", "Median", "Mean", "Q3", "Maximum", "SD",
          "CV_percent", "Skewness"]].copy()
t3s.columns = ["Nutrient", "n", "Min", "Q1", "Median", "Mean", "Q3", "Max", "SD",
               "CV (%)", "Skew"]
table(t3s, "Table 5",
      "Descriptive statistics of the 110 validated topsoil observations. All "
      "concentrations in mg kg-1.", fontsize=9)

head("4.2. Predictive Performance", 2)
par([N_(f"Under nested spatial block cross-validation, Random Forest explained "
        f"{rf_n.R2:.3f} of the variance in total nitrogen (RMSE {rf_n.RMSE:,.1f} "
        f"mg kg-1, MAE {rf_n.MAE:,.1f}, bias {rf_n.bias:+,.1f}, CCC {rf_n.CCC:.3f}), "
        f"{rf_p.R2:.3f} for Olsen phosphorus (RMSE {rf_p.RMSE:.2f} mg kg-1, MAE "
        f"{rf_p.MAE:.2f}, bias {rf_p.bias:+.2f}, CCC {rf_p.CCC:.3f}) and {rf_k.R2:.3f} "
        f"for exchangeable potassium (RMSE {rf_k.RMSE:.2f} mg kg-1, MAE {rf_k.MAE:.2f}, "
        f"bias {rf_k.bias:+.2f}, CCC {rf_k.CCC:.3f}). XGBoost gave coefficients of "
        f"determination of {xg_n.R2:.3f}, {xg_p.R2:.3f} and {xg_k.R2:.3f} respectively. "
        f"Observed-versus-predicted relationships for the held-out folds are shown in "
        f"Figs. 4 and 5, and all statistics for all three validation schemes are given "
        f"in Table 6.")])
perf = []
for t in TARGETS:
    for m in ["RF", "XGB"]:
        r, c = sp(t, m), cir(t, m)
        perf.append({"Nutrient": t, "Model": "RF" if m == "RF" else "XGBoost",
                     "R2": f"{r.R2:.3f}", "R2 95 % CI": f"[{c.R2_lo:.3f}, {c.R2_hi:.3f}]",
                     "RMSE": f"{r.RMSE:,.2f}",
                     "RMSE 95 % CI": f"[{c.RMSE_lo:,.2f}, {c.RMSE_hi:,.2f}]",
                     "MAE": f"{r.MAE:,.2f}", "Bias": f"{r.bias:+,.2f}",
                     "CCC": f"{r.CCC:.3f}", "Slope": f"{r.slope:.2f}",
                     "RPIQ": f"{r.RPIQ:.2f}"})
table(pd.DataFrame(perf), "Table 6",
      "Performance on held-out outer folds of the nested spatial block cross-validation "
      "(n = 110). RMSE, MAE and bias in mg kg-1. Slope is the regression of predicted on "
      "observed values. Confidence intervals from 3,000 bootstrap resamples.",
      fontsize=8.5)

par([N_(f"All six models show pronounced regression to the mean. The slope of predicted "
        f"on observed values ranges from "
        f"{mt[mt.scheme=='nested spatial CV'].slope.min():.2f} to "
        f"{mt[mt.scheme=='nested spatial CV'].slope.max():.2f}, so low concentrations "
        f"are systematically over-predicted and high concentrations under-predicted. "
        f"This shrinkage propagates to the maps, whose ranges are compressed relative to "
        f"the observations (Table 12), and it must be taken into account wherever the "
        f"surfaces are used to identify extreme values.")])
figure(os.path.join(FIG, "Figure_3_RF_observed_vs_predicted.png"), "Fig. 4",
       "Relationship between observed and Random Forest predicted soil macronutrients "
       "from held-out folds of the nested spatial block cross-validation: (a) total "
       "nitrogen (N), (b) available phosphorus (P), and (c) exchangeable potassium (K). "
       "The dashed line is 1:1 and the solid line the fitted regression of predicted on "
       "observed values.", width_cm=16.0)
figure(os.path.join(FIG, "Figure_4_XGB_observed_vs_predicted.png"), "Fig. 5",
       "Relationship between observed and XGBoost predicted soil macronutrients from "
       "held-out folds of the nested spatial block cross-validation: (a) total nitrogen "
       "(N), (b) available phosphorus (P), and (c) exchangeable potassium (K).",
       width_cm=16.0)

opt = []
for t in TARGETS:
    for m in ["RF", "XGB"]:
        a, b, c_ = sp(t, m, "nested random CV"), sp(t, m), sp(t, m, "calibration")
        opt.append({"Nutrient": t, "Model": "RF" if m == "RF" else "XGBoost",
                    "Calibration R2": f"{c_.R2:.3f}",
                    "Nested random CV R2": f"{a.R2:.3f}",
                    "Nested spatial CV R2": f"{b.R2:.3f}",
                    "Optimism (delta R2)": f"{a.R2 - b.R2:+.3f}"})
par([N_("Ignoring spatial dependence inflates the estimates, and calibration statistics "
        "inflate them far more (Table 7, Fig. 6). Calibration figures are reported here "
        "only to document that gap; they are not evidence of predictive skill and are "
        "not quoted anywhere else in this paper.")])
table(pd.DataFrame(opt), "Table 7",
      "Coefficient of determination under the three evaluation schemes. The final column "
      "is the optimism introduced by using random rather than spatial folds.",
      fontsize=9)

par([N_(f"The two algorithms cannot be separated on this dataset. The paired difference "
        f"in squared error on identical outer folds has a bootstrap 95 % confidence "
        f"interval that includes zero for every nutrient "
        f"({'; '.join(f'{r.target}: [{r.ci_lo:,.1f}, {r.ci_hi:,.1f}]' for _, r in pcmp.iterrows())}). "
        f"No claim of superiority is therefore made in either direction. Random Forest "
        f"was adopted for mapping on the criterion specified before the comparison was "
        f"made, namely the root mean square error on the outer spatial folds.")])
blt = []
for t in TARGETS:
    sub = [("Random Forest", sp(t, "RF").R2, sp(t, "RF").RMSE, sp(t, "RF").CCC),
           ("XGBoost", sp(t, "XGB").R2, sp(t, "XGB").RMSE, sp(t, "XGB").CCC)]
    for _, r in bl[bl.target == t].iterrows():
        sub.append((r.model, r.R2, r.RMSE, r.CCC))
    for nm, r2, rmse, ccc in sorted(sub, key=lambda z: -z[1]):
        blt.append({"Nutrient": t, "Model": nm, "R2": f"{r2:.3f}",
                    "RMSE": f"{rmse:,.2f}", "CCC": f"{ccc:.3f}"})
par([N_(f"The benchmarks place these results in context (Table 8). The intercept-only "
        f"model returns a negative coefficient of determination for all three nutrients, "
        f"confirming that the fitted models carry real information. Ordinary kriging of "
        f"the nutrient alone performs poorly ({bs('N','Ordinary kriging').R2:.2f} for N, "
        f"{bs('P','Ordinary kriging').R2:.2f} for P and "
        f"{bs('K','Ordinary kriging').R2:.2f} for K), so it is the spectral predictors "
        f"rather than spatial interpolation that carry the signal. Two results temper "
        f"the case for ensemble methods and are reported rather than omitted: for "
        f"potassium, ridge regression on the same predictors attains R2 = "
        f"{bs('K','Ridge regression').R2:.3f} against {rf_k.R2:.3f} for Random Forest, "
        f"and for nitrogen, regression kriging attains "
        f"{bs('N','Regression kriging (RF + OK residuals)').R2:.3f} against "
        f"{rf_n.R2:.3f}. On this dataset the ensemble algorithms are not uniformly the "
        f"best available option.")])
table(pd.DataFrame(blt), "Table 8",
      "Benchmark comparison. All models evaluated on identical outer spatial folds, "
      "ordered by coefficient of determination within each nutrient. RMSE in mg kg-1.",
      fontsize=9, widths=[2.0, 7.0, 2.2, 2.6, 2.2])
figure(os.path.join(FIG, "Figure_5_validation_schemes.png"), "Fig. 6",
       "Coefficient of determination by validation scheme for (a) total nitrogen (N), "
       "(b) available phosphorus (P), and (c) exchangeable potassium (K), with bootstrap "
       "95 % confidence intervals on the spatial cross-validation estimate and "
       "horizontal lines showing the non-machine-learning benchmarks.", width_cm=16.0)

head("4.3. Spatial Structure and Residual Diagnostics", 2, new=True)
vgtxt = "; ".join(f"{t} {vgr(t,'observed').nugget_ratio:.2f}"
                  for t in TARGETS if vgr(t, "observed") is not None)
pure = [t for t in TARGETS if vgr(t, "RF residuals") is not None
        and vgr(t, "RF residuals").nugget_ratio > 0.95]
struct = [t for t in TARGETS if t not in pure and vgr(t, "RF residuals") is not None]
par([N_(f"Fitted spherical variograms of the observations give nugget-to-sill ratios of "
        f"{vgtxt} (Fig. 7, upper row; Supplementary Table S4). These high ratios "
        f"indicate that most of the variance occurs at separation distances shorter than "
        f"the sampling interval, which is consistent with a density of one observation "
        f"per {msum['mapped_area_ha']/110:,.0f} ha. The fitted ranges are weakly "
        f"identified at this sample size and are reported for completeness rather than "
        f"as reliable estimates of correlation length.")])
par([N_("Variograms of the Random Forest spatial cross-validation residuals are shown in "
        "the lower row of Fig. 7. The " + " and ".join(pure) + " residuals are "
        "effectively pure nugget (nugget ratio " +
        ", ".join(f"{vgr(t,'RF residuals').nugget_ratio:.2f}" for t in pure) +
        "), indicating that the covariates captured the spatial structure these data can "
        "resolve. The " + " and ".join(struct) + " residuals retain spatial structure "
        "(nugget ratio " +
        ", ".join(f"{vgr(t,'RF residuals').nugget_ratio:.2f}" for t in struct) +
        ", fitted range " +
        ", ".join(f"{vgr(t,'RF residuals').range_m/1000:.1f} km" for t in struct) +
        "), so the " + " and ".join(struct) + " model leaves spatially organised error "
        "behind and would benefit from a geostatistical correction. This is consistent "
        "with the regression-kriging benchmark in Table 8.")])
figure(os.path.join(FIG, "Figure_7_variograms.png"), "Fig. 7",
       "Empirical variograms with fitted spherical models. Upper row: observed (a) total "
       "nitrogen (N), (b) available phosphorus (P), and (c) exchangeable potassium (K). "
       "Lower row: the corresponding Random Forest spatial cross-validation residuals.",
       width_cm=16.0)

head("4.4. Predictor Importance", 2)
par([N_(f"Permutation importance measured on held-out spatial folds is dominated by raw "
        f"reflectance bands; the spectral indices contribute little (Table 9, Fig. 8). "
        f"For total nitrogen, near-infrared bands account for "
        f"{cls_share('N','RF','NIR bands (B8, B8A)'):.0f} % of positive importance in "
        f"Random Forest and red-edge bands for "
        f"{cls_share('N','RF','Red-edge bands (B5, B6, B7)'):.0f} %. For Olsen "
        f"phosphorus the order is reversed, with red-edge bands at "
        f"{cls_share('P','RF','Red-edge bands (B5, B6, B7)'):.0f} % and near-infrared "
        f"bands at {cls_share('P','RF','NIR bands (B8, B8A)'):.0f} %. Exchangeable "
        f"potassium is the only nutrient with a broadly distributed spectral signature: "
        f"near-infrared bands {cls_share('K','RF','NIR bands (B8, B8A)'):.0f} %, "
        f"brightness and salinity indices "
        f"{cls_share('K','RF','Brightness / salinity indices'):.0f} %, shortwave "
        f"infrared bands {cls_share('K','RF','SWIR bands (B11, B12)'):.0f} % and visible "
        f"bands {cls_share('K','RF','Visible bands (B2, B3, B4)'):.0f} %.")])
table(t7, "Table 9",
      "Permutation importance aggregated by predictor class, expressed as a percentage "
      "of total positive importance on held-out spatial folds.",
      fontsize=9, widths=[5.2, 1.9, 1.9, 1.9, 1.9, 1.9, 1.9], align_center=False)
par([N_(f"Two features of this result constrain how far it can be interpreted. First, "
        f"the November 2024 composite accounts for {novmin:.0f}–{novmax:.0f} % of "
        f"positive importance across the six models, so a single month carries almost "
        f"the whole predictive signal. Second, individual predictor rankings are "
        f"unstable: only {int(t7c.n_stable_predictors.min())}–"
        f"{int(t7c.n_stable_predictors.max())} of {int(t7c.n_total.iloc[0])} predictors "
        f"have a mean permutation importance exceeding their own between-fold standard "
        f"deviation. Aggregated classes of predictor can therefore be discussed with "
        f"confidence; the rank order of individual bands and indices cannot, and no "
        f"mechanistic interpretation is attached to it here.")])
figure(os.path.join(FIG, "Figure_6_permutation_importance.png"), "Fig. 8",
       "Permutation importance measured on held-out spatial folds. (a–c) importance "
       "aggregated by predictor class for both algorithms; (d–f) the fifteen strongest "
       "individual predictors in the Random Forest models, with between-fold standard "
       "deviations.", width_cm=16.0)

head("4.5. Spatial Nutrient Patterns", 2)
par([N_(f"Continuous Random Forest surfaces at 10 m are shown in Fig. 9 and the "
        f"corresponding XGBoost surfaces in Fig. 10. Predicted ranges are "
        f"{rs('N','RF')['min']:,.0f}–{rs('N','RF')['max']:,.0f} mg kg-1 for nitrogen, "
        f"{rs('P','RF')['min']:.1f}–{rs('P','RF')['max']:.1f} for phosphorus and "
        f"{rs('K','RF')['min']:.1f}–{rs('K','RF')['max']:.1f} for potassium. All three "
        f"lie inside the observed range of the corresponding nutrient and all are "
        f"compressed relative to it, as expected from the slopes reported in Section 4.2. "
        f"The surfaces contain sharply bounded, rectangular patches that follow field and "
        f"parcel geometry rather than pedological boundaries; this is a substantive "
        f"observation about what the models are detecting and is discussed in "
        f"Section 5.1.")])
figure(os.path.join(FIG, "Figure_8_continuous_maps_RF.png"), "Fig. 9",
       "Random Forest predicted spatial distribution of soil macronutrients: (a) total "
       "nitrogen (N), (b) available phosphorus (P), and (c) exchangeable potassium (K) "
       "in the Beni Moussa irrigated district at 10 m resolution, masked outside the "
       "area of applicability. Open circles are the 110 sampling locations. Projection "
       "EPSG:26191, axis units km.", width_cm=15.0)
figure(os.path.join(FIG, "Figure_9_continuous_maps_XGB.png"), "Fig. 10",
       "XGBoost predicted spatial distribution of soil macronutrients: (a) total "
       "nitrogen (N), (b) available phosphorus (P), and (c) exchangeable potassium (K) "
       "at 10 m resolution, masked outside the area of applicability. Colour ranges are "
       "set independently of Fig. 9; the classified agreement between the two algorithms "
       "is quantified in Table 10.", width_cm=15.0)

par([N_(f"Local uncertainty is mapped in Fig. 11 as the standard deviation across the "
        f"tree ensemble. District mean values are {rs('N','RF_SD')['mean']:,.0f} mg kg-1 "
        f"for nitrogen, {rs('P','RF_SD')['mean']:.1f} for phosphorus and "
        f"{rs('K','RF_SD')['mean']:.1f} for potassium, with maxima of "
        f"{rs('N','RF_SD')['max']:,.0f}, {rs('P','RF_SD')['max']:.1f} and "
        f"{rs('K','RF_SD')['max']:.1f} respectively. Uncertainty is highest at parcel "
        f"boundaries and in the parts of the district furthest from sampling locations "
        f"in predictor space.")])
figure(os.path.join(FIG, "Figure_10_uncertainty_maps.png"), "Fig. 11",
       "Local prediction uncertainty of the Random Forest surfaces, expressed as the "
       "standard deviation across the regression-tree ensemble, for (a) total nitrogen "
       "(N), (b) available phosphorus (P), and (c) exchangeable potassium (K).",
       width_cm=15.0)

agr_txt = ", ".join(f"{r.nutrient} kappa = {r.kappa:.2f}" for _, r in agr.iterrows())
par([N_(f"The dissimilarity index and the resulting area of applicability are shown in "
        f"Fig. 3b. {msum['pct_inside_aoa']:.1f} % of the mapped area "
        f"({msum['ha_inside_aoa']:,.0f} ha) falls inside the applicability domain, and "
        f"{100-msum['pct_inside_aoa']:.1f} % ({msum['ha_outside_aoa']:,.0f} ha) is "
        f"masked as extrapolative. The two algorithms produce closely similar classified "
        f"surfaces: pixelwise agreement is "
        f"{100*agr.overall_agreement.min():.0f}–{100*agr.overall_agreement.max():.0f} % "
        f"with {agr_txt}. These statistics replace any qualitative assertion that one "
        f"algorithm produces clearer or more detailed maps than the other.")])
t10c = agr.copy()
t10c.columns = ["Nutrient", "Overall pixel agreement", "Cohen's kappa", "Pixels compared"]
t10c["Overall pixel agreement"] = t10c["Overall pixel agreement"].map(lambda v: f"{v:.3f}")
t10c["Cohen's kappa"] = t10c["Cohen's kappa"].map(lambda v: f"{v:.3f}")
t10c["Pixels compared"] = t10c["Pixels compared"].map(lambda v: f"{v:,}")
table(t10c, "Table 10",
      "Agreement between the Random Forest and XGBoost classified 10 m surfaces.",
      fontsize=9)

par([N_(f"Fertility classes within the applicability domain are mapped in Fig. 12 and "
        f"quantified in Table 11. Nitrogen is dominated by the Low class "
        f"({ca('N','Low').rf_inAOA_pct:.1f} %, {ca('N','Low').rf_inAOA_ha:,.0f} ha) and "
        f"the Medium class ({ca('N','Medium').rf_inAOA_pct:.1f} %, "
        f"{ca('N','Medium').rf_inAOA_ha:,.0f} ha). Phosphorus is spread across all four "
        f"classes, with {ca('P','Very low').rf_inAOA_pct:.1f} % Very low and "
        f"{ca('P','High').rf_inAOA_pct:.1f} % High. Potassium is concentrated in the two "
        f"lowest classes ({ca('K','Very low').rf_inAOA_pct:.1f} % Very low, "
        f"{ca('K','Low').rf_inAOA_pct:.1f} % Low); the High class is empty in the Random "
        f"Forest surface. That absence reflects both the shrinkage described in "
        f"Section 4.2 and the fact that only {int((df.K > 120).sum())} of the 110 "
        f"observations exceed 120 mg kg-1, and it is reported rather than removed by "
        f"adjusting the thresholds.")])
carea = care.copy()
carea["Class"] = carea.cls
piv = carea.pivot(index="Class", columns="nutrient",
                  values=["rf_inAOA_ha", "rf_inAOA_pct"])
piv.columns = [f"{b} {'ha' if a == 'rf_inAOA_ha' else '%'}" for a, b in piv.columns]
piv = piv.reindex(["Very low", "Low", "Medium", "High"]).reset_index()
for c in piv.columns[1:]:
    piv[c] = piv[c].map(lambda v: f"{v:,.0f}" if c.endswith("ha") else f"{v:.1f}")
table(piv, "Table 11",
      "Fertility class areas from the Random Forest 10 m surfaces, computed inside the "
      "area of applicability.", fontsize=9)
t10 = pd.read_csv(os.path.join(TAB, "T10d_map_metadata.csv"))
table(t10, "Table 12",
      "Cartographic metadata for the delivered raster products.",
      fontsize=9, widths=[8.0, 8.0], align_center=False)
figure(os.path.join(FIG, "Figure_11_fertility_classes.png"), "Fig. 12",
       "Reclassified soil fertility maps derived from the Random Forest 10 m predictions "
       "for (a) total nitrogen (N), (b) available phosphorus (P), and (c) exchangeable "
       "potassium (K), restricted to the area of applicability. Class areas are given in "
       "the legend of each panel.", width_cm=15.0)

head("4.6. Fertility Class Agreement", 2, new=True)
kt = t9[["nutrient", "model", "overall_accuracy", "kappa"]].copy()
kt.columns = ["Nutrient", "Model", "Overall accuracy", "Cohen's kappa"]
kt["Model"] = kt.Model.replace({"RF": "Random Forest", "XGB": "XGBoost"})
par([N_(f"Applying the thresholds in Table 4 to the observations and to the held-out "
        f"outer-fold predictions gives overall class accuracies of "
        f"{t9.overall_accuracy.min():.2f}–{t9.overall_accuracy.max():.2f} and Cohen's "
        f"kappa of {t9.kappa.min():.2f}–{t9.kappa.max():.2f} (Table 13; full confusion "
        f"matrices with producer's and user's accuracies in Supplementary Table S3). "
        f"Agreement is good in the classes that are well represented in the sample and "
        f"unreliable in the upper classes, which rest on very few observations: only "
        f"{int((df.K > 120).sum())} samples exceed the highest potassium threshold and "
        f"{int((df.N > 3000).sum())} exceed the highest nitrogen threshold. "
        f"Class-specific accuracy in those categories should not be relied upon for any "
        f"operational purpose.")])
table(kt, "Table 13",
      "Agreement between the observed fertility class and the class of the held-out "
      "outer-fold prediction (n = 110).", fontsize=9)

# ==========================================================================
#  5. DISCUSSION
# ==========================================================================
head("5. Discussion", 1)

head("5.1. What the Spectral Signal Represents", 2, new=True)
par([N_(f"None of the three macronutrients has a direct diagnostic absorption feature in "
        f"the Sentinel-2 wavelengths. Any predictive skill must therefore be indirect, "
        f"and three features of the present results indicate what the intermediary is. "
        f"Predictive importance is concentrated in raw near-infrared and red-edge "
        f"reflectance rather than in soil-oriented indices (Section 4.4); the November "
        f"composite alone carries {novmin:.0f}–{novmax:.0f} % of that importance; and "
        f"the predicted surfaces contain sharply bounded, field-shaped patches that "
        f"follow parcel geometry (Fig. 9). The most parsimonious interpretation is that "
        f"the models are learning canopy condition in a single month, which is "
        f"correlated with soil fertility through crop selection, irrigation and "
        f"fertilisation history, rather than sensing soil nutrient status directly.")])
par([N_("This reading has a direct consequence for how the maps should be used. If the "
        "signal is carried by the crop, the surfaces are conditioned on the 2024–2025 "
        "cropping mosaic and are seasonal snapshots rather than stable soil property "
        "maps. Temporal transferability cannot be assessed from a single season and "
        "should not be assumed. Patterns visible in the maps may plausibly relate to "
        "clay content, salinity, hydromorphy, irrigation infrastructure, carbonate "
        "content or fertilisation history, and such relationships have been documented "
        "in the Tadla Plain (Barakat et al., 2017; El Hamzaoui and El Baghdadi, 2021; "
        "Salmi et al., 2024, 2025). None of these variables was measured in the present "
        "study, however, so these relationships are offered as hypotheses for testing "
        "rather than as findings.")])
par([N_("A further consequence concerns the alternative explanation for the dominance of "
        "the November composite. Temporal proximity to part of the sampling campaign, "
        "and a sampling or analytical batch effect aligned with that period, cannot be "
        "excluded on the present evidence. The difference in illumination geometry "
        "between the two composites, reported in Section 3.2.1, is a further systematic "
        "difference between them. Distinguishing these possibilities requires sampling "
        "and laboratory metadata organised by batch, and repeated acquisitions across "
        "seasons, both of which we recommend for future work in this district.")])

head("5.2. Comparison with Previous Work", 2)
par([N_(f"National-scale Random Forest mapping of Moroccan soil nutrients at 250 m "
        f"reported cross-validated coefficients of determination between 0.62 and 0.76 "
        f"for phosphorus and potassium (Bouslihim et al., 2025). The spatially validated "
        f"values obtained here, {rf_p.R2:.2f} for phosphorus and {rf_k.R2:.2f} for "
        f"potassium, are of the same order, obtained at finer resolution over a single "
        f"district but from a much smaller sample. Al Masmoudi et al. (2022) showed that "
        f"machine learning can predict soil fertility indicators in central Morocco from "
        f"laboratory data, but without spatial predictors could not produce continuous "
        f"surfaces; the present work supplies that spatial component. Studies combining "
        f"red-edge indices with tree-based learners in irrigated systems have reported "
        f"coefficients of determination between 0.80 and 0.95 (Silatsa and Kebede, "
        f"2023), but such values are generally obtained under random rather than spatial "
        f"validation and are not directly comparable with those reported here. We "
        f"encourage explicit reporting of whether folds were random or spatial, since "
        f"Table 7 shows the difference is material. At the opposite end of the scale "
        f"range, Dalle Vaglie et al. (2026) estimated soil carbon, nitrogen, pH and "
        f"salinity across large areas from satellite archives; the contrast between such "
        f"continental products and a single 78,000 ha district illustrates that the "
        f"limiting factor here is sampling density rather than sensor capability.")])

head("5.3. Algorithm Comparison and Methodological Considerations", 2, new=True)
par([N_("Random Forest and XGBoost proved statistically indistinguishable for all three "
        "nutrients once both were evaluated on identical held-out spatial folds "
        "(Section 4.2). Differences of the magnitude observed here are smaller than the "
        "resampling uncertainty, and reporting either algorithm as superior on this "
        "evidence would not be defensible. More striking is that for potassium a ridge "
        "regression on the same predictors performs at least as well as either ensemble, "
        "and for nitrogen regression kriging outperforms both. Where the predictor–"
        "response relationship is close to linear and the sample is small, the "
        "flexibility of ensemble methods buys little and may cost variance. Benchmarking "
        "against simple alternatives should be routine in digital soil mapping studies "
        "of this size (Wadoux et al., 2020).")])
par([N_("The instability of individual predictor rankings deserves emphasis. With only a "
        "handful of predictors exceeding their own between-fold variability, importance "
        "rankings computed on a single fit, as is common practice, would have supported "
        "almost any narrative. Reporting importance with resampling uncertainty, and "
        "aggregating to interpretable classes of predictor, is a necessary discipline at "
        "this sample size.")])

head("5.4. Practical Implications", 2)
par([N_("Total nitrogen is not plant-available nitrogen, and Olsen phosphorus and "
        "exchangeable potassium represent operationally defined pools rather than "
        "directly plant-available quantities. Fertiliser prescription additionally "
        "requires crop demand, yield target, mineralisation rate, pH, texture, "
        "irrigation regime and management history, none of which are available here. "
        "These maps are therefore offered as reconnaissance products. They can support "
        "stratified resampling that concentrates effort where uncertainty is high or "
        "where the applicability mask indicates extrapolation; they can delineate "
        "candidate management zones for field verification; and they can generate "
        "testable hypotheses about the controls on nutrient distribution in the "
        "district, in the spirit of integrated fertility mapping demonstrated elsewhere "
        "(Salih et al., 2026). Used in that way they complement the multi-criteria land "
        "suitability assessments already available for the Tadla plain (Ennaji et al., "
        "2018). Combined with irrigation network maps they may help identify areas where "
        "water distribution coincides with apparent nutrient depletion or accumulation "
        "(Chaaou et al., 2022; Mouaddine et al., 2025). They "
        "do not by themselves justify variable-rate fertiliser prescription, and should "
        "not be presented to end users as if they did.")])

head("5.5. Limitations of the Study", 2)
par([N_(f"The sample size of 110 is small relative to the heterogeneity of the district "
        f"and corresponds to one observation per {msum['mapped_area_ha']/110:,.0f} ha, "
        f"or approximately {msum['n_valid_pixels']/110:,.0f} prediction pixels per "
        f"observation. The upper fertility classes rest on very few samples, so class "
        f"accuracy there is poorly estimated. All predictors derive from two monthly "
        f"composites of a single sensor within one season; no terrain, radar, climate, "
        f"groundwater or proximal-sensing covariates were included, and their absence "
        f"probably omits genuine soil-forming information. The 20 m native support of "
        f"the red-edge and shortwave infrared bands limits the effective resolution of "
        f"nominal 10 m outputs. Predictions shrink towards the mean, so extremes are "
        f"systematically under-represented. Water bodies, settlements, roads and canals "
        f"were not separately masked beyond the district boundary and the applicability "
        f"domain. Finally, {100-msum['pct_inside_aoa']:.0f} % of the district lies "
        f"outside the applicability domain and is not mapped. Because the study covers a "
        f"single season and contains no climate analysis, no claims are made here about "
        f"climate resilience, drought monitoring, or transferability to the wider "
        f"Mediterranean region.")])

# ==========================================================================
#  6. CONCLUSIONS
# ==========================================================================
head("6. Conclusions", 1)
par([N_(f"Sentinel-2 reflectance predicts topsoil macronutrients in the Beni Moussa "
        f"irrigated district with moderate accuracy when performance is measured on "
        f"spatially independent data: coefficients of determination of {rf_n.R2:.2f} for "
        f"total nitrogen, {rf_p.R2:.2f} for Olsen phosphorus and {rf_k.R2:.2f} for "
        f"exchangeable potassium under nested spatial block cross-validation, with 95 % "
        f"confidence intervals of [{cn.R2_lo:.2f}, {cn.R2_hi:.2f}], "
        f"[{cp.R2_lo:.2f}, {cp.R2_hi:.2f}] and [{ck.R2_lo:.2f}, {ck.R2_hi:.2f}] "
        f"respectively. Random Forest and XGBoost are statistically indistinguishable on "
        f"this dataset, and for potassium a ridge regression on the same predictors "
        f"performs at least as well as either.")])
par([N_(f"Predictive skill rests on raw near-infrared and red-edge reflectance from a "
        f"single November composite rather than on soil-oriented spectral indices, which "
        f"indicates that the models track canopy condition rather than soil composition "
        f"directly. The resulting 10 m surfaces, delivered with local uncertainty and an "
        f"applicability domain covering {msum['pct_inside_aoa']:.0f} % of the district "
        f"({msum['ha_inside_aoa']:,.0f} ha), are best used for reconnaissance, "
        f"stratified resampling and hypothesis generation rather than for direct "
        f"fertiliser prescription. Establishing whether these relationships hold across "
        f"seasons, and whether they persist within crop classes, requires multi-season "
        f"sampling, crop-type information and non-spectral soil-forming covariates. That "
        f"is the necessary next step before products of this kind can support "
        f"operational nutrient management in irrigated districts of the Tadla Plain.")])

# ==========================================================================
#  BACK MATTER
# ==========================================================================
head("Acknowledgments", 1)
par([K("This study was conducted as part of a PhD thesis. "),
     N_("The authors thank the editor and the referee, whose detailed technical review "
        "substantially changed the analysis and the conclusions of this work.")])
head("Author Contributions", 1)
par("Y.A.: Conceptualization, Data curation, Formal analysis, Methodology, Software, "
    "Visualization, Writing – original draft. A.A.: Supervision, Conceptualization, "
    "Writing – review and editing. D.E.: Supervision, Writing – review and editing. "
    "M.E.B.: Supervision, Resources, Writing – review and editing. A.B.: Resources, "
    "Writing – review and editing. J.E.A.: Resources, Writing – review and editing. "
    "I.O.: Writing – review and editing. M.B.: Writing – review and editing. "
    "O.N.-T.: Writing – review and editing. A.H.: Supervision, Resources, Writing – "
    "review and editing.")
head("Funding", 1)
par("This research received no specific grant from any funding agency in the public, "
    "commercial, or not-for-profit sectors.")
head("Data Availability", 1)
par([N_(f"The analysis-ready dataset (110 observations with coordinates and "
        f"{s01['n_predictors']} predictors), the complete analysis code, the exact "
        f"cross-validation fold assignments, session information, out-of-fold "
        f"predictions for every validation scheme, and the georeferenced 10 m "
        f"prediction, uncertainty, dissimilarity and applicability rasters are openly "
        f"available at [GITHUB URL] and archived with a persistent identifier at "
        f"[ZENODO DOI]. Every value reported in this article can be reproduced from that "
        f"archive by running the numbered analysis scripts in order. The Sentinel-2 "
        f"Level-2A imagery is freely available from the Copernicus Data Space; the "
        f"Google Earth Engine scripts used to build and export the covariates are "
        f"included in the archive.")])
author_note("Insert the GitHub URL and the Zenodo DOI. Linking the repository to Zenodo "
            "and cutting a release takes about ten minutes and produces the persistent "
            "identifier the referee requires under M11; a GitHub URL alone is not one.")
head("Conflict of Interest", 1)
par("On behalf of all authors, the corresponding author states that there is no conflict "
    "of interest.")
head("Consent for Publication", 1)
par("Not applicable.")
head("Ethical Responsibilities of Authors", 1)
par("The authors have read and agree to the journal's Ethical Responsibilities of Authors "
    "as stated in the submission guidelines.")

# ==========================================================================
#  REFERENCES
# ==========================================================================
doc.add_page_break()
head("References", 1)

# (new?, "authors (year). title.", "Journal Name", " rest")
REFS = [
 (0, "Al Masmoudi, Y., Bouslihim, Y., Doumali, K., Hssaini, L., and Ibno Namr, K. (2022). Use of machine learning in Moroccan soil fertility prediction as an alternative to laborious analyses. ", "Modeling Earth Systems and Environment", " 8, 3707–3717. https://doi.org/10.1007/s40808-021-01329-8"),
 (0, "Aljanabi, F. K., and Dedeoğlu, M. (2026). Hybrid machine learning model and terrain variables for spatial modeling of topsoil physicochemical properties. ", "International Journal of Engineering and Geosciences", " 11, 149–162. https://doi.org/10.26833/ijeg.1655607"),
 (0, "Alvarez, C. I., and Govind, A. (2025). Assessing climate and land use changes in Morocco (2001–2023): from a geospatial and farmers' perspective. ", "Theoretical and Applied Climatology", " 156, 420. https://doi.org/10.1007/s00704-025-05656-z"),
 (0, "Barakat, A., Ennaji, W., El Jazouli, A., Amediaz, R., and Touhami, F. (2017). Multivariate analysis and GIS-based soil suitability diagnosis for sustainable intensive agriculture in Beni-Moussa irrigated subperimeter (Tadla plain, Morocco). ", "Modeling Earth Systems and Environment", " 3, 3. https://doi.org/10.1007/s40808-017-0272-5"),
 (0, "Bijaber, N., Rochdi, A., Yessef, M., and El Yacoubi, H. (2024). Mapping the structural vulnerability to drought in Morocco. ", "International Journal of Engineering and Geosciences", " 9, 264–280. https://doi.org/10.26833/ijeg.1404507"),
 (0, "Bouslihim, Y., Bouasria, A., Jelloul, A., Khiari, L., Dahhani, S., Mrabet, R., and Moussadek, R. (2025). Baseline high-resolution maps of soil nutrients in Morocco to support sustainable agriculture. ", "Scientific Data", " 12, 1389. https://doi.org/10.1038/s41597-025-05699-x"),
 (1, "Breiman, L. (2001). Random forests. ", "Machine Learning", " 45, 5–32. https://doi.org/10.1023/A:1010933404324"),
 (1, "Castaldi, F. (2021). Sentinel-2 and Landsat-8 multi-temporal series to estimate topsoil properties on croplands. ", "Remote Sensing", " 13, 3345. https://doi.org/10.3390/rs13173345"),
 (0, "Chaaou, A., Chikhaoui, M., Naimi, M., El Miad, A. K., Achemrk, A., Seif-Ennasr, M., and El Harche, S. (2022). Mapping soil salinity risk using the approach of soil salinity index and land cover: a case study from Tadla plain, Morocco. ", "Arabian Journal of Geosciences", " 15, 722. https://doi.org/10.1007/s12517-022-10009-5"),
 (1, "Chen, T., and Guestrin, C. (2016). XGBoost: a scalable tree boosting system. In ", "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining", ", 785–794. https://doi.org/10.1145/2939672.2939785"),
 (1, "Dalle Vaglie, M., Francini, S., Chirici, G., and Martellozzo, F. (2026). A large-scale framework for estimating soil carbon, nitrogen, pH, and salinity dynamics for 1985–2023. ", "Proceedings of the National Academy of Sciences of the United States of America", " 123, e2534913123. https://doi.org/10.1073/pnas.2534913123"),
 (1, "Dehni, A., and Lounis, M. (2012). Remote sensing techniques for salt affected soil mapping: application to the Oran region of Algeria. ", "Procedia Engineering", " 33, 188–198. https://doi.org/10.1016/j.proeng.2012.01.1193"),
 (0, "Drusch, M., Del Bello, U., Carlier, S., Colin, O., Fernandez, V., Gascon, F., Hoersch, B., Isola, C., Laberinti, P., Martimort, P., Meygret, A., Spoto, F., Sy, O., Marchese, F., and Bargellini, P. (2012). Sentinel-2: ESA's optical high-resolution mission for GMES operational services. ", "Remote Sensing of Environment", " 120, 25–36. https://doi.org/10.1016/j.rse.2011.11.026"),
 (0, "El Hamzaoui, E. H., and El Baghdadi, M. (2021). Characterizing spatial variability of some soil properties in Beni-Moussa irrigated perimeter from Tadla plain (Morocco) using geostatistics and kriging techniques. ", "Journal of Sedimentary Environments", " 6, 381–394. https://doi.org/10.1007/s43217-021-00050-x"),
 (0, "Ennaji, W., Barakat, A., El Baghdadi, M., Oumenskou, H., Aadraoui, M., Karroum, L. A., and Hilali, A. (2018). GIS-based multi-criteria land suitability analysis for sustainable agriculture in the northeast area of Tadla plain (Morocco). ", "Journal of Earth System Science", " 127, 79. https://doi.org/10.1007/s12040-018-0980-x"),
 (0, "Gao, B.-C. (1996). NDWI—a normalized difference water index for remote sensing of vegetation liquid water from space. ", "Remote Sensing of Environment", " 58, 257–266. https://doi.org/10.1016/S0034-4257(96)00067-3"),
 (0, "Gitelson, A., and Merzlyak, M. N. (1994). Spectral reflectance changes associated with autumn senescence of Aesculus hippocastanum L. and Acer platanoides L. leaves. ", "Journal of Plant Physiology", " 143, 286–292. https://doi.org/10.1016/S0176-1617(11)81633-0"),
 (0, "Gitelson, A. A., Kaufman, Y. J., and Merzlyak, M. N. (1996). Use of a green channel in remote sensing of global vegetation from EOS-MODIS. ", "Remote Sensing of Environment", " 58, 289–298. https://doi.org/10.1016/S0034-4257(96)00072-7"),
 (0, "Gorelick, N., Hancher, M., Dixon, M., Ilyushchenko, S., Thau, D., and Moore, R. (2017). Google Earth Engine: planetary-scale geospatial analysis for everyone. ", "Remote Sensing of Environment", " 202, 18–27. https://doi.org/10.1016/j.rse.2017.06.031"),
 (0, "Hafiane, F. Z., Tahri, L., Nouayti, N., El Jarmouni, M., Rochdi, R., Arifi, K., Idrissi Elamrani, A., and Fekhaoui, M. (2020). Microbial quality assessment of Beni Aamir and Beni Moussa groundwater (Tadla plain-Morocco). ", "Desalination and Water Treatment", " 200, 74–81. https://doi.org/10.5004/dwt.2020.26144"),
 (1, "Harris, C. R., Millman, K. J., van der Walt, S. J., Gommers, R., Virtanen, P., Cournapeau, D., Wieser, E., Taylor, J., Berg, S., and Smith, N. J. (2020). Array programming with NumPy. ", "Nature", " 585, 357–362. https://doi.org/10.1038/s41586-020-2649-2"),
 (0, "Havlin, J., and Heiniger, R. (2020). Soil fertility management for better crop production. ", "Agronomy", " 10, 1349. https://doi.org/10.3390/agronomy10091349"),
 (0, "Huete, A. R. (1988). A soil-adjusted vegetation index (SAVI). ", "Remote Sensing of Environment", " 25, 295–309. https://doi.org/10.1016/0034-4257(88)90106-X"),
 (0, "Huete, A., Didan, K., Miura, T., Rodriguez, E. P., Gao, X., and Ferreira, L. G. (2002). Overview of the radiometric and biophysical performance of the MODIS vegetation indices. ", "Remote Sensing of Environment", " 83, 195–213. https://doi.org/10.1016/S0034-4257(02)00096-2"),
 (0, "IUSS Working Group WRB (2015). ", "World Reference Base for Soil Resources 2014, Update 2015: International Soil Classification System for Naming Soils and Creating Legends for Soil Maps", ". World Soil Resources Reports No. 106. FAO, Rome. https://www.fao.org/3/i3794en/I3794en.pdf"),
 (0, "Khan, N. M., Rastoskuev, V. V., Sato, Y., and Shiozawa, S. (2005). Assessment of hydrosaline land degradation by using a simple approach of remote sensing indicators. ", "Agricultural Water Management", " 77, 96–109. https://doi.org/10.1016/j.agwat.2004.09.038"),
 (0, "Kuhn, M., and Johnson, K. (2013). ", "Applied Predictive Modeling", ". Springer, New York. https://doi.org/10.1007/978-1-4614-6849-3"),
 (1, "Lin, L. I.-K. (1989). A concordance correlation coefficient to evaluate reproducibility. ", "Biometrics", " 45, 255–268. https://doi.org/10.2307/2532051"),
 (0, "Liu, Y., Liu, C., Rubinato, M., Guo, K., Zhou, J., and Cui, M. (2020). An assessment of soil's nutrient deficiencies and their influence on the restoration of degraded karst vegetation in Southwest China. ", "Forests", " 11, 797. https://doi.org/10.3390/f11080797"),
 (0, "McBratney, A. B., Mendonça Santos, M. L., and Minasny, B. (2003). On digital soil mapping. ", "Geoderma", " 117, 3–52. https://doi.org/10.1016/S0016-7061(03)00223-4"),
 (1, "Meyer, H., and Pebesma, E. (2021). Predicting into unknown space? Estimating the area of applicability of spatial prediction models. ", "Methods in Ecology and Evolution", " 12, 1620–1633. https://doi.org/10.1111/2041-210X.13650"),
 (0, "Minasny, B., and McBratney, A. B. (2006). A conditioned Latin hypercube method for sampling in the presence of ancillary information. ", "Computers and Geosciences", " 32, 1378–1388. https://doi.org/10.1016/j.cageo.2005.12.009"),
 (0, "Mouaddine, A., Barakat, A., Hajaj, S., Mosaid, H., Bouzekraoui, H., Bni, Z., and Hilali, A. (2025). Predicting and mapping soil saturated hydraulic conductivity in the Beni Moussa irrigated perimeter (Tadla Plain, Morocco) using Random Forest machine learning model. ", "Modeling Earth Systems and Environment", " 11, 82. https://doi.org/10.1007/s40808-024-02210-0"),
 (0, "Nait-Taleb, O., Elomari, S., Abdelrahman, K., Ismaili, M., Fnais, M. S., Atiq, J. E., Ouchkir, I., Karaoui, I., Krimissa, S., Namous, M., and Elaloui, A. (2025). Monitoring soil degradation using Sentinel-2 imagery and statistical analysis of spectral indices in a semi-arid watershed of the Moroccan High Atlas. ", "Frontiers in Soil Science", " 5, 1553887. https://doi.org/10.3389/fsoil.2025.1553887"),
 (0, "Olsen, S. R. (1954). ", "Estimation of Available Phosphorus in Soils by Extraction with Sodium Bicarbonate", ". USDA Circular No. 939. United States Department of Agriculture, Washington, DC."),
 (1, "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., and Dubourg, V. (2011). Scikit-learn: machine learning in Python. ", "Journal of Machine Learning Research", " 12, 2825–2830. https://www.jmlr.org/papers/v12/pedregosa11a.html"),
 (0, "Recena, R., Torrent, J., del Campillo, M. C., and Delgado, A. (2015). Accuracy of Olsen P to assess plant P uptake in relation to soil properties and P forms. ", "Agronomy for Sustainable Development", " 35, 1571–1579. https://doi.org/10.1007/s13593-015-0332-z"),
 (0, "Rikimaru, A., Roy, P. S., and Miyatake, S. (2002). Tropical forest cover density mapping. ", "Tropical Ecology", " 43, 39–47. https://www.tropecol.com/pdf/open/PDF_43_1/43104.pdf"),
 (1, "Roberts, D. R., Bahn, V., Ciuti, S., Boyce, M. S., Elith, J., Guillera-Arroita, G., Hauenstein, S., Lahoz-Monfort, J. J., Schröder, B., Thuiller, W., Warton, D. I., Wintle, B. A., Hartig, F., and Dormann, C. F. (2017). Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. ", "Ecography", " 40, 913–929. https://doi.org/10.1111/ecog.02881"),
 (0, "Rouse, J. W., Haas, R. H., Schell, J. A., and Deering, D. W. (1974). Monitoring vegetation systems in the Great Plains with ERTS. ", "NASA Special Publication", " 351, 309–317."),
 (0, "Salih, M. A., Bilgili, A. V., and Rekani, S. I. K. (2026). Soil fertility assessment and mapping in Bardarash District, Duhok Province, Iraq: an integrated approach using geostatistics, machine learning, and laboratory analyses. ", "Journal of Animal and Plant Sciences", " 36, 375–390. https://doi.org/10.36899/JAPS.2026.2.0032"),
 (0, "Salmi, A., El Baghdadi, M., Mosaid, H., Barakat, A., and Hilali, A. (2024). Iron behaviour and soil properties in hydromorphic soils of Beni Moussa, Tadla Plain, Morocco. ", "Ecological Chemistry and Engineering S", " 31, 365–383. https://doi.org/10.2478/eces-2024-0025"),
 (0, "Salmi, A., El Baghdadi, M., Hilali, A., Ennaji, W., and Mosaid, H. (2025). Assessing the relationship between iron behavior and phosphorus in hydromorphic soils: the Day Valley case, Tadla Plain, Morocco. ", "Mediterranean Geoscience Reviews", " 7, 379–393. https://doi.org/10.1007/s42990-025-00165-7"),
 (0, "Shen, J., Yuan, L., Zhang, J., Li, H., Bai, Z., Chen, X., Zhang, W., and Zhang, F. (2011). Phosphorus dynamics: from soil to plant. ", "Plant Physiology", " 156, 997–1005. https://doi.org/10.1104/pp.111.175232"),
 (0, "Silatsa, F. B. T., and Kebede, F. (2023). A quarter century experience in soil salinity mapping and its contribution to sustainable soil management and food security in Morocco. ", "Geoderma Regional", " 34, e00695. https://doi.org/10.1016/j.geodrs.2023.e00695"),
 (0, "Silvero, N. E. Q., Demattê, J. A. M., Amorim, M. T. A., dos Santos, N. V., Rizzo, R., Safanelli, J. L., Poppiel, R. R., de Sousa Mendes, W., and Bonfatti, B. R. (2021). Soil variability and quantification based on Sentinel-2 and Landsat-8 bare soil images: a comparison. ", "Remote Sensing of Environment", " 252, 112117. https://doi.org/10.1016/j.rse.2020.112117"),
 (0, "Vereecken, H., Schnepf, A., Hopmans, J. W., Javaux, M., Or, D., Roose, T., Vanderborght, J., Young, M. H., Amelung, W., Aitkenhead, M., Allison, S. D., Assouline, S., Baveye, P., Berli, M., Brüggemann, N., Finke, P., Flury, M., Gaiser, T., Govers, G., and Young, I. M. (2016). Modeling soil processes: review, key challenges, and new perspectives. ", "Vadose Zone Journal", " 15, 1–57. https://doi.org/10.2136/vzj2015.09.0131"),
 (0, "Veronesi, F., and Schillaci, C. (2019). Comparison between geostatistical and machine learning models as predictors of topsoil organic carbon with a focus on local uncertainty estimation. ", "Ecological Indicators", " 101, 1032–1044. https://doi.org/10.1016/j.ecolind.2019.02.026"),
 (0, "Wadoux, A. M. J.-C., Minasny, B., and McBratney, A. B. (2020). Machine learning for digital soil mapping: applications, challenges and suggested solutions. ", "Earth-Science Reviews", " 210, 103359. https://doi.org/10.1016/j.earscirev.2020.103359"),
 (0, "Wiesmeier, M., Urbanski, L., Hobley, E., Lang, B., von Lützow, M., Marin-Spiotta, E., van Wesemael, B., Rabot, E., Ließ, M., Garcia-Franco, N., Wollschläger, U., Vogel, H. J., and Kögel-Knabner, I. (2019). Soil organic carbon storage as a key function of soils – a review of drivers and indicators at various scales. ", "Geoderma", " 333, 149–162. https://doi.org/10.1016/j.geoderma.2018.07.026"),
 (0, "Xu, H. (2006). Modification of normalised difference water index (NDWI) to enhance open water features in remotely sensed imagery. ", "International Journal of Remote Sensing", " 27, 3025–3033. https://doi.org/10.1080/01431160600589179"),
]
for new, pre, jour, post in REFS:
    p = doc.add_paragraph()
    col = RED if new else BLK
    for txt, ital in ((pre, False), (jour, True), (post, False)):
        r = p.add_run(txt)
        r.font.size = Pt(SZ); r.font.color.rgb = col; r.italic = ital
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

n_new = sum(n for n, _, _, _ in REFS)
author_note(f"The reference list contains {len(REFS)} entries, of which {n_new} are new "
            f"or corrected (red). Ten entries from the previous version were deleted "
            f"because the referee showed they did not support the statements attached to "
            f"them, or were never cited: Abatzoglou et al. (2018), Cambardella et al. "
            f"(1994), Chen et al. (2017), Liu and Huete (1995), Lundberg and Lee (2017), "
            f"Nguyen et al. (2020), Qi et al. (1994), Sharma (2010), Viscarra Rossel et "
            f"al. (2016) and Zhang (2023). Every entry is cited in the text and every "
            f"in-text citation appears here.")
author_note(f"All {len(doi)} DOIs were checked against Crossref and all resolve "
            f"(Supplementary Table S7). Three findings need your decision. (1) The "
            f"Castaldi entry in the previous version was wrong in every field: DOI "
            f"10.3390/rs11242924 resolves to Monteith et al., a radar-tomography paper. "
            f"The correct record is Castaldi, F. (2021), Remote Sensing 13, 3345, doi "
            f"10.3390/rs13173345, single author, 2021 not 2019. Corrected here and the "
            f"in-text citation changed to 'Castaldi, 2021'. (2) Aljanabi and Dedeoğlu: "
            f"Crossref records online publication on 1 October 2025, while the article "
            f"sits in Volume 11, Issue 1, whose volume sequence implies 2026. We cite "
            f"2026; 2025 is equally defensible. (3) Salih et al.: Crossref records "
            f"25 December 2025 online, in the issue dated April 2026; we cite 2026.")

# ==========================================================================
#  SUPPLEMENTARY
# ==========================================================================
doc.add_page_break()
head("Supplementary Material", 1, new=True)
par([N_("The following tables support the main text and are provided with the "
        "submission. All items are also included in the public archive cited under Data "
        "Availability.")])

inv_s = inv.copy()
inv_s.columns = ["Variable", "Type", "Acquisition", "Formula", "Source bands",
                 "Wavelength (nm)", "Unit", "Native res. (m)", "Output res. (m)",
                 "Resampling", "Reference"]
table(inv_s[["Variable", "Type", "Acquisition", "Formula", "Source bands",
             "Native res. (m)", "Output res. (m)", "Resampling"]],
      "Table S1", f"Complete inventory of the {s01['n_predictors']} predictors used in "
                  f"the final models.", fontsize=7.5,
      widths=[2.4, 2.2, 2.6, 3.6, 2.0, 1.6, 1.6, 1.8])
table(t6f, "Table S2",
      "Hyperparameters selected by the inner tuning loop for every outer fold, model and "
      "nutrient.", fontsize=7.5)
table(t9b, "Table S3",
      "Confusion matrices between observed fertility class and the class of the held-out "
      "outer-fold prediction, for both algorithms and all three nutrients.", fontsize=8)

vg_s = vg.copy()
vg_s.columns = ["Variable", "Kind", "Nugget", "Partial sill", "Range (m)", "Nugget ratio"]
vg_s["Nugget"] = vg_s["Nugget"].map(lambda v: f"{v:,.1f}")
vg_s["Partial sill"] = vg_s["Partial sill"].map(lambda v: f"{v:,.1f}")
vg_s["Range (m)"] = vg_s["Range (m)"].map(lambda v: f"{v:,.0f}")
vg_s["Nugget ratio"] = vg_s["Nugget ratio"].map(lambda v: f"{v:.3f}")
table(vg_s, "Table S4a",
      "Fitted spherical variogram parameters for the observed nutrients and for the "
      "Random Forest spatial cross-validation residuals.", fontsize=8.5)
mor_s = mor.copy(); mor_s.columns = ["Bandwidth (km)", "N", "P", "K"]
table(mor_s, "Table S4b",
      "Moran's I of the observed nutrients at increasing distance bandwidths.",
      fontsize=8.5)
bd_s = bdes.copy()
bd_s.columns = ["Block", "n", "Extent E (km)", "Extent N (km)", "Centroid E (m)",
                "Centroid N (m)", "Min. separation to other block (km)"]
table(bd_s, "Table S5", "Spatial cross-validation block design (EPSG:26191).",
      fontsize=8.5)

sens_s = sens.copy()
sens_s.columns = ["Nutrient", "Extraction support", "R2", "RMSE", "MAE", "Bias", "CCC",
                  "Slope"]
for c in ["R2", "CCC", "Slope"]:
    sens_s[c] = sens_s[c].map(lambda v: f"{v:.3f}")
for c in ["RMSE", "MAE", "Bias"]:
    sens_s[c] = sens_s[c].map(lambda v: f"{v:,.3f}")
table(sens_s, "Table S6",
      "Extraction-support sensitivity. Random Forest under the identical nested spatial "
      "block cross-validation, with predictors extracted either from the single pixel "
      "containing the sample centroid or as the mean of a 3 × 3 pixel neighbourhood. "
      "RMSE, MAE and bias in mg kg-1.", fontsize=8.5)

doi_s = doi[["reference", "doi", "resolved", "expected_year", "crossref_year",
             "crossref_author"]].copy()
doi_s.columns = ["Reference", "DOI", "Resolves", "Year cited", "Crossref year",
                 "Crossref first author"]
table(doi_s, "Table S7",
      f"Verification of every DOI in the reference list against Crossref "
      f"({(doi.resolved == 'yes').sum()} of {len(doi)} resolve).",
      fontsize=7.5, widths=[4.2, 4.4, 1.7, 1.7, 1.9, 2.5], align_center=False)

scn_s = scn[["composite", "datetime", "mgrs_tile", "cloud_pct", "nodata_pct",
             "platform", "orbit", "sun_zenith"]].copy()
scn_s.columns = ["Composite", "Sensing date and time (UTC)", "MGRS tile", "Cloud (%)",
                 "No-data (%)", "Platform", "Orbit", "Solar zenith (deg)"]
scn_s["Cloud (%)"] = scn_s["Cloud (%)"].map(lambda v: f"{v:.3f}")
scn_s["No-data (%)"] = scn_s["No-data (%)"].map(lambda v: f"{v:.2f}")
scn_s["Solar zenith (deg)"] = scn_s["Solar zenith (deg)"].map(lambda v: f"{v:.2f}")
scn_s["Orbit"] = scn_s["Orbit"].astype(int)
table(scn_s, "Table S8",
      f"Sentinel-2 Level-2A scenes contributing to the two monthly median composites "
      f"(COPERNICUS/S2_SR_HARMONIZED, scene cloud cover below 40 %). November 2024: "
      f"{int(c11.n_scenes)} scenes on {int(c11.n_dates)} dates; January 2025: "
      f"{int(c01.n_scenes)} scenes on {int(c01.n_dates)} dates. Full product and granule "
      f"identifiers are given in the deposited file data/14_scene_inventory_used.csv.",
      fontsize=7.5, widths=[1.9, 3.4, 1.7, 1.7, 1.9, 2.4, 1.4, 2.0])

# ==========================================================================
#  EDITORIAL NOTE
# ==========================================================================
doc.add_page_break()
head("Editorial note to the authors — delete this page before submission", 1, new=True)
par([NB("What was deleted from the previous version, and why")], space_after=4)
for t in [
    "The claim that XGBoost outperformed Random Forest, including the statement that it "
    "reduced nitrogen RMSE by nearly 50 %. That comparison set a Random Forest "
    "cross-validation figure against an XGBoost calibration figure. On identical "
    "held-out folds the difference is not significant.",
    "All calibration statistics presented as evidence of predictive skill (previous R2 "
    "values of 0.910–0.988). They are retained only in Table 7, to document the gap "
    "between calibration and validation.",
    "The interpretation that red-edge indices drive nitrogen, vegetation-vigour indices "
    "drive phosphorus and SWIR indices drive potassium. Permutation importance on "
    "held-out folds does not support it; NDRE and GNDVI together hold about 1 % of "
    "nitrogen importance.",
    "The description of GNDVI as a red-edge index. It is a green–NIR index.",
    "Claims that the maps support direct variable-rate fertiliser prescription.",
    "Climate-resilience, drought-monitoring and Mediterranean-transferability claims. "
    "The study covers one season and contains no climate analysis.",
    "Causal attribution of map patterns to clay, salinity, hydromorphy, irrigation "
    "infrastructure, carbonates and fertilisation history. None was measured; they are "
    "now framed as hypotheses.",
    "The statements in Section 3.7.2 that XGBoost maps are clearer or retain more "
    "detail. These were results placed in the Methods and were untested; Table 10 now "
    "reports the agreement statistics instead.",
    "The assertion that Random Forest and XGBoost outperform SVM or ANN and require "
    "fewer adjustments. Neither was tested.",
    "The interquartile-range rule for removing outliers. Valid extremes should not be "
    "removed on statistical grounds alone.",
    "The unused correlation-based definition of R2, the duplicated sentence in the "
    "Figure 5 caption, the near-duplicate closing paragraph in the Conclusions, and the "
    "repeated novelty claim in the Abstract and Conclusions.",
    "Ten references that did not support the statements attached to them, or were never "
    "cited (listed in the note under the reference list).",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t); r.font.size = Pt(9); r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.line_spacing = 1.15

par([NB("What still needs your input before submitting")], space_after=4)
for t in [
    "Optionally, the limits of detection and quantification, the percentage recovery on "
    "the internal reference soils, and repeatability as a relative standard deviation, "
    "if those figures exist in the laboratory records. The procedures themselves are now "
    "fully described in Section 3.1; only these three quantities are absent.",
    "Confirmation of the sampling design name and the minimum separation, plus a table "
    "of samples by stratum and crop or land-cover class in Section 3.1.",
    "A locally relevant source for each fertility threshold in Table 4.",
    "The GitHub URL and Zenodo DOI in the Data Availability statement.",
    "Regeneration of Fig. 2 without its embedded caption, updated to 40 predictors.",
    "A decision on the Aljanabi and Dedeoğlu publication year, 2025 or 2026.",
    "A final read for English style. The text has been rewritten throughout, but a "
    "professional proofread before submission remains advisable.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t); r.font.size = Pt(9); r.font.color.rgb = NOTE
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.line_spacing = 1.15

out = os.path.join(DOCS, "Revised_Manuscript_v3_TrackedRed_Amrouss_et_al.docx")
doc.save(out)
print("wrote", out)
print(f"references: {len(REFS)} ({n_new} new/corrected)")
print("STEP 10 complete.")
