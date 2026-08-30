# -*- coding: utf-8 -*-
"""
STEP 13 - "What data you still need, and exactly what it should look like".

For each remaining gap: what the referee asked, what to collect, a worked EXAMPLE
with invented numbers, and the ready-to-paste sentence with blanks.
"""
import os, json
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FINAL = r"d:\Doctorat\article1\outputs_fast\FINAL"
RES   = os.path.join(FINAL, "results")
DOCS  = os.path.join(FINAL, "docs")
os.makedirs(DOCS, exist_ok=True)

ACC  = RGBColor(0x1F, 0x5F, 0x6B)
RED  = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x55, 0x5F, 0x5B)
GRN  = RGBColor(0x1B, 0x6E, 0x3C)
ORG  = RGBColor(0xB5, 0x52, 0x2F)
FONT = "Calibri"

msum = json.load(open(os.path.join(RES, "03_map_summary.json")))

doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = s.bottom_margin = Cm(1.8)
st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12
for lvl, sz in [(1, 15), (2, 12), (3, 10.5)]:
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = FONT; h.font.size = Pt(sz); h.font.bold = True
    h.font.color.rgb = ACC
    h.paragraph_format.space_before = Pt(14); h.paragraph_format.space_after = Pt(4)


def P(text="", bold=False, italic=False, size=10.5, color=None, align=None,
      space_after=6, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    r.font.name = FONT
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), hexcolor)
    tcPr.append(s)


def boxed(label, text, colour, fill):
    p = doc.add_paragraph()
    r = p.add_run(label + "  "); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = colour; r.font.name = FONT
    r2 = p.add_run(text); r2.font.size = Pt(9.8); r2.font.name = FONT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    lf = OxmlElement("w:left")
    lf.set(qn("w:val"), "single"); lf.set(qn("w:sz"), "18")
    lf.set(qn("w:space"), "8")
    lf.set(qn("w:color"), f"{colour.__str__()}")
    pbdr.append(lf); pPr.append(pbdr)
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill)
    pPr.append(sh)
    return p


def ask(t):    boxed("REFEREE ASKS", t, ACC, "EEF3F4")
def collect(t): boxed("WHAT TO COLLECT", t, ORG, "FBF1EC")


def example(t):
    p = doc.add_paragraph()
    r = p.add_run("EXAMPLE (invented numbers — replace with yours)  ")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GRN; r.font.name = FONT
    r2 = p.add_run(t); r2.font.size = Pt(9.8); r2.font.name = FONT
    r2.font.color.rgb = GRN
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "EDF6EF")
    pPr.append(sh)
    return p


def table(dfx, caption=None, fontsize=8.4, widths=None, green=False):
    cols = list(dfx.columns)
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]; cell.text = ""
        r = cell.paragraphs[0].add_run(str(c)); r.bold = True
        r.font.size = Pt(fontsize); r.font.name = FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        shade(cell, "1B6E3C" if green else "1F5F6B")
    for _, row in dfx.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run("" if pd.isna(v) else str(v))
            r.font.size = Pt(fontsize); r.font.name = FONT
            if green:
                r.font.color.rgb = GRN
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
    for ri, row in enumerate(t.rows):
        if ri > 0 and ri % 2 == 0:
            for c in row.cells:
                shade(c, "EDF6EF" if green else "F2F5F1")
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    if caption:
        P(caption, size=8.6, italic=True, color=GREY, space_after=10)
    else:
        P("", size=4, space_after=6)
    return t


def code(lines):
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.2)
        r.font.color.rgb = RGBColor(0x1A, 0x21, 0x1E)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        pPr = p._p.get_or_add_pPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "F4F6F3")
        pPr.append(sh)
    P("", size=4, space_after=8)


# ==========================================================================
P("Data you still need to supply", bold=True, size=20, color=ACC,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P("with a worked example of each, ready to paste into the manuscript",
  size=11.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

P("Six items remain before the manuscript can be submitted. Three of the nine listed "
  "earlier are now closed: the Sentinel-2 scene inventory has been recovered in full "
  "(54 scenes, Supplementary Table S8); the 3 × 3 extraction-support sensitivity test "
  "has been run (Section 3.3, Supplementary Table S6); and all 46 DOIs have been "
  "verified against Crossref (Supplementary Table S7).")
P("For each remaining item below you will find what the referee asked for, exactly what "
  "to collect, a worked example in green with invented numbers, and a sentence you can "
  "paste into the manuscript once you replace the placeholders. The green numbers are "
  "illustrations of format only — they are not your data and must not be submitted.",
  space_after=4)
P("Do not invent any of these values. If a measurement was not made, say so plainly — "
  "\"blanks were not run\" is a weaker paper but an honest one, and a referee who has "
  "already caught three arithmetic problems will check.", bold=True, color=RED)

status = pd.DataFrame([
    ["1", "Laboratory QA/QC (§3.1)", "You — laboratory records", "Highest"],
    ["2", "Sampling design + stratum table (§3.1)", "You — field records", "High"],
    ["3", "GitHub URL + Zenodo DOI", "You — 10 min", "High"],
    ["4", "Fertility threshold sources (Table 3)", "You — local guideline", "Medium"],
    ["5", "Figure 2 regenerated", "You — 30 min", "Medium"],
    ["6", "Aljanabi year decision", "You — one decision", "Low"],
    ["—", "Sentinel-2 scene metadata", "DONE — Table S8, 54 scenes", "closed"],
    ["—", "3 × 3 sensitivity test", "DONE — Table S6", "closed"],
    ["—", "DOI verification (46 refs)", "DONE — Table S7", "closed"],
    ["—", "English rewrite", "DONE — full rewrite", "closed"],
], columns=["#", "Item", "Who / effort", "Priority"])
table(status, widths=[1.2, 7.4, 5.4, 2.6])

P("One further correction came out of your Earth Engine output, and it changes the "
  "Methods rather than adding to your to-do list.", bold=True, color=RED, space_after=3)
P("Your export script builds monthly median composites, not single scenes. The "
  "manuscript said \"two cloud-free images\". In fact the November 2024 composite is the "
  "per-pixel median of 30 scenes acquired on 11 dates, and the January 2025 composite "
  "the median of 24 scenes on 9 dates, across four MGRS tiles (29SPR, 29SPS, 29SQR, "
  "29SQS) and two relative orbits. The collection is COPERNICUS/S2_SR_HARMONIZED, not "
  "COPERNICUS/S2_SR. Section 3.2.1 has been rewritten accordingly and the full scene "
  "inventory added as Supplementary Table S8. Nothing you need to do — but read the new "
  "§3.2.1 before submitting, because it now describes your actual pipeline.",
  color=RED, space_after=12)

# ---------------------------------------------------------------- 1
doc.add_heading("1  Laboratory QA/QC — the biggest gap", 1)
ask("\"Provide the number of analytical replicates, blanks, certified or internal "
    "reference materials, calibration procedures, detection/quantification limits, "
    "recovery, repeatability/precision, and any rejected or incomplete measurements.\" "
    "(comment M4)")
collect("From your laboratory notebooks or the INRA Tadla laboratory: how many samples "
        "were run in duplicate; how many reagent blanks per batch; what reference or "
        "control material was used and what recovery it gave; how the calibration curve "
        "was built and its r²; the detection and quantification limits; the relative "
        "standard deviation between duplicates; and how many samples had to be re-run "
        "or were discarded.")

P("Suggested table to add as Table 3 in the manuscript", bold=True, space_after=3)
qa = pd.DataFrame([
    ["Method", "Kjeldahl digestion", "Olsen (0.5 M NaHCO₃, pH 8.5)",
     "1 N NH₄OAc, pH 7"],
    ["Determination", "Titration after distillation", "Molybdenum blue, 882 nm",
     "Flame photometry"],
    ["Samples analysed", "110", "110", "110"],
    ["Analytical duplicates", "12 (11 %)", "12 (11 %)", "12 (11 %)"],
    ["Reagent blanks", "2 per batch of 20", "2 per batch of 20", "2 per batch of 20"],
    ["Reference material", "Internal control soil IC-01", "Internal control soil IC-01",
     "Internal control soil IC-01"],
    ["Recovery (%)", "97.4 ± 3.1", "101.2 ± 4.6", "98.8 ± 2.9"],
    ["Calibration", "0.5–5.0 mg N, r² = 0.999", "0–2.0 mg L⁻¹ P, r² = 0.998",
     "0–40 mg L⁻¹ K, r² = 0.999"],
    ["LOD (mg kg⁻¹)", "35", "0.6", "1.8"],
    ["LOQ (mg kg⁻¹)", "105", "1.9", "5.5"],
    ["Repeatability (RSD %)", "4.2", "6.8", "3.5"],
    ["Samples re-run", "3", "5", "2"],
    ["Samples excluded", "0", "0", "0"],
], columns=["Quality control item", "Total N", "Available P", "Exchangeable K"])
table(qa, "Example only — replace every value with your laboratory's records.",
      widths=[4.4, 4.2, 4.2, 4.0], green=True)

P("Sentence to paste into §3.1 after the methods description", bold=True, space_after=3)
example("Analytical quality control followed the laboratory's standard protocol. "
        "Twelve samples (11 %) were analysed in duplicate for each determination, and "
        "two reagent blanks were included in every batch of twenty samples. An internal "
        "control soil of known composition was included in each batch; mean recovery was "
        "97.4 ± 3.1 % for total N, 101.2 ± 4.6 % for available P and 98.8 ± 2.9 % for "
        "exchangeable K. Calibration curves were prepared over the ranges given in "
        "Table 3 and all had r² ≥ 0.998. Limits of detection and quantification, "
        "estimated as three and ten times the standard deviation of the blanks, are "
        "reported in Table 3. Repeatability, expressed as the relative standard "
        "deviation between duplicates, was 4.2 % for total N, 6.8 % for available P and "
        "3.5 % for exchangeable K. Ten samples required re-analysis because of an "
        "out-of-range reading; none was excluded, and all 110 observations passed "
        "quality control.")
P("If some of this was not recorded, write instead: \"Duplicates and reagent blanks were "
  "included in each batch, but recovery against a certified reference material was not "
  "determined; this is a limitation of the present dataset.\" That is a defensible "
  "answer. An invented recovery figure is not.", italic=True, color=RED, space_after=12)

# ---------------------------------------------------------------- 2
doc.add_heading("2  Sampling design and stratum allocation", 1)
ask("\"Name the design, give allocation by stratum, describe randomization, define "
    "minimum separation, and state inclusion/exclusion criteria … Add a sampling-design "
    "map and a table of samples by stratum and crop/land-cover class.\" (comment M3)")
collect("How you actually chose the 110 locations. If you used conditioned Latin "
        "hypercube sampling, list the conditioning covariates and the software. If you "
        "used a stratified or grid-based design, say so — the map in Fig. 3a shows a "
        "partly regular arrangement, which the referee noticed, so an honest description "
        "of a stratified-systematic design will be more credible than a cLHS claim the "
        "layout does not support. Also record the crop present at each point.")
P("Suggested table to add as a new table in the manuscript", bold=True, space_after=3)
strat = pd.DataFrame([
    ["Cereals (wheat, barley)", "38", "34.5", "26,900", "33.8"],
    ["Forage (alfalfa, maize)", "27", "24.5", "19,400", "24.4"],
    ["Sugar beet", "18", "16.4", "13,100", "16.5"],
    ["Olive and citrus orchards", "15", "13.6", "11,600", "14.6"],
    ["Market garden", "8", "7.3", "5,200", "6.5"],
    ["Fallow / bare", "4", "3.6", "3,300", "4.2"],
    ["Total", "110", "100.0", f"{msum['mapped_area_ha']:,.0f}", "100.0"],
], columns=["Crop / land-cover stratum", "Samples (n)", "Samples (%)",
            "Stratum area (ha)", "Area (%)"])
table(strat, "Example only — replace with your field records. The area column can come "
             "from the national land-cover layer or from your own classification.",
      widths=[6.0, 2.6, 2.6, 3.0, 2.4], green=True)

P("Sentence to paste into §3.1", bold=True, space_after=3)
example("Sampling locations were allocated by a stratified-systematic design. The "
        "district was first divided into six strata defined by dominant crop and land "
        "cover (Table 2), and within each stratum points were placed on a systematic "
        "grid with a random origin, giving an allocation approximately proportional to "
        "stratum area. A minimum separation of 500 m was imposed between locations. "
        "Water bodies, settlements, roads and canal corridors were excluded from the "
        "sampling frame. Crop type at the time of sampling was recorded at every point "
        "and is reported in Table 2. The realised configuration is shown in Fig. 3a.")
P("If it really was cLHS, replace with: \"Sampling locations were allocated by "
  "conditioned Latin hypercube sampling (Minasny and McBratney, 2006) conditioned on "
  "[list the exact covariate layers], implemented in [software], with a minimum "
  "separation of [x] m.\" If it was not cLHS, delete the Minasny and McBratney "
  "citation — the referee checked it specifically.", italic=True, color=RED,
  space_after=12)

# ---------------------------------------------------------------- 4
doc.add_heading("3  GitHub repository and Zenodo DOI", 1)
ask("\"Deposit the analysis-ready data and scripts in a stable repository with a "
    "persistent identifier.\" (comment M11)")
collect("A GitHub URL is not a persistent identifier — repositories can be renamed or "
        "deleted. Link the repository to Zenodo and cut a release; Zenodo then mints a "
        "DOI that will resolve permanently. This takes about ten minutes.")
P("Steps", bold=True, space_after=3)
for i, s in enumerate([
    "Create a public GitHub repository, e.g. beni-moussa-npk-mapping.",
    "Upload the contents of outputs_fast\\FINAL\\deposit. The eleven GeoTIFFs total "
    "271 MB and exceed GitHub's 100 MB per-file limit — either enable Git LFS for "
    "maps/*.tif, or upload the rasters to Zenodo directly and keep code, tables and "
    "figures on GitHub.",
    "Sign in to zenodo.org with your GitHub account, open Settings → GitHub, and switch "
    "the repository toggle on.",
    "Back on GitHub, create a release tagged v1.0. Zenodo captures it automatically and "
    "issues a DOI within a minute.",
    "Paste the DOI into the Data availability statement, replacing [ZENODO DOI].",
], 1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(s); r.font.size = Pt(10); r.font.name = FONT
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.space_after = Pt(3)

P("Sentence to paste into Data availability", bold=True, space_after=3)
example("The analysis-ready dataset, the complete analysis code, the cross-validation "
        "fold assignments, session information, out-of-fold predictions and the "
        "georeferenced 10 m prediction, uncertainty, dissimilarity and applicability "
        "rasters are openly available at https://github.com/yamrouss/"
        "beni-moussa-npk-mapping and archived at https://doi.org/10.5281/zenodo.14832991. "
        "Every value reported in this article can be reproduced from that archive by "
        "running the numbered analysis scripts in order.")

# ---------------------------------------------------------------- 5
doc.add_heading("4  A source for each fertility threshold", 1)
ask("\"Provide one verifiable, locally relevant source for each exact threshold, show "
    "the conversion used for total N, and write non-overlapping boundaries … Explain and "
    "validate any local adaptation rather than choosing thresholds merely to avoid empty "
    "classes.\" (comment M10)")
collect("A Moroccan or regional fertility interpretation guide — an INRA Morocco "
        "extension bulletin, the ORMVAT fertilisation guide for the Tadla perimeter, or "
        "a national soil-survey interpretation manual. One citation per nutrient is "
        "enough. The non-overlapping boundaries and the N conversion are already done in "
        "Table 3 of the manuscript.")
P("Sentence to paste into §3.7.3", bold=True, space_after=3)
example("Class boundaries follow the interpretation ranges used by the national "
        "agricultural extension service for irrigated soils of the Tadla perimeter "
        "(ORMVAT, 2019) for phosphorus and potassium, and the total-nitrogen ranges of "
        "Havlin and Heiniger (2020) converted to mg kg⁻¹ using 1,000 mg kg⁻¹ = 0.10 % N. "
        "No threshold was adjusted to alter the resulting class distribution.")
P("If no local guideline exists in print, say so and cite the international source you "
  "did use: \"In the absence of a published local interpretation scheme for the Tadla "
  "perimeter, internationally used ranges were adopted (Havlin and Heiniger, 2020; "
  "Recena et al., 2015) and applied without local adjustment.\" That is honest and the "
  "referee will accept it.", italic=True, color=RED, space_after=12)

# ---------------------------------------------------------------- 6
doc.add_heading("5  Figure 2 regenerated", 1)
ask("\"Remove the caption embedded inside Figure 2 if a document caption is retained, "
    "and ensure the two versions do not carry different wording.\" (minor comments)")
collect("Redraw the workflow diagram in PowerPoint, Inkscape or draw.io and export at "
        "600 dpi or as vector PDF. Remove the title text baked into the image, and "
        "update the content so it matches the revised analysis.")
P("What the boxes should now say", bold=True, space_after=3)
flow = pd.DataFrame([
    ["1", "Field sampling", "110 composite topsoil samples, 0–20 cm, Nov 2024 – Jan 2025"],
    ["2", "Laboratory", "Kjeldahl N · Olsen P · NH₄OAc K, with QA/QC"],
    ["3", "Sentinel-2 (GEE)", "2 L2A scenes · SCL masking · ÷10,000 scaling · 20 m → 10 m"],
    ["4", "Predictors", "10 bands + 10 indices × 2 dates = 40 predictors"],
    ["5", "Extraction", "Predictors at sample locations (3 × 3 sensitivity tested)"],
    ["6", "Validation", "Nested CV: 10 spatial blocks outer, GroupKFold(5) inner"],
    ["7", "Models", "Random Forest · XGBoost · 5 benchmarks on identical folds"],
    ["8", "Prediction", "10 m surfaces · tree-ensemble uncertainty · area of applicability"],
    ["9", "Products", "Continuous maps · fertility classes · class areas inside the AOA"],
], columns=["Step", "Box", "Content"])
table(flow, "Replace '41 predictors', any 'R version' box and any 10-fold random CV box "
            "with these.", widths=[1.4, 4.0, 11.0], green=True)

# ---------------------------------------------------------------- 7
doc.add_heading("6  One decision on a publication year", 1)
ask("\"The Aljanabi and Dedeoğlu item also requires year verification, because the "
    "journal landing page and the issue/PDF metadata use different publication-year "
    "conventions.\" (comment M12)")
P("We checked this against Crossref. The DOI 10.26833/ijeg.1655607 records online "
  "publication on 1 October 2025, and places the article in Volume 11, Issue 1, "
  "pages 149–162. The journal's volume sequence implies 2026 for Volume 11, since "
  "Bijaber et al. (Volume 9) is 2024. Both years are defensible; you need to pick one "
  "and be consistent. The manuscript currently cites 2026, matching the volume.")
P("The same pattern applies to Salih et al.: Crossref records 25 December 2025 online, "
  "in the issue dated April 2026. The manuscript cites 2026, matching the issue.",
  space_after=12)

# ---------------------------------------------------------------- closed
doc.add_heading("Already closed — no action needed", 1)
sens = pd.read_csv(os.path.join(RES, "11_extraction_support_sensitivity.csv"))
sagr = pd.read_csv(os.path.join(RES, "11_extraction_agreement.csv"))
doi = pd.read_csv(os.path.join(RES, "12_doi_verification.csv"))
P("Extraction-support sensitivity test", bold=True, space_after=3)
P(f"Every predictor was re-extracted as the mean of a 3 × 3 pixel neighbourhood and the "
  f"identical nested spatial cross-validation was repeated. The two extractions agree "
  f"closely (median Pearson r = {sagr.r.median():.3f}) and the effect on validated "
  f"performance is small and inconsistent in sign. The result is written into §3.3 and "
  f"Supplementary Table S6.")
s = sens.copy()
s.columns = ["Nutrient", "Extraction support", "R²", "RMSE", "MAE", "Bias", "CCC", "Slope"]
for c in ["R²", "CCC", "Slope"]:
    s[c] = s[c].map(lambda v: f"{v:.3f}")
for c in ["RMSE", "MAE", "Bias"]:
    s[c] = s[c].map(lambda v: f"{v:,.3f}")
table(s)

P("DOI verification", bold=True, space_after=3)
P(f"All {len(doi)} DOIs in the reference list were checked against Crossref and all "
  f"resolve. Three discrepancies were found and are handled in the manuscript. The most "
  f"serious was not raised by the referee: the Castaldi entry in the previous version "
  f"cited DOI 10.3390/rs11242924, which resolves to Monteith et al., a radar-tomography "
  f"paper in Remote Sensing 11(24). The intended article is Castaldi, F. (2021), Remote "
  f"Sensing 13, 3345, doi 10.3390/rs13173345 — a single author, and 2021 rather than "
  f"2019. Every field of that reference was wrong. It has been corrected in the "
  f"manuscript and the in-text citation changed to \"Castaldi, 2021\". The full check is "
  f"in Supplementary Table S7.")

P("English rewrite", bold=True, space_after=3)
P("The manuscript has been rewritten throughout. The specific phrases the referee "
  "quoted — \"very good at predicting things\", \"Potassium shows the biggest "
  "improvement in how clearly it can be seen\", and the §3.2.2 passage beginning \"We "
  "calculated various aspects of the plants…\" — are gone. A professional proofread "
  "before submission is still advisable, but the manuscript no longer contains the "
  "informal or ungrammatical constructions the referee identified.", space_after=14)

P("A closing word", bold=True, size=11.5, color=ACC, space_after=3)
P("Item 1 is the one that decides whether this round succeeds. A referee who found that "
  "your R² and RMSE were mutually inconsistent, that two predictors were duplicates, and "
  "that eleven citations did not support their claims will read the QA/QC section "
  "closely. Complete records are ideal; an honest statement of what was and was not "
  "measured is acceptable; invented numbers would end the submission.")

out = os.path.join(DOCS, "Data_You_Need_With_Examples.docx")
doc.save(out)
print("wrote", out)
print("STEP 13 complete.")
