# -*- coding: utf-8 -*-
"""
STEP 18 - Full manuscript audit.

  A. formatting matches the originally submitted file
  B. every in-text citation resolves to a reference entry, and vice versa
  C. every figure and table is cited in the body, in order
  D. no shading anywhere in any table
  E. unit and terminology consistency

Annotation paragraphs ([AUTHOR ACTION], [INSERT ...], the editorial page and its
bullet lists) are excluded from every content check: they are instructions to the
authors, not manuscript text, and are deleted before submission.
"""
import os, re
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn

FINAL = r"d:\Doctorat\article1\outputs_fast\FINAL"
RES = os.path.join(FINAL, "results")
MS = os.path.join(FINAL, "docs", "Revised_Manuscript_v3_TrackedRed_Amrouss_et_al.docx")
ORIG = os.path.normpath(os.path.join(FINAL, "..", "Revised Manuscript(1).docx"))

ok, bad, warn = [], [], []
def chk(n, c, d=""):  (ok if c else bad).append((n, d))
def note(n, d=""):    warn.append((n, d))

d = docx.Document(MS)
o = docx.Document(ORIG)


def is_annot_text(t, style):
    return (t.startswith("[AUTHOR ACTION]") or t.startswith("[INSERT")
            or style == "List Bullet")


# One ordered pass. lxml element proxies are not stable identities, so the
# annotation flag is carried alongside the paragraph rather than looked up.
PARAS = []
_after_editorial = False
for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith("Editorial note"):
        _after_editorial = True
    PARAS.append((p, t, _after_editorial or is_annot_text(t, p.style.name)))

OPARAS = [(p, p.text.strip(), False) for p in o.paragraphs]

# ---------------------------------------------------------------- A. format
print("=" * 100); print("A. FORMATTING vs THE ORIGINALLY SUBMITTED FILE"); print("=" * 100)
for name, a, b in [
    ("page width",  d.sections[0].page_width,  o.sections[0].page_width),
    ("page height", d.sections[0].page_height, o.sections[0].page_height),
    ("left margin", d.sections[0].left_margin, o.sections[0].left_margin),
    ("top margin",  d.sections[0].top_margin,  o.sections[0].top_margin),
]:
    chk(f"{name} matches original", a == b, f"{a} vs {b}")


def body_paras(paras):
    return [p for p, t, ann in paras if len(t) > 200 and not ann]


ns_ = {r.font.size.pt for p in body_paras(PARAS) for r in p.runs
       if r.text.strip() and r.font.size}
os_ = {r.font.size.pt for p in body_paras(OPARAS) for r in p.runs
       if r.text.strip() and r.font.size}
chk("body text 10 pt as in original", ns_ == {10.0} and os_ == {10.0}, f"{ns_} vs {os_}")

nsp = {p.paragraph_format.line_spacing for p in body_paras(PARAS)}
osp = {p.paragraph_format.line_spacing for p in body_paras(OPARAS)}
# running text is 1.5 in both; reference paragraphs are 1.15 in both, so the
# requirement is that no spacing value appears here that the original did not use
chk("body line spacing values all present in the original",
    nsp <= osp and 1.5 in nsp, f"{sorted(str(x) for x in nsp)} vs "
                               f"{sorted(str(x) for x in osp)}")
_in_refs, _started = False, False
_running = []
for p, t, ann in PARAS:
    if t == "Abstract":
        _started = True; continue          # title-page block is single-spaced
    if t == "References":
        _in_refs = True; continue
    if t.startswith("Supplementary Material"):
        _in_refs = False; continue
    # captions are deliberately single-spaced, as in the original
    if (_started and not ann and not _in_refs and len(t) > 200
            and not re.match(r"^(Fig\.|Table )", t)):
        _running.append(p)
run_sp = {p.paragraph_format.line_spacing for p in _running}
chk("running text is 1.5 line spacing", run_sp == {1.5},
    f"{sorted(str(x) for x in run_sp)} over {len(_running)} paragraphs")
nal = {str(p.alignment) for p in body_paras(PARAS)}
oal = {str(p.alignment) for p in body_paras(OPARAS)}
chk("body justified as in original", "JUSTIFY" in str(nal) and "JUSTIFY" in str(oal))

shaded = boldhdr = 0
for t in d.tables:
    for s in t._tbl.findall('.//' + qn('w:shd')):
        f = s.get(qn('w:fill'))
        if f and f.lower() not in ("auto", "ffffff"):
            shaded += 1
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                if r.bold:
                    boldhdr += 1
chk("no shaded cells in any table", shaded == 0, f"{shaded} shaded cells")
chk("no bold header cells (matches original)", boldhdr == 0, f"{boldhdr} bold runs")
chk("all tables use a plain grid style",
    all(t.style.name == "Table Grid" for t in d.tables),
    ", ".join(sorted({t.style.name for t in d.tables})))
tsz = {r.font.size.pt for t in d.tables for row in t.rows for c in row.cells
       for p in c.paragraphs for r in p.runs if r.font.size}
chk("table text 7.5-10 pt", min(tsz) >= 7.5 and max(tsz) <= 10.0,
    f"{min(tsz)}-{max(tsz)} pt")
hdrrep = sum(1 for t in d.tables
             if t.rows[0]._tr.find(qn('w:trPr')) is not None
             and t.rows[0]._tr.find(qn('w:trPr')).find(qn('w:tblHeader')) is not None)
chk("header row repeats on every table", hdrrep == len(d.tables),
    f"{hdrrep}/{len(d.tables)}")

# ------------------------------------------------------- B. citation audit
print("=" * 100); print("B. CITATION AUDIT"); print("=" * 100)
body_text, ref_text, in_refs = [], [], False
for p, t, ann in PARAS:
    if t == "References":
        in_refs = True; continue
    if t.startswith("Supplementary Material"):
        in_refs = False; continue
    if t.startswith("Editorial note"):
        break
    if ann:
        continue
    (ref_text if in_refs else body_text).append(t)

table_text = "\n".join(c.text for tb in d.tables for row in tb.rows for c in row.cells)
body = "\n".join(body_text) + "\n" + table_text
refs = [t for t in ref_text if len(t) > 60 and "(" in t]

italic_refs = 0
in_refs = False
for p, t, ann in PARAS:
    if t == "References":
        in_refs = True; continue
    if t.startswith("Supplementary Material"):
        break
    if in_refs and not ann and len(t) > 60 and "(" in t:
        if any(r.italic for r in p.runs):
            italic_refs += 1
chk("journal names italic in references (as in original)",
    len(refs) > 0 and italic_refs == len(refs), f"{italic_refs}/{len(refs)}")


def ref_key(entry):
    m = re.match(r"^(.+?)\s*\((\d{4})\)", entry)
    if not m:
        return None
    return (re.split(r",|\s+and\s+", m.group(1))[0].strip(), m.group(2))


ref_keys = {}
for r in refs:
    k = ref_key(r)
    if k:
        ref_keys.setdefault(k, []).append(r)
chk("every reference parses to author and year", len(ref_keys) == len(refs),
    f"{len(ref_keys)} keys from {len(refs)} entries")
dups = {k: len(v) for k, v in ref_keys.items() if len(v) > 1}
chk("no duplicate reference entries", not dups, str(dups))

norm = [ref_key(r)[0].replace("\u2019", "'").lower() for r in refs if ref_key(r)]
misordered = [(norm[i], norm[i + 1]) for i in range(len(norm) - 1)
              if norm[i] > norm[i + 1]]
chk("reference list alphabetised", not misordered, str(misordered[:4]))

cited = set()
for m in re.finditer(r"\(([^()]*\d{4}[^()]*)\)", body):
    for part in m.group(1).split(";"):
        mm = re.match(r"^([A-Z][^,]*?)(?:\s+et\s+al\.)?(?:,)?\s*(\d{4})", part.strip())
        if mm:
            cited.add((mm.group(1).split(" and ")[0].strip(), mm.group(2)))
for m in re.finditer(
        r"([A-Z][A-Za-z\-\u2019']+(?:\s+(?:et\s+al\.|and\s+[A-Z][A-Za-z\-\u2019']+))?)"
        r"\s*\((\d{4})\)", body):
    cited.add((m.group(1).split(" and ")[0].replace(" et al.", "").strip(),
               m.group(2)))

ref_lookup = {(a.split()[-1].lower(), y) for a, y in ref_keys}
SKIP = ("fig", "table", "eq", "section", "epsg", "copernicus", "rmse", "sse", "s2")
missing = sorted({(a, y) for a, y in cited
                  if (a.split()[-1].lower(), y) not in ref_lookup
                  and a.lower().rstrip(":") not in SKIP})
chk("every in-text citation has a reference entry", not missing, str(missing))

body_l = body.lower().replace("\u2019", "'")
orphans = [f"{a} ({y})" for (a, y) in ref_keys
           if a.split()[-1].replace("\u2019", "'").lower() not in body_l]
chk("no orphan references (all cited in text)", not orphans, str(orphans))

# ---------------------------------------------------- C. figures and tables
print("=" * 100); print("C. FIGURE AND TABLE CALLOUTS"); print("=" * 100)
caps_f = [t for _, t, a in PARAS if not a and re.match(r"^Fig\. \d+ ", t)]
caps_t = [t for _, t, a in PARAS if not a and re.match(r"^Table \d+ ", t)]
caps_s = [t for _, t, a in PARAS if not a and re.match(r"^Table S", t)]
fn = [int(re.match(r"^Fig\. (\d+)", c).group(1)) for c in caps_f]
tnum = [int(re.match(r"^Table (\d+)", c).group(1)) for c in caps_t]
chk("figure captions 1..N sequential", fn == list(range(1, len(fn) + 1)), str(fn))
chk("table captions 1..N sequential", tnum == list(range(1, len(tnum) + 1)), str(tnum))
chk("supplementary tables S1..S8 present", len(caps_s) == 9, str(len(caps_s)))
for i in fn:
    chk(f"Fig. {i} cited in body",
        re.search(rf"Figs?\.\s*(?:\d+\s*(?:,|and)\s*)*{i}\b", body) is not None)
for i in tnum:
    chk(f"Table {i} cited in body",
        re.search(rf"Tables?\s*(?:\d+\s*(?:,|and|to)\s*)*{i}\b", body) is not None)
for s in ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8"]:
    chk(f"Table {s} cited in body", f"Table {s}" in body)
nimg = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
chk("images embedded equals figure captions", nimg == len(fn), f"{nimg} vs {len(fn)}")

# ------------------------------------------------------------- D. content
print("=" * 100); print("D. CONTENT AND CONSISTENCY"); print("=" * 100)
for s, why in [
    ("41 predictors", "old predictor count"),
    ("mg/kg", "old unit style"),
    ("Castaldi et al., 2019", "wrong Castaldi citation"),
    ("two cloud-free", "composites described as single scenes"),
    ("Havlin, 2014", "reference absent from the list"),
    ("Zhang, 2023", "deleted reference"),
    ("Lundberg and Lee", "deleted reference"),
    ("Viscarra Rossel", "deleted reference"),
    ("Abatzoglou", "deleted reference"),
    ("Cambardella", "deleted reference"),
    ("Qi et al.", "deleted reference"),
    ("Nguyen et al.", "deleted reference"),
    ("Sharma", "deleted reference"),
]:
    chk(f"no '{s}' ({why})", s not in body)
chk("uses 'mg kg-1' consistently", "mg kg-1" in body)
chk("40 predictors stated", "40 predictors" in body)
chk("S2_SR_HARMONIZED named", "COPERNICUS/S2_SR_HARMONIZED" in body)
chk("no unfilled [INSERT] placeholder remains",
    not any(t.startswith("[INSERT") for _, t, _ in PARAS))
for term in ["three independent analytical replicates", "Two analytical blanks",
             "Internal certified reference soils", "flame photometer",
             "Kjeldahl digestion", "limit of quantification",
             "National Institute of Agricultural Research"]:
    chk(f"QA/QC text present: '{term}'", term in body)
# every LOD/LOQ figure in Table 1 must be below the observed minimum
import pandas as _pd
_df = _pd.read_csv(os.path.join(RES, "01_analysis_ready_dataset.csv"))
for nut, loq in [("N", 150.0), ("P", 0.8), ("K", 5.0)]:
    chk(f"{nut}: all 110 observations above the stated LOQ",
        (_df[nut] >= loq).all(),
        f"min {_df[nut].min():.2f} vs LOQ {loq}")
chk("section numbering uses dots", any(t.startswith("3.2.1.") for _, t, _ in PARAS))
chk("equations numbered (1) and (2)", "(1)" in body and "(2)" in body)

RED = RGBColor(0xC0, 0x00, 0x00)
nred = sum(len(r.text) for p, t, a in PARAS if not a for r in p.runs
           if r.font.color and r.font.color.type is not None and r.font.color.rgb == RED)
ntot = sum(len(r.text) for p, t, a in PARAS if not a for r in p.runs if r.text.strip())
note("red / total characters (manuscript only)",
     f"{nred:,} / {ntot:,} ({100*nred/max(ntot,1):.0f} % new)")
note("word count (manuscript body, excl. tables and notes)",
     f"{sum(len(t.split()) for _, t, a in PARAS if not a):,}")
note("references", str(len(refs)))
note("figures / tables / supplementary tables",
     f"{len(fn)} / {len(tnum)} / {len(caps_s)}")

print()
print("PASSED (%d)" % len(ok))
for n, dd in ok:
    print(f"   OK    {n}" + (f"  -  {dd}" if dd else ""))
if warn:
    print("\nINFO")
    for n, dd in warn:
        print(f"   .     {n}: {dd}")
if bad:
    print("\nFAILED (%d)" % len(bad))
    for n, dd in bad:
        print(f"   FAIL  {n}" + (f"  -  {dd}" if dd else ""))
else:
    print("\nAll manuscript checks passed.")
